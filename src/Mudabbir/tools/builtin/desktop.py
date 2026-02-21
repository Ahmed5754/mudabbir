# pyright: reportMissingImports=false, reportMissingModuleSource=false
"""Desktop automation and system control helpers for Windows."""

from __future__ import annotations

import hashlib
import ipaddress
import json
import logging
import os
import base64
import re
import subprocess
import time
import wave
from datetime import datetime
from pathlib import Path
from typing import Any

from Mudabbir.bus.media import get_media_dir
from Mudabbir.config import get_settings
from Mudabbir.tools.builtin.desktop_audio import AUDIO_ACTIONS
from Mudabbir.tools.builtin.desktop_display import DISPLAY_ACTIONS
from Mudabbir.tools.builtin.desktop_files import FILE_ACTIONS
from Mudabbir.tools.builtin.desktop_network import NETWORK_ACTIONS
from Mudabbir.tools.builtin.desktop_process import PROCESS_ACTIONS
from Mudabbir.tools.builtin.desktop_registry import REGISTRY_ACTIONS
from Mudabbir.tools.builtin.desktop_security import SECURITY_ACTIONS
from Mudabbir.tools.builtin.desktop_system import SYSTEM_ACTIONS
from Mudabbir.tools.protocol import BaseTool


class _NoisyDisplayWarningFilter(logging.Filter):
    """Drop known noisy EDID warnings from third-party brightness discovery."""

    _patterns = (
        "exception parsing edid str",
        "edidparseerror",
        "detected but not present in monitor_uids",
        "maybe it is asleep",
    )

    def filter(self, record: logging.LogRecord) -> bool:
        try:
            message = str(record.getMessage() or "").lower()
        except Exception:
            message = ""
        if any(pattern in message for pattern in self._patterns):
            return False
        return True


_DISPLAY_WARNING_FILTER_INSTALLED = False


def _install_display_warning_filter_once() -> None:
    global _DISPLAY_WARNING_FILTER_INSTALLED
    if _DISPLAY_WARNING_FILTER_INSTALLED:
        return
    filt = _NoisyDisplayWarningFilter()
    for logger_name in (
        "screen_brightness_control",
        "screen_brightness_control.windows",
        "screen_brightness_control.helpers",
    ):
        logging.getLogger(logger_name).addFilter(filt)
    logging.getLogger().addFilter(filt)
    _DISPLAY_WARNING_FILTER_INSTALLED = True


def _json(data: dict[str, Any] | list[Any]) -> str:
    return json.dumps(data, ensure_ascii=False)


def _clamp(value: int | float, minimum: int, maximum: int) -> int:
    return max(minimum, min(maximum, int(value)))


def _timestamp_id() -> str:
    return time.strftime("%Y%m%d_%H%M%S")


def _filetime_to_iso(value: Any) -> str | None:
    """Convert Windows FILETIME ticks to ISO timestamp (local time)."""
    try:
        raw = int(str(value or "").strip())
    except Exception:
        return None
    if raw <= 0:
        return None
    try:
        unix_ts = (raw / 10_000_000.0) - 11644473600.0
        if unix_ts <= 0:
            return None
        return datetime.fromtimestamp(unix_ts).isoformat()
    except Exception:
        return None


def _run_powershell(command: str, timeout: int = 15) -> tuple[bool, str]:
    try:
        proc = subprocess.run(
            ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", command],
            capture_output=True,
            text=True,
            timeout=timeout,
        )
    except Exception as e:
        return False, str(e)

    out = (proc.stdout or "").strip()
    err = (proc.stderr or "").strip()
    if proc.returncode == 0:
        return True, out
    if err:
        return False, err
    if out:
        return False, out
    return False, f"PowerShell exited with code {proc.returncode}"


def _iter_search_roots() -> list[Path]:
    home = Path.home()
    roots = [
        home / "Desktop",
        home / "Documents",
        home / "Downloads",
        home / "Pictures",
        home / "Videos",
    ]
    existing = [p for p in roots if p.exists() and p.is_dir()]
    return existing or [home]


def _contains_text(haystack: str, needle: str) -> bool:
    return needle.casefold() in haystack.casefold()


def _contains_arabic(text: str) -> bool:
    return bool(re.search(r"[\u0600-\u06FF]", text or ""))


def _normalize_query(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip())


def _serialize_windows(
    *,
    include_untitled: bool = True,
    limit: int = 200,
) -> list[dict[str, Any]]:
    windows: list[dict[str, Any]] = []
    try:
        import psutil
        import win32gui
        import win32process

        def enum_handler(hwnd: int, _: Any) -> None:
            if len(windows) >= limit:
                return
            if not win32gui.IsWindowVisible(hwnd):
                return

            title = (win32gui.GetWindowText(hwnd) or "").strip()
            if not include_untitled and not title:
                return

            left, top, right, bottom = win32gui.GetWindowRect(hwnd)
            width = max(0, right - left)
            height = max(0, bottom - top)
            if width < 40 or height < 40:
                return

            _, pid = win32process.GetWindowThreadProcessId(hwnd)
            process_name = ""
            try:
                process_name = psutil.Process(pid).name()
            except Exception:
                process_name = ""

            windows.append(
                {
                    "hwnd": int(hwnd),
                    "pid": int(pid),
                    "process_name": process_name,
                    "title": title,
                    "left": int(left),
                    "top": int(top),
                    "width": int(width),
                    "height": int(height),
                }
            )

        win32gui.EnumWindows(enum_handler, None)
    except Exception:
        try:
            import pygetwindow as gw

            for win in gw.getAllWindows():
                if len(windows) >= limit:
                    break
                title = str(getattr(win, "title", "") or "").strip()
                if not include_untitled and not title:
                    continue
                width = int(getattr(win, "width", 0) or 0)
                height = int(getattr(win, "height", 0) or 0)
                if width < 40 or height < 40:
                    continue
                windows.append(
                    {
                        "hwnd": 0,
                        "pid": 0,
                        "process_name": "",
                        "title": title,
                        "left": int(getattr(win, "left", 0) or 0),
                        "top": int(getattr(win, "top", 0) or 0),
                        "width": width,
                        "height": height,
                    }
                )
        except Exception:
            return []

    windows.sort(key=lambda item: (item.get("title") or "").casefold())
    return windows[:limit]


class DesktopTool(BaseTool):
    BLOCKED_AUTOMATION_PROCESSES = {
        "keepass.exe",
        "1password.exe",
        "lastpass.exe",
        "bitwarden.exe",
        "authy desktop.exe",
        "secpol.msc",
        "gpedit.msc",
        "regedit.exe",
    }
    BLOCKED_HOTKEY_COMBINATIONS = {
        frozenset({"ctrl", "alt", "delete"}),
        frozenset({"win", "r"}),
        frozenset({"win", "x"}),
    }
    BLOCKED_WINDOW_CLASSES = {
        "#32770",
        "credential dialog",
        "credential dialog xaml host",
        "windows security",
    }
    ACTION_GROUPS = {
        "system": SYSTEM_ACTIONS,
        "audio": AUDIO_ACTIONS,
        "display": DISPLAY_ACTIONS,
        "files": FILE_ACTIONS,
        "network": NETWORK_ACTIONS,
        "process": PROCESS_ACTIONS,
        "registry": REGISTRY_ACTIONS,
        "security": SECURITY_ACTIONS,
    }

    @property
    def name(self) -> str:
        return "desktop"

    @property
    def description(self) -> str:
        return "Desktop automation, app control, media capture, and Windows system actions."

    @property
    def trust_level(self) -> str:
        return "high"

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "action": {"type": "string"},
                "mode": {"type": "string"},
                "query": {"type": "string"},
                "process_name": {"type": "string"},
                "x": {"type": "number"},
                "y": {"type": "number"},
                "x2": {"type": "number"},
                "y2": {"type": "number"},
                "button": {"type": "string"},
                "clicks": {"type": "integer"},
                "keys": {"type": "array", "items": {"type": "string"}},
                "key": {"type": "string"},
                "level": {"type": "integer"},
                "delta": {"type": "integer"},
                "max_results": {"type": "integer"},
                "seconds": {"type": "number"},
                "frames": {"type": "integer"},
                "interval_sec": {"type": "number"},
                "camera_index": {"type": "integer"},
                "only_windowed": {"type": "boolean"},
                "include_untitled": {"type": "boolean"},
                "limit": {"type": "integer"},
                "page": {"type": "string"},
                "force": {"type": "boolean"},
                "window_title": {"type": "string"},
                "control_name": {"type": "string"},
                "auto_id": {"type": "string"},
                "control_type": {"type": "string"},
                "text": {"type": "string"},
                "index": {"type": "integer"},
                "press_enter": {"type": "boolean"},
                "timeout_sec": {"type": "number"},
                "duration": {"type": "number"},
                "interval": {"type": "number"},
                "interaction": {"type": "string"},
                "path": {"type": "string"},
                "target": {"type": "string"},
                "name": {"type": "string"},
                "pattern": {"type": "string"},
                "ext": {"type": "string"},
                "permanent": {"type": "boolean"},
                "host": {"type": "string"},
                "username": {"type": "string"},
                "password": {"type": "string"},
                "group": {"type": "string"},
                "command": {"type": "string"},
                "trigger": {"type": "string"},
                "value_name": {"type": "string"},
                "value_data": {"type": "string"},
                "value_type": {"type": "string"},
                "drive": {"type": "string"},
                "port": {"type": "integer"},
                "rule_name": {"type": "string"},
                "url": {"type": "string"},
                "city": {"type": "string"},
                "minutes": {"type": "integer"},
                "pid": {"type": "integer"},
                "app": {"type": "string"},
                "repeat_count": {"type": "integer"},
                "timezone": {"type": "string"},
                "folder": {"type": "string"},
                "size_mb": {"type": "number"},
                "urls": {"type": "array", "items": {"type": "string"}},
                "date": {"type": "string"},
                "width": {"type": "integer"},
                "height": {"type": "integer"},
                "content": {"type": "string"},
                "replace_with": {"type": "string"},
                "priority": {"type": "string"},
                "opacity": {"type": "integer"},
                "threshold": {"type": "number"},
                "monitor_seconds": {"type": "number"},
                "editor": {"type": "string"},
                "notify": {"type": "boolean"},
            },
            "required": ["action"],
        }

    async def execute(self, action: str, **params: Any) -> str:
        action_normalized = (action or "").strip().lower()
        if not action_normalized:
            return self._error("Missing action")

        try:
            if action_normalized == "screen_snapshot":
                return self._screen_snapshot()
            if action_normalized == "screen_watch":
                return self._screen_watch(
                    frames=int(params.get("frames", 3) or 3),
                    interval_sec=float(params.get("interval_sec", 0.7) or 0.7),
                )
            if action_normalized == "desktop_overview":
                return self._desktop_overview()
            if action_normalized == "battery_status":
                return self._battery_status()
            if action_normalized == "list_windows":
                return self._list_windows(
                    include_untitled=bool(params.get("include_untitled", True)),
                    limit=int(params.get("limit", 120) or 120),
                )
            if action_normalized == "focus_window":
                return self._focus_window(
                    window_title=str(params.get("window_title", "") or params.get("query", "") or ""),
                    process_name=str(params.get("process_name", "") or ""),
                    timeout_sec=float(params.get("timeout_sec", 4.0) or 4.0),
                )
            if action_normalized == "ui_list_controls":
                return self._ui_list_controls(
                    window_title=str(params.get("window_title", "") or params.get("query", "") or ""),
                    process_name=str(params.get("process_name", "") or ""),
                    control_name=str(params.get("control_name", "") or params.get("query", "") or ""),
                    control_type=str(params.get("control_type", "") or ""),
                    max_results=int(params.get("max_results", 40) or 40),
                    timeout_sec=float(params.get("timeout_sec", 4.0) or 4.0),
                )
            if action_normalized == "ui_click":
                return self._ui_click(
                    window_title=str(params.get("window_title", "") or params.get("query", "") or ""),
                    process_name=str(params.get("process_name", "") or ""),
                    control_name=str(params.get("control_name", "") or ""),
                    auto_id=str(params.get("auto_id", "") or ""),
                    control_type=str(params.get("control_type", "") or ""),
                    index=int(params.get("index", 0) or 0),
                    timeout_sec=float(params.get("timeout_sec", 4.0) or 4.0),
                )
            if action_normalized == "ui_set_text":
                return self._ui_set_text(
                    text=str(params.get("text", "") or ""),
                    window_title=str(params.get("window_title", "") or params.get("query", "") or ""),
                    process_name=str(params.get("process_name", "") or ""),
                    control_name=str(params.get("control_name", "") or ""),
                    auto_id=str(params.get("auto_id", "") or ""),
                    control_type=str(params.get("control_type", "") or ""),
                    index=int(params.get("index", 0) or 0),
                    press_enter=bool(params.get("press_enter", False)),
                    timeout_sec=float(params.get("timeout_sec", 4.0) or 4.0),
                )
            if action_normalized in {"ui_target", "ui_move_to"}:
                interaction = str(params.get("interaction", "") or "").strip().lower()
                if action_normalized == "ui_move_to" and not interaction:
                    interaction = "move"
                return self._ui_target(
                    window_title=str(params.get("window_title", "") or params.get("query", "") or ""),
                    process_name=str(params.get("process_name", "") or ""),
                    control_name=str(params.get("control_name", "") or params.get("query", "") or ""),
                    auto_id=str(params.get("auto_id", "") or ""),
                    control_type=str(params.get("control_type", "") or ""),
                    index=int(params.get("index", 0) or 0),
                    interaction=interaction or "move",
                    duration=float(params.get("duration", 0.25) or 0.25),
                    timeout_sec=float(params.get("timeout_sec", 5.0) or 5.0),
                )
            if action_normalized == "list_processes":
                return self._list_processes(
                    only_windowed=bool(params.get("only_windowed", False)),
                    max_results=int(params.get("max_results", 25) or 25),
                )
            if action_normalized == "close_app":
                return self._close_app(
                    process_name=str(params.get("process_name", "") or ""),
                    force=bool(params.get("force", True)),
                )
            if action_normalized == "search_start_apps":
                return self._search_start_apps(
                    query=str(params.get("query", "") or ""),
                    max_results=int(params.get("max_results", 10) or 10),
                )
            if action_normalized == "launch_start_app":
                return self._launch_start_app(query=str(params.get("query", "") or ""))
            if action_normalized == "search_files":
                return self._search_files(
                    query=str(params.get("query", "") or ""),
                    max_results=int(params.get("max_results", 20) or 20),
                )
            if action_normalized == "move_mouse_to_desktop_file":
                return self._move_mouse_to_desktop_file(
                    query=str(params.get("query", "") or ""),
                    duration=float(params.get("duration", 0.25) or 0.25),
                    timeout_sec=float(params.get("timeout_sec", 4.0) or 4.0),
                )
            if action_normalized == "list_installed_apps":
                return self._list_installed_apps(
                    query=str(params.get("query", "") or ""),
                    max_results=int(params.get("max_results", 25) or 25),
                )
            if action_normalized == "volume":
                return self._volume_control(
                    mode=str(params.get("mode", "get") or "get"),
                    level=params.get("level"),
                    delta=params.get("delta"),
                )
            if action_normalized == "brightness":
                return self._brightness_control(
                    mode=str(params.get("mode", "get") or "get"),
                    level=params.get("level"),
                    delta=params.get("delta"),
                )
            if action_normalized == "media_control":
                return self._media_control(mode=str(params.get("mode", "play_pause") or "play_pause"))
            if action_normalized == "mouse_move":
                return self._mouse_move(
                    x=int(params.get("x", 0) or 0),
                    y=int(params.get("y", 0) or 0),
                    duration=float(params.get("duration", 0.2) or 0.2),
                )
            if action_normalized == "click":
                return self._click(
                    x=params.get("x"),
                    y=params.get("y"),
                    button=str(params.get("button", "left") or "left"),
                    clicks=int(params.get("clicks", 1) or 1),
                )
            if action_normalized == "press_key":
                return self._press_key(key=str(params.get("key", "") or ""))
            if action_normalized == "type_text":
                return self._type_text(
                    text=str(params.get("text", "") or ""),
                    press_enter=bool(params.get("press_enter", False)),
                    interval=float(params.get("interval", 0.01) or 0.01),
                )
            if action_normalized == "hotkey":
                keys = params.get("keys") or []
                keys = [str(k) for k in keys if str(k).strip()]
                return self._hotkey(keys=keys)
            if action_normalized == "camera_snapshot":
                return self._camera_snapshot(camera_index=int(params.get("camera_index", 0) or 0))
            if action_normalized == "microphone_record":
                return self._microphone_record(seconds=float(params.get("seconds", 3.0) or 3.0))
            if action_normalized == "open_settings_page":
                return self._open_settings_page(page=str(params.get("page", "") or ""))
            if action_normalized == "bluetooth_control":
                return self._bluetooth_control(mode=str(params.get("mode", "open_settings") or "open_settings"))
            if action_normalized == "system_power":
                return self._system_power(
                    mode=str(params.get("mode", "lock") or "lock"),
                    name=str(params.get("name", "") or ""),
                )
            if action_normalized == "shutdown_schedule":
                return self._shutdown_schedule(
                    mode=str(params.get("mode", "set") or "set"),
                    minutes=params.get("minutes"),
                )
            if action_normalized == "system_info":
                return self._system_info(mode=str(params.get("mode", "windows_version") or "windows_version"))
            if action_normalized == "network_tools":
                return self._network_tools(
                    mode=str(params.get("mode", "ip_internal") or "ip_internal"),
                    host=str(params.get("host", "") or ""),
                    port=params.get("port"),
                )
            if action_normalized == "file_tools":
                return self._file_tools(
                    mode=str(params.get("mode", "open_documents") or "open_documents"),
                    path=str(params.get("path", "") or ""),
                    target=str(params.get("target", "") or ""),
                    name=str(params.get("name", "") or ""),
                    pattern=str(params.get("pattern", "") or ""),
                    ext=str(params.get("ext", "") or ""),
                    permanent=bool(params.get("permanent", False)),
                )
            if action_normalized == "window_control":
                return self._window_control(
                    mode=str(params.get("mode", "show_desktop") or "show_desktop"),
                    app=str(params.get("app", "") or params.get("query", "") or ""),
                    x=params.get("x"),
                    y=params.get("y"),
                    width=params.get("width"),
                    height=params.get("height"),
                    opacity=params.get("opacity"),
                    name=str(params.get("name", "") or ""),
                    text=str(params.get("text", "") or ""),
                )
            if action_normalized == "process_tools":
                return self._process_tools(
                    mode=str(params.get("mode", "list") or "list"),
                    pid=params.get("pid"),
                    name=str(params.get("name", "") or params.get("process_name", "") or ""),
                    other_name=str(params.get("other_name", "") or params.get("target", "") or ""),
                    dry_run=bool(params.get("dry_run", False)),
                    max_kill=params.get("max_kill"),
                    resource=str(params.get("resource", "") or ""),
                    stage=str(params.get("stage", "") or ""),
                    priority=str(params.get("priority", "") or ""),
                    threshold=params.get("threshold"),
                )
            if action_normalized == "service_tools":
                return self._service_tools(
                    mode=str(params.get("mode", "list") or "list"),
                    name=str(params.get("name", "") or ""),
                    startup=str(params.get("startup", "") or ""),
                )
            if action_normalized == "background_tools":
                return self._background_tools(
                    mode=str(params.get("mode", "count_background") or "count_background"),
                    max_results=int(params.get("max_results", 50) or 50),
                )
            if action_normalized == "startup_tools":
                return self._startup_tools(
                    mode=str(params.get("mode", "list") or "list"),
                    name=str(params.get("name", "") or ""),
                    seconds=params.get("seconds"),
                    monitor_seconds=params.get("monitor_seconds"),
                    notify=bool(params.get("notify", False)),
                )
            if action_normalized == "clipboard_tools":
                return self._clipboard_tools(mode=str(params.get("mode", "clear") or "clear"))
            if action_normalized == "browser_control":
                return self._browser_control(mode=str(params.get("mode", "new_tab") or "new_tab"))
            if action_normalized == "user_tools":
                return self._user_tools(
                    mode=str(params.get("mode", "list") or "list"),
                    username=str(params.get("username", "") or ""),
                    password=str(params.get("password", "") or ""),
                    group=str(params.get("group", "") or ""),
                )
            if action_normalized == "task_tools":
                return self._task_tools(
                    mode=str(params.get("mode", "list") or "list"),
                    name=str(params.get("name", "") or ""),
                    command=str(params.get("command", "") or ""),
                    trigger=str(params.get("trigger", "") or ""),
                )
            if action_normalized == "registry_tools":
                return self._registry_tools(
                    mode=str(params.get("mode", "query") or "query"),
                    key=str(params.get("key", "") or ""),
                    value_name=str(params.get("value_name", "") or ""),
                    value_data=str(params.get("value_data", "") or ""),
                    value_type=str(params.get("value_type", "REG_SZ") or "REG_SZ"),
                )
            if action_normalized == "disk_tools":
                return self._disk_tools(
                    mode=str(params.get("mode", "smart_status") or "smart_status"),
                    drive=str(params.get("drive", "") or ""),
                )
            if action_normalized == "security_tools":
                return self._security_tools(
                    mode=str(params.get("mode", "firewall_status") or "firewall_status"),
                    target=str(params.get("target", "") or ""),
                    port=params.get("port"),
                    rule_name=str(params.get("rule_name", "") or ""),
                )
            if action_normalized == "web_tools":
                return self._web_tools(
                    mode=str(params.get("mode", "open_url") or "open_url"),
                    url=str(params.get("url", "") or ""),
                    city=str(params.get("city", "") or ""),
                )
            if action_normalized == "hardware_tools":
                return self._hardware_tools(
                    mode=str(params.get("mode", "cpu_info") or "cpu_info"),
                    drive=str(params.get("drive", "") or ""),
                )
            if action_normalized == "update_tools":
                return self._update_tools(
                    mode=str(params.get("mode", "list_updates") or "list_updates"),
                    target=str(params.get("target", "") or ""),
                )
            if action_normalized == "ui_tools":
                return self._ui_tools(mode=str(params.get("mode", "dark_mode") or "dark_mode"))
            if action_normalized == "automation_tools":
                return self._automation_tools(
                    mode=str(params.get("mode", "delay") or "delay"),
                    seconds=params.get("seconds"),
                    monitor_seconds=params.get("monitor_seconds"),
                    text=str(params.get("text", "") or ""),
                    key=str(params.get("key", "") or ""),
                    repeat_count=params.get("repeat_count"),
                    x=params.get("x"),
                    y=params.get("y"),
                    x2=params.get("x2"),
                    y2=params.get("y2"),
                    width=params.get("width"),
                    height=params.get("height"),
                )
            if action_normalized == "app_tools":
                return self._app_tools(
                    mode=str(params.get("mode", "open_default_browser") or "open_default_browser"),
                    app=str(params.get("app", "") or ""),
                    dry_run=bool(params.get("dry_run", False)),
                    max_kill=params.get("max_kill"),
                )
            if action_normalized == "info_tools":
                return self._info_tools(
                    mode=str(params.get("mode", "timezone_get") or "timezone_get"),
                    timezone=str(params.get("timezone", "") or ""),
                )
            if action_normalized == "dev_tools":
                return self._dev_tools(
                    mode=str(params.get("mode", "open_cmd_admin") or "open_cmd_admin"),
                    drive=str(params.get("drive", "") or ""),
                    path=str(params.get("path", "") or ""),
                    editor=str(params.get("editor", "") or ""),
                    max_results=int(params.get("max_results", 20) or 20),
                    target=str(params.get("target", "") or ""),
                    text=str(params.get("text", "") or ""),
                    execute=bool(params.get("force", True)),
                )
            if action_normalized == "shell_tools":
                return self._shell_tools(mode=str(params.get("mode", "quick_settings") or "quick_settings"))
            if action_normalized == "office_tools":
                return self._office_tools(
                    mode=str(params.get("mode", "open_word_new") or "open_word_new"),
                    path=str(params.get("path", "") or ""),
                    target=str(params.get("target", "") or ""),
                )
            if action_normalized == "remote_tools":
                return self._remote_tools(
                    mode=str(params.get("mode", "rdp_open") or "rdp_open"),
                    host=str(params.get("host", "") or ""),
                )
            if action_normalized == "search_tools":
                return self._search_tools(
                    mode=str(params.get("mode", "search_text") or "search_text"),
                    folder=str(params.get("folder", "") or ""),
                    pattern=str(params.get("pattern", "") or ""),
                    ext=str(params.get("ext", "") or ""),
                    size_mb=params.get("size_mb"),
                )
            if action_normalized == "performance_tools":
                return self._performance_tools(
                    mode=str(params.get("mode", "top_cpu") or "top_cpu"),
                    threshold=params.get("threshold"),
                )
            if action_normalized == "media_tools":
                return self._media_tools(
                    mode=str(params.get("mode", "stop_all_media") or "stop_all_media"),
                    url=str(params.get("url", "") or ""),
                )
            if action_normalized == "browser_deep_tools":
                return self._browser_deep_tools(
                    mode=str(params.get("mode", "multi_open") or "multi_open"),
                    urls=params.get("urls"),
                )
            if action_normalized == "maintenance_tools":
                return self._maintenance_tools(mode=str(params.get("mode", "empty_ram") or "empty_ram"))
            if action_normalized == "driver_tools":
                return self._driver_tools(mode=str(params.get("mode", "drivers_list") or "drivers_list"))
            if action_normalized == "power_user_tools":
                return self._power_user_tools(mode=str(params.get("mode", "airplane_on") or "airplane_on"))
            if action_normalized == "screenshot_tools":
                return self._screenshot_tools(
                    mode=str(params.get("mode", "full") or "full"),
                    x=params.get("x"),
                    y=params.get("y"),
                    width=params.get("width"),
                    height=params.get("height"),
                    path=str(params.get("path", "") or ""),
                )
            if action_normalized == "text_tools":
                return self._text_tools(
                    mode=str(params.get("mode", "text_to_file") or "text_to_file"),
                    path=str(params.get("path", "") or ""),
                    content=str(params.get("content", "") or ""),
                    folder=str(params.get("folder", "") or ""),
                    pattern=str(params.get("pattern", "") or ""),
                    replace_with=str(params.get("replace_with", "") or ""),
                )
            if action_normalized == "api_tools":
                return self._api_tools(
                    mode=str(params.get("mode", "currency") or "currency"),
                    target=str(params.get("target", "") or ""),
                    city=str(params.get("city", "") or ""),
                    text=str(params.get("text", "") or ""),
                )
            if action_normalized == "vision_tools":
                return self._vision_tools(
                    mode=str(params.get("mode", "describe_screen") or "describe_screen"),
                    path=str(params.get("path", "") or ""),
                    x=params.get("x"),
                    y=params.get("y"),
                    width=params.get("width"),
                    height=params.get("height"),
                )
            if action_normalized == "threat_tools":
                return self._threat_tools(
                    mode=str(params.get("mode", "suspicious_connections") or "suspicious_connections"),
                    path=str(params.get("path", "") or ""),
                    target=str(params.get("target", "") or ""),
                    max_results=int(params.get("max_results", 50) or 50),
                )
            if action_normalized == "content_tools":
                return self._content_tools(
                    mode=str(params.get("mode", "draft_reply") or "draft_reply"),
                    content=str(params.get("content", "") or ""),
                    path=str(params.get("path", "") or ""),
                    target=str(params.get("target", "") or ""),
                )
        except Exception as e:
            return self._error(str(e))

        return self._error(f"Unknown desktop action: {action_normalized}")

    def _screen_snapshot(self) -> str:
        try:
            import pyautogui

            pyautogui.FAILSAFE = False
            media_dir = get_media_dir()
            path = media_dir / f"screen_{_timestamp_id()}.png"
            screenshot = pyautogui.screenshot()
            screenshot.save(path)
            return f"Screenshot saved to {path}"
        except Exception as e:
            return self._error(f"screenshot failed: {e}")

    def _screen_watch(self, frames: int, interval_sec: float) -> str:
        try:
            import pyautogui

            pyautogui.FAILSAFE = False
            frames = _clamp(frames, 1, 12)
            interval = max(0.05, min(5.0, float(interval_sec)))
            media_dir = get_media_dir()
            paths: list[str] = []
            for idx in range(frames):
                shot = pyautogui.screenshot()
                path = media_dir / f"watch_{_timestamp_id()}_{idx + 1}.png"
                shot.save(path)
                paths.append(str(path))
                if idx < frames - 1:
                    time.sleep(interval)
            return _json(
                {
                    "captured_frames": len(paths),
                    "interval_sec": interval,
                    "paths": paths,
                }
            )
        except Exception as e:
            return self._error(f"screen watch failed: {e}")

    def _desktop_overview(self) -> str:
        try:
            import pyautogui

            pyautogui.FAILSAFE = False
            size = pyautogui.size()
            pos = pyautogui.position()
            windows = _serialize_windows(include_untitled=False, limit=40)
            return _json(
                {
                    "screen": {"width": int(size.width), "height": int(size.height)},
                    "mouse": {"x": int(pos.x), "y": int(pos.y)},
                    "windows": {"count": len(windows), "items": windows[:20]},
                    "timestamp": time.time(),
                }
            )
        except Exception as e:
            return self._error(f"desktop overview failed: {e}")

    def _battery_status(self) -> str:
        """Return battery status if available."""
        try:
            import psutil

            battery = psutil.sensors_battery()
            if battery is not None:
                secs_left_raw = getattr(battery, "secsleft", None)
                secs_left: int | None
                if secs_left_raw is None:
                    secs_left = None
                else:
                    try:
                        value = int(secs_left_raw)
                    except Exception:
                        value = -2
                    if value < 0 or value in {psutil.POWER_TIME_UNKNOWN, psutil.POWER_TIME_UNLIMITED}:
                        secs_left = None
                    else:
                        secs_left = max(0, value)

                return _json(
                    {
                        "available": True,
                        "percent": max(0.0, min(100.0, float(getattr(battery, "percent", 0.0) or 0.0))),
                        "plugged": bool(getattr(battery, "power_plugged", False)),
                        "secs_left": secs_left,
                        "source": "psutil",
                    }
                )
        except Exception:
            pass

        # Fallback for systems where psutil battery API is unavailable.
        cmd = (
            "Get-CimInstance Win32_Battery | Select-Object -First 1 "
            "EstimatedChargeRemaining,BatteryStatus,Name | ConvertTo-Json -Compress"
        )
        ok, out = _run_powershell(cmd, timeout=12)
        if ok and out and out.strip() not in {"null", ""}:
            try:
                data = json.loads(out)
                if isinstance(data, dict):
                    percent_raw = data.get("EstimatedChargeRemaining")
                    percent = None
                    try:
                        if percent_raw is not None:
                            percent = max(0.0, min(100.0, float(percent_raw)))
                    except Exception:
                        percent = None
                    status = int(data.get("BatteryStatus", 0) or 0)
                    plugged = status in {2, 6, 7, 8, 9, 11}
                    return _json(
                        {
                            "available": percent is not None,
                            "percent": percent,
                            "plugged": plugged,
                            "secs_left": None,
                            "status_code": status,
                            "name": str(data.get("Name", "") or "").strip(),
                            "source": "win32_battery",
                        }
                    )
            except Exception:
                pass

        return _json(
            {
                "available": False,
                "percent": None,
                "plugged": None,
                "secs_left": None,
                "message": "Battery information is not available on this machine/session.",
            }
        )

    def _list_windows(self, include_untitled: bool, limit: int) -> str:
        windows = _serialize_windows(include_untitled=include_untitled, limit=max(1, limit))
        return _json({"count": len(windows), "windows": windows})

    def _list_processes(self, only_windowed: bool, max_results: int) -> str:
        max_results = _clamp(max_results, 1, 200)
        if only_windowed:
            windows = _serialize_windows(include_untitled=False, limit=max_results * 3)
            rows: list[dict[str, Any]] = []
            seen: set[tuple[int, str]] = set()
            for w in windows:
                name = str(w.get("process_name") or "").strip()
                title = str(w.get("title") or "").strip()
                pid = int(w.get("pid") or 0)
                if not title:
                    continue
                key = (pid, title.casefold())
                if key in seen:
                    continue
                seen.add(key)
                rows.append({"Name": name or "unknown", "Id": pid, "MainWindowTitle": title})
                if len(rows) >= max_results:
                    break
            return _json({"count": len(rows), "processes": rows})

        try:
            import psutil

            rows = []
            for p in psutil.process_iter(["pid", "name", "memory_info"]):
                try:
                    mem = int(getattr(p.info.get("memory_info"), "rss", 0) or 0)
                    rows.append(
                        {
                            "Name": p.info.get("name") or "unknown",
                            "Id": int(p.info.get("pid") or 0),
                            "MemoryMB": round(mem / (1024 * 1024), 1),
                            "MainWindowTitle": "",
                        }
                    )
                except Exception:
                    continue
            rows.sort(key=lambda x: float(x.get("MemoryMB") or 0), reverse=True)
            return _json({"count": min(len(rows), max_results), "processes": rows[:max_results]})
        except Exception as e:
            return self._error(f"list_processes failed: {e}")

    def _close_app(self, process_name: str, force: bool) -> str:
        query = _normalize_query(process_name)
        if not query:
            return self._error("process_name is required")
        try:
            import psutil

            closed = 0
            pids: list[int] = []
            query_cf = query.casefold()
            query_noext = query_cf[:-4] if query_cf.endswith(".exe") else query_cf
            terms = [t for t in re.split(r"\s+", query_noext) if t]
            current_pid = os.getpid()
            for p in psutil.process_iter(["pid", "name", "exe", "cmdline"]):
                try:
                    pid = int(p.info.get("pid") or 0)
                    if pid <= 0 or pid == current_pid:
                        continue
                    name = str(p.info.get("name") or "").casefold()
                    exe = str(p.info.get("exe") or "")
                    exe_base = Path(exe).name.casefold() if exe else ""
                    name_noext = name[:-4] if name.endswith(".exe") else name
                    exe_noext = exe_base[:-4] if exe_base.endswith(".exe") else exe_base
                    candidates = [name, exe_base, name_noext, exe_noext]

                    matched = any(query_noext and query_noext in c for c in candidates if c)
                    if not matched and terms:
                        joined = " ".join([c for c in (name_noext, exe_noext) if c])
                        matched = all(t in joined for t in terms)
                    if not matched:
                        continue

                    if force:
                        p.kill()
                    else:
                        p.terminate()
                    closed += 1
                    pids.append(pid)
                except Exception:
                    continue
            return _json({"closed": closed, "matched_pids": pids})
        except Exception as e:
            return self._error(f"close_app failed: {e}")

    def _search_start_apps(self, query: str, max_results: int) -> str:
        max_results = _clamp(max_results, 1, 2000)
        query_norm = _normalize_query(query)
        query_cf = query_norm.casefold()
        query_terms = {query_cf} if query_cf else set()
        if query_cf and (
            "telegram" in query_cf
            or "تلجرام" in query_cf
            or "تيليجرام" in query_cf
            or "تليجرام" in query_cf
        ):
            query_terms.update({"telegram", "unigram", "تلجرام", "تيليجرام", "تليجرام"})

        command = "Get-StartApps | Sort-Object Name | ConvertTo-Json -Depth 4 -Compress"
        ok, output = _run_powershell(command, timeout=20)
        if not ok:
            return self._error(f"search_start_apps failed: {output}")
        try:
            parsed = json.loads(output) if output else []
            if isinstance(parsed, dict):
                parsed = [parsed]
            apps = []
            for item in parsed:
                if not isinstance(item, dict):
                    continue
                name = str(item.get("Name", "") or "").strip()
                app_id = str(item.get("AppID", "") or "").strip()
                if not name:
                    continue
                if query_terms:
                    name_cf = name.casefold()
                    if not any(term in name_cf for term in query_terms if term):
                        continue
                apps.append({"Name": name, "AppID": app_id})
                if len(apps) >= max_results:
                    break
            return _json({"count": len(apps), "apps": apps})
        except Exception as e:
            return self._error(f"search_start_apps parse failed: {e}")

    @staticmethod
    def _parse_json_dict(raw: str) -> dict[str, Any] | None:
        try:
            parsed = json.loads(raw)
            return parsed if isinstance(parsed, dict) else None
        except Exception:
            return None

    def _try_focus_app(
        self,
        *,
        title_query: str = "",
        process_query: str = "",
        timeout_sec: float = 1.5,
    ) -> dict[str, Any] | None:
        raw = self._focus_window(
            window_title=str(title_query or ""),
            process_name=str(process_query or ""),
            timeout_sec=max(0.4, min(8.0, float(timeout_sec))),
        )
        if not isinstance(raw, str) or raw.lower().startswith("error:"):
            return None
        payload = self._parse_json_dict(raw)
        if isinstance(payload, dict) and payload.get("focused"):
            return payload
        return None

    def _launch_start_app(self, query: str) -> str:
        query_norm = _normalize_query(query)
        if not query_norm:
            return self._error("query is required")

        query_cf = query_norm.casefold()

        def _add_unique(values: list[str], extra: list[str]) -> None:
            for item in extra:
                value = _normalize_query(item)
                if value and value not in values:
                    values.append(value)

        alias_rules: dict[str, dict[str, list[str]]] = {
            "telegram": {
                "search": ["telegram", "unigram"],
                "focus_titles": ["telegram", "unigram"],
                "focus_processes": ["telegram.exe", "unigram.exe"],
                "uri": ["telegram:"],
                "shell": [],
            },
            "whatsapp": {
                "search": ["whatsapp"],
                "focus_titles": ["whatsapp"],
                "focus_processes": ["whatsapp.exe"],
                "uri": ["whatsapp:"],
                "shell": [],
            },
            "settings": {
                "search": ["settings"],
                "focus_titles": ["settings"],
                "focus_processes": ["systemsettings.exe"],
                "uri": ["ms-settings:"],
                "shell": ["ms-settings:"],
            },
            "bluetooth": {
                "search": ["bluetooth"],
                "focus_titles": ["settings", "bluetooth"],
                "focus_processes": ["systemsettings.exe"],
                "uri": ["ms-settings:bluetooth"],
                "shell": ["ms-settings:bluetooth"],
            },
            "calculator": {
                "search": ["calculator", "calc"],
                "focus_titles": ["calculator", "calc"],
                "focus_processes": ["calculatorapp.exe", "calculator.exe"],
                "uri": ["calculator:"],
                "shell": ["calc.exe"],
            },
            "calc": {
                "search": ["calculator", "calc"],
                "focus_titles": ["calculator", "calc"],
                "focus_processes": ["calculatorapp.exe", "calculator.exe"],
                "uri": ["calculator:"],
                "shell": ["calc.exe"],
            },
            "notepad": {
                "search": ["notepad"],
                "focus_titles": ["notepad"],
                "focus_processes": ["notepad.exe"],
                "uri": [],
                "shell": ["notepad.exe"],
            },
            "paint": {
                "search": ["paint", "mspaint"],
                "focus_titles": ["paint"],
                "focus_processes": ["mspaint.exe"],
                "uri": [],
                "shell": ["mspaint.exe"],
            },
            "cmd": {
                "search": ["command prompt", "cmd"],
                "focus_titles": ["command prompt", "cmd"],
                "focus_processes": ["cmd.exe"],
                "uri": [],
                "shell": ["cmd.exe"],
            },
            "command prompt": {
                "search": ["command prompt", "cmd"],
                "focus_titles": ["command prompt", "cmd"],
                "focus_processes": ["cmd.exe"],
                "uri": [],
                "shell": ["cmd.exe"],
            },
            "powershell": {
                "search": ["powershell", "windows powershell"],
                "focus_titles": ["powershell"],
                "focus_processes": ["powershell.exe", "pwsh.exe"],
                "uri": [],
                "shell": ["powershell.exe"],
            },
            "task manager": {
                "search": ["task manager", "taskmgr"],
                "focus_titles": ["task manager"],
                "focus_processes": ["taskmgr.exe"],
                "uri": [],
                "shell": ["taskmgr.exe"],
            },
            "taskmgr": {
                "search": ["task manager", "taskmgr"],
                "focus_titles": ["task manager"],
                "focus_processes": ["taskmgr.exe"],
                "uri": [],
                "shell": ["taskmgr.exe"],
            },
            "file explorer": {
                "search": ["file explorer", "explorer"],
                "focus_titles": ["file explorer", "explorer"],
                "focus_processes": ["explorer.exe"],
                "uri": [],
                "shell": ["explorer.exe"],
            },
            "explorer": {
                "search": ["file explorer", "explorer"],
                "focus_titles": ["file explorer", "explorer"],
                "focus_processes": ["explorer.exe"],
                "uri": [],
                "shell": ["explorer.exe"],
            },
        }

        search_terms: list[str] = [query_norm]
        focus_titles: list[str] = [query_norm]
        focus_processes: list[str] = []
        uri_targets: list[str] = []
        shell_targets: list[str] = []

        for token, bundle in alias_rules.items():
            if token in query_cf:
                _add_unique(search_terms, bundle.get("search", []))
                _add_unique(focus_titles, bundle.get("focus_titles", []))
                _add_unique(focus_processes, bundle.get("focus_processes", []))
                _add_unique(uri_targets, bundle.get("uri", []))
                _add_unique(shell_targets, bundle.get("shell", []))

        # If process-style query is provided directly.
        if query_cf.endswith(".exe"):
            _add_unique(focus_processes, [query_norm])
            _add_unique(shell_targets, [query_norm])

        # Focus already-open window first to avoid duplicate launches.
        for title in focus_titles[:4]:
            focused = self._try_focus_app(title_query=title, timeout_sec=1.3)
            if focused:
                return _json(
                    {
                        "launched": True,
                        "method": "focus_existing",
                        "query": query_norm,
                        "title": focused.get("title", ""),
                        "process_name": focused.get("process_name", ""),
                        "pid": focused.get("pid", 0),
                    }
                )
        for proc in focus_processes[:4]:
            focused = self._try_focus_app(process_query=proc, timeout_sec=1.2)
            if focused:
                return _json(
                    {
                        "launched": True,
                        "method": "focus_existing_process",
                        "query": query_norm,
                        "title": focused.get("title", ""),
                        "process_name": focused.get("process_name", ""),
                        "pid": focused.get("pid", 0),
                    }
                )

        # Search Start apps with ranking.
        search_terms_cf = [s.casefold() for s in search_terms if s]
        query_tokens = [tok for tok in re.split(r"\s+", query_cf) if tok]
        ranked_apps: list[dict[str, Any]] = []
        seen_app_ids: set[str] = set()
        for term in search_terms[:6]:
            raw_apps = self._search_start_apps(query=term, max_results=12)
            if not isinstance(raw_apps, str) or raw_apps.lower().startswith("error:"):
                continue
            payload = self._parse_json_dict(raw_apps)
            apps = payload.get("apps", []) if isinstance(payload, dict) else []
            for item in apps:
                if not isinstance(item, dict):
                    continue
                name = str(item.get("Name", "") or "").strip()
                app_id = str(item.get("AppID", "") or "").strip()
                if not app_id or app_id in seen_app_ids:
                    continue
                seen_app_ids.add(app_id)
                name_cf = name.casefold()
                score = 0
                if name_cf == query_cf:
                    score += 240
                if query_cf and query_cf in name_cf:
                    score += 140
                for token in search_terms_cf:
                    if token and token in name_cf:
                        score += 75
                for token in query_tokens:
                    if token in name_cf:
                        score += 20
                ranked_apps.append(
                    {
                        "score": score,
                        "Name": name,
                        "AppID": app_id,
                    }
                )
        ranked_apps.sort(key=lambda row: int(row.get("score", 0)), reverse=True)

        for candidate in ranked_apps[:8]:
            app_id = str(candidate.get("AppID", "") or "").strip()
            name = str(candidate.get("Name", "") or "").strip() or query_norm
            if not app_id:
                continue
            try:
                subprocess.Popen(
                    ["cmd", "/c", "start", "", f"shell:AppsFolder\\{app_id}"],
                    shell=False,
                )
                focused = (
                    self._try_focus_app(title_query=name, timeout_sec=3.8)
                    or self._try_focus_app(title_query=query_norm, timeout_sec=2.2)
                )
                if not focused:
                    for proc in focus_processes[:3]:
                        focused = self._try_focus_app(process_query=proc, timeout_sec=1.6)
                        if focused:
                            break
                payload: dict[str, Any] = {
                    "launched": True,
                    "method": "start_apps",
                    "Name": name,
                    "AppID": app_id,
                    "focused": bool(focused),
                }
                if focused:
                    payload.update(
                        {
                            "title": focused.get("title", ""),
                            "process_name": focused.get("process_name", ""),
                            "pid": focused.get("pid", 0),
                        }
                    )
                return _json(payload)
            except Exception:
                continue

        for uri in uri_targets:
            try:
                subprocess.Popen(["cmd", "/c", "start", "", uri], shell=False)
                focused = (
                    self._try_focus_app(title_query=query_norm, timeout_sec=2.4)
                    or self._try_focus_app(title_query="settings", timeout_sec=2.0)
                )
                return _json(
                    {
                        "launched": True,
                        "method": "uri",
                        "target": uri,
                        "focused": bool(focused),
                    }
                )
            except Exception:
                continue

        for shell_target in shell_targets:
            try:
                subprocess.Popen(["cmd", "/c", "start", "", shell_target], shell=False)
                focused = (
                    self._try_focus_app(title_query=query_norm, timeout_sec=2.8)
                    or self._try_focus_app(process_query=shell_target, timeout_sec=1.8)
                )
                return _json(
                    {
                        "launched": True,
                        "method": "shell_alias",
                        "target": shell_target,
                        "focused": bool(focused),
                    }
                )
            except Exception:
                continue

        try:
            subprocess.Popen(["cmd", "/c", "start", "", query_norm], shell=False)
            focused = self._try_focus_app(title_query=query_norm, timeout_sec=2.5)
            return _json(
                {
                    "launched": True,
                    "method": "shell_start",
                    "target": query_norm,
                    "focused": bool(focused),
                }
            )
        except Exception as e:
            return self._error(f"launch_start_app failed: {e}")

    def _search_files(self, query: str, max_results: int) -> str:
        query_norm = _normalize_query(query)
        if not query_norm:
            return self._error("query is required")
        max_results = _clamp(max_results, 1, 200)
        query_cf = query_norm.casefold()
        results: list[str] = []
        roots = _iter_search_roots()
        skip_dirs = {
            "node_modules",
            ".git",
            "__pycache__",
            ".venv",
            "venv",
            "appdata",
            "programdata",
            "$recycle.bin",
            "windows",
        }

        for root in roots:
            for current_root, dirs, files in os.walk(root, topdown=True):
                dirs[:] = [d for d in dirs if d.casefold() not in skip_dirs]
                for file_name in files:
                    if query_cf not in file_name.casefold():
                        continue
                    full = str(Path(current_root) / file_name)
                    results.append(full)
                    if len(results) >= max_results:
                        return _json({"count": len(results), "results": results})
        return _json({"count": len(results), "results": results})

    def _move_mouse_to_desktop_file(self, query: str, duration: float, timeout_sec: float) -> str:
        blocked = self._guard_interactive_action("move_mouse_to_desktop_file")
        if blocked:
            return blocked

        query_norm = _normalize_query(query)
        if not query_norm:
            return self._error("query is required")

        desktop_dir = Path.home() / "Desktop"
        if not desktop_dir.is_dir():
            return self._error("Desktop folder not found")

        query_cf = query_norm.casefold()
        candidates: list[Path] = []
        try:
            for item in desktop_dir.iterdir():
                name_cf = item.name.casefold()
                stem_cf = item.stem.casefold()
                if query_cf == name_cf or query_cf == stem_cf:
                    candidates.append(item)
                    continue
                if query_cf in name_cf or query_cf in stem_cf:
                    candidates.append(item)
        except Exception as e:
            return self._error(f"failed to inspect desktop items: {e}")

        if not candidates:
            return self._error(f"No desktop file matched: {query_norm}")

        candidates.sort(key=lambda p: (len(p.name), p.name.casefold()))
        target = candidates[0]
        label = target.stem or target.name

        attempts = [
            {"window_title": "Desktop", "control_type": "ListItem"},
            {"window_title": "Desktop", "control_type": ""},
            {"window_title": "", "control_type": "ListItem"},
        ]
        last_err = ""
        for attempt in attempts:
            raw = self._ui_target(
                window_title=attempt["window_title"],
                process_name="",
                control_name=label,
                auto_id="",
                control_type=attempt["control_type"],
                index=0,
                interaction="move",
                duration=duration,
                timeout_sec=timeout_sec,
            )
            if str(raw).lower().startswith("error:"):
                last_err = str(raw).replace("Error: ", "", 1)
                continue
            try:
                data = json.loads(raw)
            except Exception:
                last_err = "invalid response from ui_target"
                continue
            if isinstance(data, dict) and data.get("ok"):
                data["file_name"] = target.name
                data["file_path"] = str(target)
                return _json(data)
            last_err = "ui_target did not return success"

        try:
            subprocess.Popen(["explorer.exe", str(desktop_dir)], shell=False)
            time.sleep(0.8)
            raw = self._ui_target(
                window_title="Desktop",
                process_name="",
                control_name=label,
                auto_id="",
                control_type="ListItem",
                index=0,
                interaction="move",
                duration=duration,
                timeout_sec=max(4.0, timeout_sec),
            )
            if str(raw).lower().startswith("error:"):
                last_err = str(raw).replace("Error: ", "", 1)
            else:
                data = json.loads(raw)
                if isinstance(data, dict) and data.get("ok"):
                    data["file_name"] = target.name
                    data["file_path"] = str(target)
                    return _json(data)
        except Exception as e:
            last_err = f"fallback explorer path failed: {e}"

        return self._error(last_err or "Could not locate desktop icon on screen")

    def _list_installed_apps(self, query: str, max_results: int) -> str:
        query_norm = _normalize_query(query)
        max_results = _clamp(max_results, 1, 400)
        entries: dict[str, dict[str, Any]] = {}

        def add_entry(name: str, source: str) -> None:
            cleaned = name.strip()
            if not cleaned:
                return
            key = cleaned.casefold()
            if query_norm and query_norm.casefold() not in key:
                return
            if key not in entries:
                entries[key] = {"DisplayName": cleaned, "Source": source}

        raw_apps = self._search_start_apps(query="", max_results=600)
        if not raw_apps.lower().startswith("error:"):
            try:
                payload = json.loads(raw_apps)
                for item in (payload.get("apps", []) if isinstance(payload, dict) else []):
                    add_entry(str((item or {}).get("Name", "") or ""), "StartApps")
            except Exception:
                pass

        try:
            import winreg

            roots = [
                (winreg.HKEY_CURRENT_USER, r"SOFTWARE\Microsoft\Windows\CurrentVersion\Uninstall"),
                (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows\CurrentVersion\Uninstall"),
                (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\WOW6432Node\Microsoft\Windows\CurrentVersion\Uninstall"),
            ]
            for hive, key_path in roots:
                try:
                    with winreg.OpenKey(hive, key_path) as key:
                        index = 0
                        while True:
                            try:
                                sub_name = winreg.EnumKey(key, index)
                                index += 1
                            except OSError:
                                break
                            try:
                                with winreg.OpenKey(key, sub_name) as sub:
                                    name, _ = winreg.QueryValueEx(sub, "DisplayName")
                                    add_entry(str(name), "Registry")
                            except Exception:
                                continue
                except Exception:
                    continue
        except Exception:
            pass

        apps = sorted(entries.values(), key=lambda item: str(item.get("DisplayName", "")).casefold())
        return _json({"count": min(len(apps), max_results), "apps": apps[:max_results]})

    def _volume_state(self) -> dict[str, Any]:
        try:
            from pycaw.pycaw import AudioUtilities

            speakers = AudioUtilities.GetSpeakers()
            endpoint = getattr(speakers, "EndpointVolume", None)
            if endpoint is None:
                return self._legacy_volume_state()
            level = float(endpoint.GetMasterVolumeLevelScalar())
            muted = bool(endpoint.GetMute())
            return {"endpoint": endpoint, "level_percent": int(round(level * 100.0)), "muted": muted}
        except Exception as e:
            raise RuntimeError(f"volume backend unavailable: {e}") from e

    def _legacy_volume_state(self) -> dict[str, Any]:
        from ctypes import POINTER, cast

        from comtypes import CLSCTX_ALL
        from pycaw.pycaw import AudioUtilities, IAudioEndpointVolume

        speakers = AudioUtilities.GetSpeakers()
        interface = speakers.Activate(IAudioEndpointVolume._iid_, CLSCTX_ALL, None)
        endpoint = cast(interface, POINTER(IAudioEndpointVolume))
        level = float(endpoint.GetMasterVolumeLevelScalar())
        muted = bool(endpoint.GetMute())
        return {"endpoint": endpoint, "level_percent": int(round(level * 100.0)), "muted": muted}

    def _volume_control(self, mode: str, level: Any = None, delta: Any = None) -> str:
        mode_norm = (mode or "get").strip().lower()
        try:
            state = self._volume_state()
            endpoint = state["endpoint"]
            current = int(state["level_percent"])
            muted = bool(state["muted"])

            if mode_norm == "get":
                return _json({"level_percent": current, "muted": muted})
            if mode_norm in {"mute", "silence"}:
                endpoint.SetMute(1, None)
            elif mode_norm in {"unmute"}:
                endpoint.SetMute(0, None)
            elif mode_norm in {"max", "maximum"}:
                endpoint.SetMasterVolumeLevelScalar(1.0, None)
                endpoint.SetMute(0, None)
            elif mode_norm in {"min", "minimum"}:
                endpoint.SetMasterVolumeLevelScalar(0.0, None)
                endpoint.SetMute(0, None)
            elif mode_norm == "set":
                if level is None:
                    return self._error("level is required for mode=set")
                new_level = _clamp(int(level), 0, 100)
                endpoint.SetMasterVolumeLevelScalar(float(new_level) / 100.0, None)
                endpoint.SetMute(0 if new_level > 0 else int(muted), None)
            elif mode_norm == "up":
                d = _clamp(int(delta or 8), 1, 100)
                new_level = _clamp(current + d, 0, 100)
                endpoint.SetMasterVolumeLevelScalar(float(new_level) / 100.0, None)
                endpoint.SetMute(0, None)
            elif mode_norm == "down":
                d = _clamp(int(delta or 8), 1, 100)
                new_level = _clamp(current - d, 0, 100)
                endpoint.SetMasterVolumeLevelScalar(float(new_level) / 100.0, None)
            else:
                return self._error(f"unsupported volume mode: {mode_norm}")

            refreshed = self._volume_state()
            return _json(
                {
                    "level_percent": int(refreshed["level_percent"]),
                    "muted": bool(refreshed["muted"]),
                }
            )
        except Exception as e:
            return self._error(str(e))

    def _brightness_control(self, mode: str, level: Any = None, delta: Any = None) -> str:
        mode_norm = (mode or "get").strip().lower()
        try:
            _install_display_warning_filter_once()
            import screen_brightness_control as sbc

            def get_level() -> int:
                values = sbc.get_brightness()
                if isinstance(values, list):
                    if not values:
                        raise RuntimeError("No displays detected")
                    return int(round(sum(float(v) for v in values) / len(values)))
                return int(values)

            if mode_norm == "get":
                return _json({"brightness_percent": _clamp(get_level(), 0, 100)})

            current = _clamp(get_level(), 0, 100)
            if mode_norm in {"max", "maximum"}:
                target = 100
            elif mode_norm in {"min", "minimum"}:
                target = 0
            elif mode_norm == "set":
                if level is None:
                    return self._error("level is required for brightness set mode")
                target = _clamp(int(level), 0, 100)
            elif mode_norm == "up":
                target = _clamp(current + _clamp(int(delta or 10), 1, 100), 0, 100)
            elif mode_norm == "down":
                target = _clamp(current - _clamp(int(delta or 10), 1, 100), 0, 100)
            else:
                return self._error(f"unsupported brightness mode: {mode_norm}")

            sbc.set_brightness(target)
            return _json({"brightness_percent": _clamp(get_level(), 0, 100)})
        except Exception as e:
            return self._error(str(e))

    def _media_control(self, mode: str) -> str:
        mode_norm = (mode or "play_pause").strip().lower()
        try:
            import pyautogui

            pyautogui.FAILSAFE = False
            if mode_norm in {"play", "pause", "play_pause", "toggle"}:
                pyautogui.press("playpause")
                return "Media play/pause toggled."
            if mode_norm in {"next", "next_track"}:
                pyautogui.press("nexttrack")
                return "Skipped to next track."
            if mode_norm in {"previous", "prev", "previous_track"}:
                pyautogui.press("prevtrack")
                return "Returned to previous track."
            if mode_norm in {"stop"}:
                pyautogui.press("stop")
                return "Media playback stopped."
            return self._error(f"unsupported media mode: {mode_norm}")
        except Exception as e:
            return self._error(f"media_control failed: {e}")

    def _is_blocked_process(self, process_name: str) -> bool:
        name = str(process_name or "").strip().lower()
        if not name:
            return False
        return name in self.BLOCKED_AUTOMATION_PROCESSES

    def _is_blocked_window_class(self, class_name: str) -> bool:
        normalized = str(class_name or "").strip().casefold()
        if not normalized:
            return False
        return normalized in self.BLOCKED_WINDOW_CLASSES

    def _get_wrapper_class_name(self, wrapper: Any) -> str:
        try:
            info = wrapper.element_info
            return str(getattr(info, "class_name", "") or "").strip()
        except Exception:
            return ""

    def _foreground_process(self) -> tuple[str, int]:
        try:
            import psutil
            import win32gui
            import win32process

            hwnd = int(win32gui.GetForegroundWindow() or 0)
            if hwnd <= 0:
                return "", 0
            _, pid = win32process.GetWindowThreadProcessId(hwnd)
            if not pid:
                return "", 0
            name = psutil.Process(pid).name()
            return str(name or ""), int(pid)
        except Exception:
            return "", 0

    def _guard_interactive_action(self, action: str, keys: list[str] | None = None) -> str | None:
        name, pid = self._foreground_process()
        if self._is_blocked_process(name):
            return self._error(f"interactive action blocked on protected process: {name} (PID {pid})")

        if action == "hotkey" and keys:
            keyset = frozenset(str(k).strip().lower() for k in keys if str(k).strip())
            if keyset in self.BLOCKED_HOTKEY_COMBINATIONS:
                combo = " + ".join(sorted(keyset))
                return self._error(f"blocked hotkey combination: {combo}")
        return None

    def _resolve_window(
        self,
        *,
        window_title: str = "",
        process_name: str = "",
        timeout_sec: float = 4.0,
    ) -> dict[str, Any]:
        from pywinauto import Desktop

        title_query = _normalize_query(window_title).casefold()
        process_query = _normalize_query(process_name).casefold()
        timeout_sec = max(0.3, min(20.0, float(timeout_sec)))
        deadline = time.time() + timeout_sec

        last_err = "No matching window found."
        while time.time() < deadline:
            try:
                import psutil
                import win32gui
                import win32process

                desktop = Desktop(backend="uia")

                if not title_query and not process_query:
                    hwnd = int(win32gui.GetForegroundWindow() or 0)
                    if hwnd > 0:
                        try:
                            _, pid = win32process.GetWindowThreadProcessId(hwnd)
                            p_name = psutil.Process(pid).name() if pid else ""
                            if self._is_blocked_process(p_name):
                                raise RuntimeError(f"target process is protected: {p_name}")
                            wrapper = desktop.window(handle=hwnd)
                            class_name = self._get_wrapper_class_name(wrapper)
                            if self._is_blocked_window_class(class_name):
                                raise RuntimeError(
                                    f"target window class is blocked: {class_name or 'unknown'}"
                                )
                            return {
                                "wrapper": wrapper,
                                "title": str(wrapper.window_text() or ""),
                                "process_name": str(p_name or ""),
                                "pid": int(pid or 0),
                                "hwnd": hwnd,
                                "class_name": class_name,
                            }
                        except Exception:
                            pass

                best: dict[str, Any] | None = None
                best_score = -1
                for wrapper in desktop.windows():
                    try:
                        title = str(wrapper.window_text() or "").strip()
                        if not title:
                            continue
                        info = wrapper.element_info
                        pid = int(getattr(info, "process_id", 0) or 0)
                        class_name = str(getattr(info, "class_name", "") or "").strip()
                        if self._is_blocked_window_class(class_name):
                            continue
                        p_name = ""
                        if pid:
                            try:
                                p_name = str(psutil.Process(pid).name() or "")
                            except Exception:
                                p_name = ""
                        if self._is_blocked_process(p_name):
                            continue

                        title_cf = title.casefold()
                        process_cf = p_name.casefold()
                        if title_query and title_query not in title_cf:
                            continue
                        if process_query and process_query not in process_cf:
                            continue

                        score = 0
                        if title_query:
                            if title_cf == title_query:
                                score += 120
                            else:
                                score += 80
                        if process_query:
                            if process_cf == process_query:
                                score += 80
                            else:
                                score += 50
                        if not title_query and not process_query:
                            score += 10
                        if score > best_score:
                            best_score = score
                            hwnd = int(getattr(info, "handle", 0) or 0)
                            best = {
                                "wrapper": wrapper,
                                "title": title,
                                "process_name": p_name,
                                "pid": pid,
                                "hwnd": hwnd,
                                "class_name": class_name,
                            }
                    except Exception:
                        continue

                if best is not None:
                    return best
            except Exception as e:
                last_err = str(e)

            time.sleep(0.15)

        raise RuntimeError(last_err)

    def _enumerate_window_controls(self, window_wrapper: Any, max_items: int = 400) -> list[dict[str, Any]]:
        controls: list[dict[str, Any]] = []
        seen: set[tuple[str, str, str]] = set()

        candidates: list[Any] = [window_wrapper]
        try:
            descendants = window_wrapper.descendants()
            candidates.extend(descendants[: max(1, max_items * 2)])
        except Exception:
            pass

        for wrapper in candidates:
            try:
                info = wrapper.element_info
                name = str(wrapper.window_text() or "").strip()
                auto_id = str(getattr(info, "automation_id", "") or "").strip()
                control_type = str(getattr(info, "control_type", "") or "").strip()
                class_name = str(getattr(info, "class_name", "") or "").strip()
                left = top = right = bottom = width = height = 0
                try:
                    rect = wrapper.rectangle()
                    left = int(getattr(rect, "left", 0) or 0)
                    top = int(getattr(rect, "top", 0) or 0)
                    right = int(getattr(rect, "right", 0) or 0)
                    bottom = int(getattr(rect, "bottom", 0) or 0)
                    width = max(0, right - left)
                    height = max(0, bottom - top)
                except Exception:
                    pass
                if not any((name, auto_id, control_type, class_name)):
                    continue
                key = (name.casefold(), auto_id.casefold(), control_type.casefold())
                if key in seen:
                    continue
                seen.add(key)
                controls.append(
                    {
                        "wrapper": wrapper,
                        "name": name,
                        "auto_id": auto_id,
                        "control_type": control_type,
                        "class_name": class_name,
                        "left": left,
                        "top": top,
                        "right": right,
                        "bottom": bottom,
                        "width": width,
                        "height": height,
                    }
                )
                if len(controls) >= max_items:
                    break
            except Exception:
                continue
        return controls

    def _pick_control(
        self,
        controls: list[dict[str, Any]],
        *,
        control_name: str = "",
        auto_id: str = "",
        control_type: str = "",
        index: int = 0,
    ) -> dict[str, Any]:
        if not controls:
            raise RuntimeError("No controls found in target window.")

        name_q = _normalize_query(control_name).casefold()
        auto_q = _normalize_query(auto_id).casefold()
        type_q = _normalize_query(control_type).casefold()
        name_tokens = [t for t in re.split(r"\s+", name_q) if t]

        ranked: list[tuple[int, dict[str, Any]]] = []
        for item in controls:
            name = str(item.get("name", "") or "")
            aid = str(item.get("auto_id", "") or "")
            ctype = str(item.get("control_type", "") or "")
            klass = str(item.get("class_name", "") or "")
            left = int(item.get("left", 0) or 0)
            top = int(item.get("top", 0) or 0)
            width = int(item.get("width", 0) or 0)
            height = int(item.get("height", 0) or 0)

            if width <= 1 or height <= 1:
                continue
            if abs(left) > 10000 or abs(top) > 10000:
                continue

            name_cf = name.casefold()
            aid_cf = aid.casefold()
            ctype_cf = ctype.casefold()
            klass_cf = klass.casefold()

            token_ratio = 0.0
            if name_q:
                if "\n" in name or len(name) > 120:
                    continue
                if not type_q and ctype_cf in {"text", "document"}:
                    continue
                exact_match = name_cf == name_q
                contains_match = bool(name_cf) and (name_q in name_cf or name_cf in name_q)
                token_hits = sum(1 for t in name_tokens if t in name_cf)
                token_ratio = (token_hits / len(name_tokens)) if name_tokens else 0.0
                if not (exact_match or contains_match or token_ratio >= 0.5):
                    continue
            if auto_q and auto_q not in aid_cf:
                continue
            if type_q and type_q not in ctype_cf and type_q not in klass_cf:
                continue

            score = 0
            if name_q:
                if name_cf == name_q:
                    score += 120
                elif name_q in name_cf or name_cf in name_q:
                    score += 90
                else:
                    score += int(60 * token_ratio)
            if auto_q:
                score += 90 if aid_cf == auto_q else 55
            if type_q:
                score += 40
            if name and not name_q:
                score += 5
            if "\n" in name:
                score -= 40
            if len(name) > 180:
                score -= 30
            clickable_types = {
                "button",
                "hyperlink",
                "menuitem",
                "tabitem",
                "listitem",
                "treeitem",
                "checkbox",
                "radiobutton",
                "splitbutton",
            }
            if ctype_cf in clickable_types:
                score += 25
            if ctype_cf in {"text", "document"}:
                score -= 15

            ranked.append((score, item))

        if not ranked:
            raise RuntimeError("No matching control found.")

        ranked.sort(key=lambda pair: pair[0], reverse=True)
        safe_index = _clamp(index, 0, len(ranked) - 1)
        return ranked[safe_index][1]

    def _control_center(self, wrapper: Any) -> tuple[int, int]:
        try:
            rect = wrapper.rectangle()
            left = int(getattr(rect, "left", 0) or 0)
            top = int(getattr(rect, "top", 0) or 0)
            right = int(getattr(rect, "right", 0) or 0)
            bottom = int(getattr(rect, "bottom", 0) or 0)
            if right <= left or bottom <= top:
                raise RuntimeError("control has invalid bounds")
            x = left + ((right - left) // 2)
            y = top + ((bottom - top) // 2)
            return x, y
        except Exception as e:
            raise RuntimeError(f"failed to resolve control bounds: {e}") from e

    def _focus_window(self, window_title: str, process_name: str, timeout_sec: float) -> str:
        try:
            target = self._resolve_window(
                window_title=window_title,
                process_name=process_name,
                timeout_sec=timeout_sec,
            )
            wrapper = target["wrapper"]
            try:
                if hasattr(wrapper, "is_minimized") and wrapper.is_minimized():
                    wrapper.restore()
            except Exception:
                pass
            try:
                wrapper.set_focus()
            except Exception:
                try:
                    wrapper.click_input()
                except Exception:
                    pass
            return _json(
                {
                    "focused": True,
                    "title": target.get("title", ""),
                    "process_name": target.get("process_name", ""),
                    "pid": target.get("pid", 0),
                }
            )
        except Exception as e:
            return self._error(f"focus_window failed: {e}")

    def _ui_list_controls(
        self,
        *,
        window_title: str,
        process_name: str,
        control_name: str,
        control_type: str,
        max_results: int,
        timeout_sec: float,
    ) -> str:
        try:
            target = self._resolve_window(
                window_title=window_title,
                process_name=process_name,
                timeout_sec=timeout_sec,
            )
            controls = self._enumerate_window_controls(target["wrapper"], max_items=max(50, max_results * 5))

            name_q = _normalize_query(control_name).casefold()
            type_q = _normalize_query(control_type).casefold()
            filtered: list[dict[str, Any]] = []
            for item in controls:
                name = str(item.get("name", "") or "")
                ctype = str(item.get("control_type", "") or "")
                klass = str(item.get("class_name", "") or "")
                if name_q and name_q not in name.casefold():
                    continue
                if type_q and type_q not in ctype.casefold() and type_q not in klass.casefold():
                    continue
                filtered.append(
                    {
                        "name": name,
                        "auto_id": str(item.get("auto_id", "") or ""),
                        "control_type": ctype,
                        "class_name": klass,
                    }
                )
                if len(filtered) >= _clamp(max_results, 1, 200):
                    break

            return _json(
                {
                    "window_title": target.get("title", ""),
                    "process_name": target.get("process_name", ""),
                    "count": len(filtered),
                    "controls": filtered,
                }
            )
        except Exception as e:
            return self._error(f"ui_list_controls failed: {e}")

    def _ui_click(
        self,
        *,
        window_title: str,
        process_name: str,
        control_name: str,
        auto_id: str,
        control_type: str,
        index: int,
        timeout_sec: float,
    ) -> str:
        blocked = self._guard_interactive_action("ui_click")
        if blocked:
            return blocked
        try:
            target = self._resolve_window(
                window_title=window_title,
                process_name=process_name,
                timeout_sec=timeout_sec,
            )
            controls = self._enumerate_window_controls(target["wrapper"])
            selected = self._pick_control(
                controls,
                control_name=control_name,
                auto_id=auto_id,
                control_type=control_type,
                index=index,
            )
            wrapper = selected["wrapper"]
            try:
                if hasattr(wrapper, "invoke"):
                    wrapper.invoke()
                else:
                    wrapper.click_input()
            except Exception:
                wrapper.click_input()
            return _json(
                {
                    "clicked": True,
                    "window_title": target.get("title", ""),
                    "process_name": target.get("process_name", ""),
                    "control_name": selected.get("name", ""),
                    "auto_id": selected.get("auto_id", ""),
                    "control_type": selected.get("control_type", ""),
                }
            )
        except Exception as e:
            return self._error(f"ui_click failed: {e}")

    def _ui_set_text(
        self,
        *,
        text: str,
        window_title: str,
        process_name: str,
        control_name: str,
        auto_id: str,
        control_type: str,
        index: int,
        press_enter: bool,
        timeout_sec: float,
    ) -> str:
        blocked = self._guard_interactive_action("ui_set_text")
        if blocked:
            return blocked
        text_value = str(text or "")
        if not text_value:
            return self._error("text is required")
        try:
            import pyautogui

            pyautogui.FAILSAFE = False
            target = self._resolve_window(
                window_title=window_title,
                process_name=process_name,
                timeout_sec=timeout_sec,
            )
            controls = self._enumerate_window_controls(target["wrapper"])
            selected = self._pick_control(
                controls,
                control_name=control_name,
                auto_id=auto_id,
                control_type=control_type,
                index=index,
            )
            wrapper = selected["wrapper"]
            used_method = "typewrite"
            try:
                if hasattr(wrapper, "set_edit_text"):
                    wrapper.set_edit_text(text_value)
                    used_method = "set_edit_text"
                else:
                    wrapper.click_input()
                    pyautogui.hotkey("ctrl", "a")
                    pyautogui.press("backspace")
                    pyautogui.write(text_value, interval=0.01)
            except Exception:
                wrapper.click_input()
                pyautogui.hotkey("ctrl", "a")
                pyautogui.press("backspace")
                pyautogui.write(text_value, interval=0.01)
            if press_enter:
                pyautogui.press("enter")
            return _json(
                {
                    "text_set": True,
                    "method": used_method,
                    "window_title": target.get("title", ""),
                    "process_name": target.get("process_name", ""),
                    "control_name": selected.get("name", ""),
                    "chars": len(text_value),
                    "pressed_enter": bool(press_enter),
                }
            )
        except Exception as e:
            return self._error(f"ui_set_text failed: {e}")

    def _ui_target(
        self,
        *,
        window_title: str,
        process_name: str,
        control_name: str,
        auto_id: str,
        control_type: str,
        index: int,
        interaction: str,
        duration: float,
        timeout_sec: float,
    ) -> str:
        blocked = self._guard_interactive_action("ui_target")
        if blocked:
            return blocked
        try:
            import pyautogui

            pyautogui.FAILSAFE = False

            interaction_norm = (interaction or "move").strip().lower()
            interaction_aliases = {
                "hover": "move",
                "point": "move",
                "double": "double_click",
                "double-click": "double_click",
                "right": "right_click",
                "context": "right_click",
            }
            interaction_norm = interaction_aliases.get(interaction_norm, interaction_norm)
            if interaction_norm not in {"move", "click", "double_click", "right_click"}:
                return self._error(f"unsupported ui_target interaction: {interaction_norm}")

            target = self._resolve_window(
                window_title=window_title,
                process_name=process_name,
                timeout_sec=timeout_sec,
            )
            controls = self._enumerate_window_controls(target["wrapper"])
            screen = pyautogui.size()
            selected: dict[str, Any] | None = None
            x = y = 0
            limit = min(len(controls), 20)
            start_index = max(0, int(index))
            coord_error = ""
            for idx in range(start_index, limit):
                candidate = self._pick_control(
                    controls,
                    control_name=control_name,
                    auto_id=auto_id,
                    control_type=control_type,
                    index=idx,
                )
                candidate_name = str(candidate.get("name", "") or "").strip()
                if control_name and not candidate_name and not str(candidate.get("auto_id", "") or "").strip():
                    continue
                cx, cy = self._control_center(candidate["wrapper"])
                if abs(int(cx)) > int(screen.width * 3) or abs(int(cy)) > int(screen.height * 3):
                    coord_error = f"target coordinates out of expected bounds: ({cx}, {cy})"
                    continue
                selected = candidate
                x, y = int(cx), int(cy)
                break

            if selected is None:
                if coord_error:
                    raise RuntimeError(coord_error)
                raise RuntimeError("No viable control match found.")

            move_duration = max(0.0, min(3.0, float(duration)))

            if interaction_norm == "move":
                pyautogui.moveTo(x, y, duration=move_duration)
            elif interaction_norm == "click":
                pyautogui.click(x=x, y=y, button="left", clicks=1)
            elif interaction_norm == "double_click":
                pyautogui.click(x=x, y=y, button="left", clicks=2)
            elif interaction_norm == "right_click":
                pyautogui.click(x=x, y=y, button="right", clicks=1)

            return _json(
                {
                    "ok": True,
                    "interaction": interaction_norm,
                    "x": x,
                    "y": y,
                    "window_title": target.get("title", ""),
                    "process_name": target.get("process_name", ""),
                    "control_name": selected.get("name", ""),
                    "auto_id": selected.get("auto_id", ""),
                    "control_type": selected.get("control_type", ""),
                }
            )
        except Exception as e:
            return self._error(f"ui_target failed: {e}")

    def _mouse_move(self, x: int, y: int, duration: float) -> str:
        blocked = self._guard_interactive_action("mouse_move")
        if blocked:
            return blocked
        try:
            import pyautogui

            pyautogui.FAILSAFE = False
            pyautogui.moveTo(int(x), int(y), duration=max(0.0, min(3.0, float(duration))))
            return _json({"ok": True, "x": int(x), "y": int(y)})
        except Exception as e:
            return self._error(f"mouse_move failed: {e}")

    def _click(self, x: Any, y: Any, button: str, clicks: int) -> str:
        blocked = self._guard_interactive_action("click")
        if blocked:
            return blocked
        try:
            import pyautogui

            pyautogui.FAILSAFE = False
            btn = (button or "left").strip().lower()
            if btn not in {"left", "right", "middle"}:
                btn = "left"
            c = _clamp(clicks, 1, 5)
            if x is not None and y is not None:
                pyautogui.click(x=int(x), y=int(y), clicks=c, button=btn)
                return _json({"ok": True, "x": int(x), "y": int(y), "button": btn, "clicks": c})
            pyautogui.click(clicks=c, button=btn)
            pos = pyautogui.position()
            return _json({"ok": True, "x": int(pos.x), "y": int(pos.y), "button": btn, "clicks": c})
        except Exception as e:
            return self._error(f"click failed: {e}")

    def _press_key(self, key: str) -> str:
        key_norm = (key or "").strip().lower()
        if not key_norm:
            return self._error("key is required")
        blocked = self._guard_interactive_action("press_key")
        if blocked:
            return blocked
        try:
            import pyautogui

            pyautogui.FAILSAFE = False
            pyautogui.press(key_norm)
            return _json({"ok": True, "key": key_norm})
        except Exception as e:
            return self._error(f"press_key failed: {e}")

    def _type_text(self, text: str, press_enter: bool, interval: float) -> str:
        text_value = str(text or "")
        if not text_value:
            return self._error("text is required")
        blocked = self._guard_interactive_action("type_text")
        if blocked:
            return blocked
        try:
            import pyautogui

            pyautogui.FAILSAFE = False
            interval_sec = max(0.0, min(0.4, float(interval)))
            pyautogui.write(text_value, interval=interval_sec)
            if press_enter:
                pyautogui.press("enter")
            return _json(
                {
                    "ok": True,
                    "typed": True,
                    "chars": len(text_value),
                    "pressed_enter": bool(press_enter),
                }
            )
        except Exception as e:
            return self._error(f"type_text failed: {e}")

    def _hotkey(self, keys: list[str]) -> str:
        usable = [k.strip().lower() for k in keys if str(k).strip()]
        if not usable:
            return self._error("keys are required")
        blocked = self._guard_interactive_action("hotkey", keys=usable)
        if blocked:
            return blocked
        try:
            import pyautogui

            pyautogui.FAILSAFE = False
            pyautogui.hotkey(*usable)
            return _json({"ok": True, "keys": usable})
        except Exception as e:
            return self._error(f"hotkey failed: {e}")

    def _camera_snapshot(self, camera_index: int) -> str:
        try:
            import cv2

            media_dir = get_media_dir()
            out_path = media_dir / f"camera_{_timestamp_id()}.jpg"
            cap = cv2.VideoCapture(camera_index, cv2.CAP_DSHOW)
            if not cap or not cap.isOpened():
                cap = cv2.VideoCapture(camera_index)
            if not cap or not cap.isOpened():
                return self._error("camera is not available")
            try:
                time.sleep(0.2)
                ok, frame = cap.read()
                if not ok or frame is None:
                    return self._error("failed to read camera frame")
                cv2.imwrite(str(out_path), frame)
            finally:
                cap.release()
            return f"Camera snapshot saved to {out_path}"
        except Exception as e:
            return self._error(f"camera_snapshot failed: {e}")

    def _microphone_record(self, seconds: float) -> str:
        sec = max(0.4, min(30.0, float(seconds or 3.0)))
        try:
            import numpy as np
            import sounddevice as sd

            samplerate = 16000
            total_samples = int(sec * samplerate)
            if total_samples <= 0:
                return self._error("recording duration must be positive")
            media_dir = get_media_dir()
            out_path = media_dir / f"mic_{_timestamp_id()}.wav"
            audio = sd.rec(total_samples, samplerate=samplerate, channels=1, dtype="int16")
            sd.wait()
            if audio is None:
                return self._error("no audio data captured")
            if not isinstance(audio, np.ndarray):
                return self._error("audio capture returned invalid data")
            with wave.open(str(out_path), "wb") as wf:
                wf.setnchannels(1)
                wf.setsampwidth(2)
                wf.setframerate(samplerate)
                wf.writeframes(audio.tobytes())
            return f"Microphone recording saved to {out_path}"
        except Exception as e:
            return self._error(f"microphone_record failed: {e}")

    def _open_settings_page(self, page: str) -> str:
        page_norm = _normalize_query(page).casefold()
        mapping = {
            "": "ms-settings:",
            "settings": "ms-settings:",
            "bluetooth": "ms-settings:bluetooth",
            "wifi": "ms-settings:network-wifi",
            "network": "ms-settings:network",
            "display": "ms-settings:display",
            "sound": "ms-settings:sound",
            "volume": "ms-settings:sound",
            "apps": "ms-settings:appsfeatures",
            "default apps": "ms-settings:defaultapps",
            "defaultapps": "ms-settings:defaultapps",
            "privacy": "ms-settings:privacy",
            "update": "ms-settings:windowsupdate",
            "language": "ms-settings:regionlanguage",
            "keyboard": "ms-settings:typing",
        }
        uri = mapping.get(page_norm, None)
        if uri is None:
            if "default" in page_norm and "app" in page_norm:
                uri = "ms-settings:defaultapps"
            elif "bluetooth" in page_norm:
                uri = "ms-settings:bluetooth"
            elif "app" in page_norm:
                uri = "ms-settings:appsfeatures"
            else:
                uri = "ms-settings:"
        try:
            subprocess.Popen(["cmd", "/c", "start", "", uri], shell=False)
            return _json({"opened": True, "uri": uri})
        except Exception as e:
            return self._error(f"open_settings_page failed: {e}")

    def _bluetooth_control(self, mode: str) -> str:
        mode_norm = (mode or "open_settings").strip().lower()
        if mode_norm in {"open", "open_settings", "settings"}:
            return self._open_settings_page("bluetooth")

        desired = None
        if mode_norm in {"on", "enable", "start"}:
            desired = "On"
        elif mode_norm in {"off", "disable", "stop"}:
            desired = "Off"
        elif mode_norm in {"toggle"}:
            desired = "Toggle"

        if desired is None:
            return self._error(f"unsupported bluetooth mode: {mode_norm}")

        set_script = f"""
$ErrorActionPreference='Stop'
Add-Type -AssemblyName System.Runtime.WindowsRuntime
$asTask = ([System.WindowsRuntimeSystemExtensions].GetMethods() | Where-Object {{
    $_.Name -eq 'AsTask' -and $_.IsGenericMethod -and $_.GetParameters().Count -eq 1
}} | Select-Object -First 1)
if (-not $asTask) {{ throw 'AsTask bridge unavailable' }}
$null = [Windows.Devices.Radios.Radio, Windows.System.Devices, ContentType=WindowsRuntime]
$accessOp = [Windows.Devices.Radios.Radio]::RequestAccessAsync()
$accessTask = $asTask.MakeGenericMethod([Windows.Devices.Radios.RadioAccessStatus]).Invoke($null, @($accessOp))
$accessTask.Wait()
if ($accessTask.Result -ne [Windows.Devices.Radios.RadioAccessStatus]::Allowed) {{
    throw 'Bluetooth access denied'
}}
$radiosOp = [Windows.Devices.Radios.Radio]::GetRadiosAsync()
$radiosTask = $asTask.MakeGenericMethod([System.Collections.Generic.IReadOnlyList[Windows.Devices.Radios.Radio]]).Invoke($null, @($radiosOp))
$radiosTask.Wait()
$radio = $radiosTask.Result | Where-Object {{ $_.Kind -eq [Windows.Devices.Radios.RadioKind]::Bluetooth }} | Select-Object -First 1
if (-not $radio) {{ throw 'Bluetooth radio not found' }}
$current = $radio.State.ToString()
$target = '{desired}'
if ($target -eq 'Toggle') {{
    if ($current -eq 'On') {{ $target = 'Off' }} else {{ $target = 'On' }}
}}
$setOp = $radio.SetStateAsync([Windows.Devices.Radios.RadioState]::$target)
$setTask = $asTask.MakeGenericMethod([Windows.Devices.Radios.RadioAccessStatus]).Invoke($null, @($setOp))
$setTask.Wait()
$stateNow = $radio.State.ToString()
$obj = [ordered]@{{ ok = $true; requested = $target.ToLower(); state = $stateNow.ToLower() }}
$obj | ConvertTo-Json -Compress
"""
        ok, output = _run_powershell(set_script, timeout=25)
        if ok and output:
            try:
                data = json.loads(output)
                if isinstance(data, dict):
                    return _json(data)
            except Exception:
                return _json({"ok": True, "state": output.strip().casefold(), "requested": mode_norm})

        fallback = self._open_settings_page("bluetooth")
        if fallback.lower().startswith("error:"):
            return self._error(f"bluetooth control failed: {output}")
        return _json(
            {
                "ok": False,
                "requested": mode_norm,
                "error": output or "Bluetooth state change failed",
                "opened_settings": True,
            }
        )

    def _system_power(self, mode: str, name: str = "") -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm in {"power_plan_balanced", "power_plan_saver", "power_plan_high"}:
            guid_map = {
                "power_plan_balanced": "381b4222-f694-41f0-9685-ff5bb260df2e",
                "power_plan_saver": "a1841308-3541-4fab-bc81-f71556f20b4a",
                "power_plan_high": "8c5e7fda-e8bf-4a96-9a85-a6e23a8c635c",
            }
            guid = guid_map[mode_norm]
            ok, out = _run_powershell(f"powercfg /setactive {guid}", timeout=10)
            if ok:
                return _json({"ok": True, "action": mode_norm, "guid": guid})
            return self._error(out or "set power plan failed")
        if mode_norm in {"lock", "lock_screen"}:
            ok, out = _run_powershell("rundll32.exe user32.dll,LockWorkStation", timeout=6)
            if ok:
                return _json({"ok": True, "action": "lock"})
            return self._error(out or "lock failed")
        if mode_norm in {"sleep"}:
            ok, out = _run_powershell("rundll32.exe powrprof.dll,SetSuspendState 0,1,0", timeout=8)
            if ok:
                return _json({"ok": True, "action": "sleep"})
            return self._error(out or "sleep failed")
        if mode_norm in {"hibernate"}:
            ok, out = _run_powershell("shutdown /h", timeout=8)
            if ok:
                return _json({"ok": True, "action": "hibernate"})
            return self._error(out or "hibernate failed")
        if mode_norm in {"logoff", "logout"}:
            ok, out = _run_powershell("shutdown /l", timeout=8)
            if ok:
                return _json({"ok": True, "action": "logoff"})
            return self._error(out or "logoff failed")
        if mode_norm in {"shutdown", "poweroff"}:
            ok, out = _run_powershell("shutdown /s /t 0", timeout=8)
            if ok:
                return _json({"ok": True, "action": "shutdown"})
            return self._error(out or "shutdown failed")
        if mode_norm in {"restart", "reboot"}:
            ok, out = _run_powershell("shutdown /r /t 0", timeout=8)
            if ok:
                return _json({"ok": True, "action": "restart"})
            return self._error(out or "restart failed")
        if mode_norm in {"bios", "reboot_bios"}:
            ok, out = _run_powershell("shutdown /r /fw /t 0", timeout=8)
            if ok:
                return _json({"ok": True, "action": "reboot_bios"})
            return self._error(out or "bios reboot failed")
        if mode_norm in {"screen_off"}:
            ps = (
                "$sig='[DllImport(\"user32.dll\")]public static extern IntPtr SendMessage(IntPtr hWnd,int Msg,IntPtr wParam,IntPtr lParam);'; "
                "Add-Type -Name NativeMethods -Namespace Win32 -MemberDefinition $sig; "
                "[Win32.NativeMethods]::SendMessage([intptr]0xffff,0x0112,[intptr]0xF170,[intptr]2) | Out-Null; "
                "@{ok=$true; action='screen_off'} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=10)
            return out if ok and out else self._error(out or "screen off failed")
        if mode_norm in {"rename_pc", "rename_computer"}:
            new_name = (name or "").strip()
            if not new_name:
                return self._error("name is required")
            ps = (
                "$ErrorActionPreference='Stop'; "
                f"Rename-Computer -NewName '{new_name}' -Force; "
                f"@{{ok=$true; action='rename_computer'; name='{new_name}'}} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=15)
            if ok and out:
                return out
            return self._error(out or "rename computer failed")
        return self._error(f"unsupported system_power mode: {mode_norm}")

    def _shutdown_schedule(self, mode: str, minutes: Any) -> str:
        mode_norm = (mode or "set").strip().lower()
        if mode_norm in {"cancel", "clear"}:
            ok, out = _run_powershell("shutdown /a", timeout=8)
            if ok:
                return _json({"ok": True, "action": "cancel_shutdown"})
            return self._error(out or "cancel shutdown failed")
        mins = 0
        try:
            mins = int(minutes if minutes is not None else 0)
        except Exception:
            mins = 0
        if mins <= 0:
            return self._error("minutes must be > 0")
        seconds = mins * 60
        ok, out = _run_powershell(f"shutdown /s /t {seconds}", timeout=8)
        if ok:
            return _json({"ok": True, "action": "schedule_shutdown", "minutes": mins, "seconds": seconds})
        return self._error(out or "schedule shutdown failed")

    def _system_info(self, mode: str) -> str:
        mode_norm = (mode or "windows_version").strip().lower()
        if mode_norm in {"uptime", "boot_time"}:
            ps = (
                "$os = Get-CimInstance Win32_OperatingSystem; "
                "$boot = $os.LastBootUpTime; "
                "$uptime = (Get-Date) - $boot; "
                "@{ok=$true; mode='uptime'; last_boot=$boot.ToString('s'); "
                "uptime_days=[math]::Round($uptime.TotalDays,2)} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=10)
            if ok and out:
                return out
            return self._error(out or "uptime failed")
        if mode_norm in {"windows_version", "version"}:
            ps = (
                "$cv = Get-ItemProperty 'HKLM:\\SOFTWARE\\Microsoft\\Windows NT\\CurrentVersion'; "
                "@{ok=$true; mode='windows_version'; "
                "product_name=$cv.ProductName; display_version=$cv.DisplayVersion; build=$cv.CurrentBuild} "
                "| ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=10)
            if ok and out:
                return out
            return self._error(out or "windows_version failed")
        if mode_norm in {"about_page", "about"}:
            return self._open_settings_page("about")
        if mode_norm in {"battery"}:
            return self._battery_status()
        if mode_norm in {"battery_minutes"}:
            return self._hardware_tools("battery_minutes")
        return self._error(f"unsupported system_info mode: {mode_norm}")

    def _network_tools(self, mode: str, host: str = "", port: Any = None) -> str:
        mode_norm = (mode or "ip_internal").strip().lower()
        if mode_norm in {"ip_internal", "local_ip"}:
            ps = (
                "Get-NetIPAddress -AddressFamily IPv4 | "
                "Where-Object {$_.IPAddress -notlike '169.254*' -and $_.PrefixOrigin -ne 'WellKnown'} | "
                "Select-Object -First 6 InterfaceAlias,IPAddress | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=10)
            if ok and out:
                return _json({"ok": True, "mode": "ip_internal", "data": json.loads(out)})
            return self._error(out or "ip internal lookup failed")
        if mode_norm in {"ip_external", "public_ip"}:
            ps = "(Invoke-RestMethod -Uri 'https://api.ipify.org?format=json' -TimeoutSec 8) | ConvertTo-Json -Compress"
            ok, out = _run_powershell(ps, timeout=12)
            if ok and out:
                return _json({"ok": True, "mode": "ip_external", "data": json.loads(out)})
            return self._error(out or "external ip lookup failed")
        if mode_norm in {"ipconfig_all"}:
            ok, out = _run_powershell("ipconfig /all", timeout=18)
            return _json({"ok": True, "mode": "ipconfig_all", "output": out[:3000]}) if ok else self._error(out or "ipconfig all failed")
        if mode_norm in {"flush_dns"}:
            ok, out = _run_powershell("ipconfig /flushdns", timeout=10)
            if ok:
                return _json({"ok": True, "mode": "flush_dns"})
            return self._error(out or "flush dns failed")
        if mode_norm in {"renew_ip", "release_renew"}:
            ok1, out1 = _run_powershell("ipconfig /release", timeout=15)
            ok2, out2 = _run_powershell("ipconfig /renew", timeout=20)
            if ok1 and ok2:
                return _json({"ok": True, "mode": "release_renew"})
            return self._error((out2 or out1 or "release/renew failed").strip())
        if mode_norm in {"ping"}:
            target = (host or "").strip() or "8.8.8.8"
            ps = f"ping -n 4 {target}"
            ok, out = _run_powershell(ps, timeout=20)
            if ok:
                return _json({"ok": True, "mode": "ping", "host": target, "output": out[:1200]})
            return self._error(out or f"ping failed: {target}")
        if mode_norm in {"wifi_on"}:
            ok, out = _run_powershell("netsh interface set interface name='Wi-Fi' admin=enabled", timeout=12)
            if ok:
                return _json({"ok": True, "mode": "wifi_on"})
            return self._error(out or "wifi on failed")
        if mode_norm in {"wifi_off"}:
            ok, out = _run_powershell("netsh interface set interface name='Wi-Fi' admin=disabled", timeout=12)
            if ok:
                return _json({"ok": True, "mode": "wifi_off"})
            return self._error(out or "wifi off failed")
        if mode_norm in {"wifi_passwords"}:
            ps = (
                "$profiles=(netsh wlan show profiles) | Select-String 'All User Profile'; "
                "$names=@(); foreach($p in $profiles){$names += ($p -split ':')[1].Trim()}; "
                "$out=@(); foreach($n in $names){ "
                "$k=(netsh wlan show profile name=\"$n\" key=clear) | Select-String 'Key Content'; "
                "$pwd=''; if($k){$pwd=($k -split ':')[1].Trim()} ; "
                "$out += [pscustomobject]@{ssid=$n; password=$pwd} }; "
                "$out | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=25)
            if ok and out:
                return _json({"ok": True, "mode": "wifi_passwords", "data": json.loads(out)})
            return self._error(out or "wifi passwords failed")
        if mode_norm in {"disconnect_wifi"}:
            ok, out = _run_powershell("netsh wlan disconnect", timeout=10)
            return _json({"ok": True, "mode": "disconnect_wifi"}) if ok else self._error(out or "wifi disconnect failed")
        if mode_norm in {"disconnect_current_network"}:
            return self._network_tools("disconnect_wifi")
        if mode_norm in {"connect_wifi"}:
            ssid = (host or "").strip()
            if not ssid:
                return self._error("host/ssid is required")
            ok, out = _run_powershell(f"netsh wlan connect name=\"{ssid}\"", timeout=15)
            return _json({"ok": True, "mode": "connect_wifi", "ssid": ssid}) if ok else self._error(out or "wifi connect failed")
        if mode_norm in {"route_table"}:
            ok, out = _run_powershell("route print", timeout=12)
            return _json({"ok": True, "mode": "route_table", "output": out[:2000]}) if ok else self._error(out or "route table failed")
        if mode_norm in {"tracert", "trace_route"}:
            target = (host or "").strip() or "8.8.8.8"
            ok, out = _run_powershell(f"tracert -d {target}", timeout=35)
            return _json({"ok": True, "mode": "tracert", "host": target, "output": out[:2500]}) if ok else self._error(out or "tracert failed")
        if mode_norm in {"pathping"}:
            target = (host or "").strip() or "8.8.8.8"
            ok, out = _run_powershell(f"pathping -n {target}", timeout=45)
            return _json({"ok": True, "mode": "pathping", "host": target, "output": out[:3000]}) if ok else self._error(out or "pathping failed")
        if mode_norm in {"nslookup", "dns_lookup"}:
            target = (host or "").strip() or "google.com"
            ok, out = _run_powershell(f"nslookup {target}", timeout=15)
            return _json({"ok": True, "mode": "nslookup", "host": target, "output": out[:2000]}) if ok else self._error(out or "nslookup failed")
        if mode_norm in {"netstat_active", "netstat"}:
            ok, out = _run_powershell("netstat -ano", timeout=15)
            return _json({"ok": True, "mode": "netstat_active", "output": out[:3000]}) if ok else self._error(out or "netstat failed")
        if mode_norm in {"display_dns", "dns_cache"}:
            ok, out = _run_powershell("ipconfig /displaydns", timeout=15)
            return _json({"ok": True, "mode": "display_dns", "output": out[:3000]}) if ok else self._error(out or "display dns failed")
        if mode_norm in {"getmac"}:
            ok, out = _run_powershell("getmac", timeout=12)
            return _json({"ok": True, "mode": "getmac", "output": out[:2000]}) if ok else self._error(out or "getmac failed")
        if mode_norm in {"arp_table"}:
            ok, out = _run_powershell("arp -a", timeout=12)
            return _json({"ok": True, "mode": "arp_table", "output": out[:2000]}) if ok else self._error(out or "arp table failed")
        if mode_norm in {"nbtstat_cache"}:
            ok, out = _run_powershell("nbtstat -c", timeout=12)
            return _json({"ok": True, "mode": "nbtstat_cache", "output": out[:2000]}) if ok else self._error(out or "nbtstat cache failed")
        if mode_norm in {"nbtstat_host"}:
            target = (host or "").strip() or "127.0.0.1"
            ok, out = _run_powershell(f"nbtstat -a {target}", timeout=15)
            return _json({"ok": True, "mode": "nbtstat_host", "host": target, "output": out[:2500]}) if ok else self._error(out or "nbtstat host failed")
        if mode_norm in {"net_view"}:
            ok, out = _run_powershell("net view", timeout=12)
            return _json({"ok": True, "mode": "net_view", "output": out[:2000]}) if ok else self._error(out or "net view failed")
        if mode_norm in {"netstat_binary"}:
            ok, out = _run_powershell("netstat -b", timeout=20)
            return _json({"ok": True, "mode": "netstat_binary", "output": out[:3000]}) if ok else self._error(out or "netstat -b failed")
        if mode_norm in {"wifi_profiles"}:
            ok, out = _run_powershell("netsh wlan show profiles", timeout=12)
            return _json({"ok": True, "mode": "wifi_profiles", "output": out[:2500]}) if ok else self._error(out or "wifi profiles failed")
        if mode_norm in {"net_scan", "arp_scan"}:
            ok, out = _run_powershell("arp -a", timeout=10)
            return _json({"ok": True, "mode": "net_scan", "output": out[:2000]}) if ok else self._error(out or "net scan failed")
        if mode_norm in {"hotspot_on", "hotspot_off"}:
            state = "$true" if mode_norm.endswith("_on") else "$false"
            ps = (
                "$conn='Local Area Connection* 10'; "
                "$ad=Get-NetAdapter -Name $conn -ErrorAction SilentlyContinue; "
                f"if($ad){{ Set-NetAdapter -Name $conn -AdminStatus {'Up' if mode_norm.endswith('_on') else 'Down'} -Confirm:$false }}; "
                f"@{{ok=$true; mode='{mode_norm}'}} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=15)
            return out if ok and out else self._error(out or "hotspot toggle failed")
        if mode_norm in {"file_sharing_on", "file_sharing_off"}:
            enabled = "$true" if mode_norm.endswith("_on") else "$false"
            ps = (
                f"Set-NetFirewallRule -DisplayGroup 'File and Printer Sharing' -Enabled {enabled}; "
                f"@{{ok=$true; mode='{mode_norm}'}} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=15)
            return out if ok and out else self._error(out or "file sharing toggle failed")
        if mode_norm in {"shared_folders"}:
            ps = "Get-SmbShare | Select-Object Name,Path,Description | ConvertTo-Json -Compress"
            ok, out = _run_powershell(ps, timeout=12)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "shared folders failed")
        if mode_norm in {"server_online"}:
            target = (host or "").strip() or "8.8.8.8"
            ps = f"Test-Connection -ComputerName '{target}' -Count 1 -Quiet | ConvertTo-Json -Compress"
            ok, out = _run_powershell(ps, timeout=12)
            if ok and out:
                try:
                    status = json.loads(out)
                except Exception:
                    status = out.strip()
                return _json({"ok": True, "mode": mode_norm, "host": target, "online": status})
            return self._error(out or "server online check failed")
        if mode_norm in {"last_login_events"}:
            ps = (
                "Get-WinEvent -FilterHashtable @{LogName='Security'; Id=4624,4625} -MaxEvents 20 | "
                "Select-Object Id,TimeCreated,Message | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=18)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "last login events failed")
        if mode_norm in {"open_settings", "settings"}:
            return self._open_settings_page("network")
        if mode_norm in {"open_ports", "ports"}:
            ps = (
                "Get-NetTCPConnection -State Listen | "
                "Select-Object -First 120 LocalAddress,LocalPort,OwningProcess | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=12)
            if ok and out:
                return _json({"ok": True, "mode": "open_ports", "data": json.loads(out)})
            return self._error(out or "open ports failed")
        if mode_norm in {"port_owner", "who_uses_port"}:
            try:
                p = int(port if port is not None else (host or "").strip())
            except Exception:
                return self._error("valid port is required")
            ps = (
                f"$rows=Get-NetTCPConnection -LocalPort {p} -ErrorAction SilentlyContinue | "
                "Select-Object LocalAddress,LocalPort,RemoteAddress,RemotePort,State,OwningProcess; "
                "$out=@(); "
                "foreach($r in $rows){ "
                " $proc=Get-Process -Id $r.OwningProcess -ErrorAction SilentlyContinue; "
                " $out += [pscustomobject]@{"
                "local_address=$r.LocalAddress;local_port=$r.LocalPort;remote_address=$r.RemoteAddress;remote_port=$r.RemotePort;"
                "state=$r.State;pid=$r.OwningProcess;process_name=($proc.ProcessName);path=($proc.Path)} "
                "}; "
                "$out | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=15)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "port": p, "data": json.loads(out)})
            return self._error(out or f"port owner lookup failed for port {p}")
        return self._error(f"unsupported network_tools mode: {mode_norm}")

    def _file_tools(
        self,
        mode: str,
        path: str = "",
        target: str = "",
        name: str = "",
        pattern: str = "",
        ext: str = "",
        permanent: bool = False,
    ) -> str:
        mode_norm = (mode or "").strip().lower()
        home = Path.home()
        quick_dirs = {
            "open_documents": home / "Documents",
            "open_downloads": home / "Downloads",
            "open_pictures": home / "Pictures",
            "open_videos": home / "Videos",
        }
        if mode_norm in quick_dirs:
            p = quick_dirs[mode_norm]
            if p.exists():
                os.startfile(str(p))  # type: ignore[attr-defined]
                return _json({"ok": True, "mode": mode_norm, "path": str(p)})
            return self._error(f"path not found: {p}")

        base_path = Path(path).expanduser() if path else Path.cwd()
        if mode_norm == "organize_desktop":
            desktop = Path.home() / "Desktop"
            if not desktop.exists():
                return self._error(f"desktop not found: {desktop}")
            buckets: dict[str, set[str]] = {
                "Documents": {".pdf", ".doc", ".docx", ".txt", ".rtf", ".ppt", ".pptx", ".xls", ".xlsx"},
                "Images": {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".svg"},
                "Videos": {".mp4", ".mkv", ".avi", ".mov", ".wmv", ".webm"},
                "Archives": {".zip", ".rar", ".7z", ".tar", ".gz"},
                "Code": {".py", ".js", ".ts", ".tsx", ".java", ".cpp", ".c", ".go", ".rs", ".ps1", ".sh"},
            }
            moved: list[dict[str, str]] = []
            for item in desktop.iterdir():
                try:
                    if not item.is_file():
                        continue
                    ext = item.suffix.lower()
                    target_bucket = "Others"
                    for bucket, exts in buckets.items():
                        if ext in exts:
                            target_bucket = bucket
                            break
                    target_dir = desktop / target_bucket
                    target_dir.mkdir(parents=True, exist_ok=True)
                    target = target_dir / item.name
                    if target.exists():
                        stem = item.stem
                        suffix = item.suffix
                        target = target_dir / f"{stem}_{_timestamp_id()}{suffix}"
                    item.rename(target)
                    moved.append({"from": str(item), "to": str(target)})
                except Exception:
                    continue
            return _json({"ok": True, "mode": mode_norm, "moved": len(moved), "items": moved[:200]})
        if mode_norm in {"organize_desktop_semantic", "semantic_organize_desktop"}:
            desktop = Path.home() / "Desktop"
            if not desktop.exists():
                return self._error(f"desktop not found: {desktop}")

            text_like_exts = {
                ".txt",
                ".md",
                ".csv",
                ".log",
                ".json",
                ".xml",
                ".html",
                ".htm",
                ".docx",
                ".pdf",
            }
            media_exts = {
                ".png",
                ".jpg",
                ".jpeg",
                ".webp",
                ".gif",
                ".bmp",
                ".svg",
                ".mp4",
                ".mkv",
                ".avi",
                ".mov",
                ".wmv",
                ".webm",
            }
            archive_exts = {".zip", ".rar", ".7z", ".tar", ".gz"}
            code_exts = {".py", ".js", ".ts", ".tsx", ".cpp", ".c", ".java", ".go", ".rs", ".ps1", ".sh"}

            groups = {
                "Finance": ("invoice", "receipt", "bill", "payment", "bank", "فاتورة", "إيصال", "دفع"),
                "Work": ("project", "meeting", "client", "proposal", "contract", "مشروع", "عميل", "عقد", "عمل"),
                "Study": ("course", "lecture", "homework", "assignment", "exam", "study", "دراسة", "محاضرة", "اختبار"),
                "Family": ("family", "birthday", "wedding", "vacation", "photo", "عائلة", "زواج", "رحلة", "ذكريات"),
            }

            def _guess_semantic_bucket(item: Path) -> str:
                ext = item.suffix.lower()
                if ext in media_exts:
                    return "Media"
                if ext in archive_exts:
                    return "Archives"
                if ext in code_exts:
                    return "Code"

                raw_text = item.stem
                if ext in text_like_exts:
                    try:
                        if ext == ".docx":
                            from docx import Document

                            doc = Document(item)
                            raw_text = "\n".join(p.text for p in doc.paragraphs if p.text)[:3000]
                        elif ext == ".pdf":
                            raw_text = item.stem
                        else:
                            raw_text = item.read_text(encoding="utf-8", errors="ignore")[:3000]
                    except Exception:
                        raw_text = item.stem
                low = str(raw_text or "").casefold()
                for bucket, keywords in groups.items():
                    if any(kw.casefold() in low for kw in keywords):
                        return bucket
                return "Others"

            moved: list[dict[str, str]] = []
            for item in desktop.iterdir():
                try:
                    if not item.is_file():
                        continue
                    bucket = _guess_semantic_bucket(item)
                    target_dir = desktop / bucket
                    target_dir.mkdir(parents=True, exist_ok=True)
                    target = target_dir / item.name
                    if target.exists():
                        target = target_dir / f"{item.stem}_{_timestamp_id()}{item.suffix}"
                    item.rename(target)
                    moved.append({"from": str(item), "to": str(target), "bucket": bucket})
                except Exception:
                    continue
            return _json({"ok": True, "mode": mode_norm, "moved": len(moved), "items": moved[:200]})
        if mode_norm == "smart_rename":
            folder = Path(path).expanduser() if path else Path.cwd()
            if not folder.exists() or not folder.is_dir():
                return self._error(f"folder not found: {folder}")
            prefix = (name or pattern or "file").strip()
            files = [p for p in folder.iterdir() if p.is_file()]
            files.sort(key=lambda p: p.stat().st_mtime)
            changed: list[dict[str, str]] = []
            for idx, item in enumerate(files, start=1):
                try:
                    ts = datetime.fromtimestamp(item.stat().st_mtime).strftime("%Y%m%d")
                    new_name = f"{prefix}_{ts}_{idx:03d}{item.suffix.lower()}"
                    target = item.with_name(new_name)
                    if target.exists():
                        target = item.with_name(f"{prefix}_{ts}_{idx:03d}_{_timestamp_id()}{item.suffix.lower()}")
                    item.rename(target)
                    changed.append({"from": str(item), "to": str(target)})
                except Exception:
                    continue
            return _json({"ok": True, "mode": mode_norm, "renamed": len(changed), "items": changed[:200]})
        if mode_norm in {"smart_rename_content", "content_rename"}:
            folder = Path(path).expanduser() if path else Path.cwd()
            if not folder.exists() or not folder.is_dir():
                return self._error(f"folder not found: {folder}")
            files = [p for p in folder.iterdir() if p.is_file()]
            files.sort(key=lambda p: p.stat().st_mtime)
            files = files[:100]

            text_like_exts = {
                ".txt",
                ".md",
                ".log",
                ".csv",
                ".json",
                ".xml",
                ".html",
                ".htm",
                ".ini",
                ".cfg",
                ".yaml",
                ".yml",
                ".py",
                ".js",
                ".ts",
            }

            def _slug(value: str) -> str:
                base = re.sub(r"[^\w\s-]", " ", value or "", flags=re.UNICODE)
                base = re.sub(r"[_\s-]+", "_", base).strip("_")
                if not base:
                    return "file"
                return base[:60]

            def _extract_date_token(text_val: str) -> str:
                patterns = (
                    r"\b(20\d{2})[-/](\d{1,2})[-/](\d{1,2})\b",
                    r"\b(\d{1,2})[-/](\d{1,2})[-/](20\d{2})\b",
                )
                for pat in patterns:
                    m = re.search(pat, text_val)
                    if not m:
                        continue
                    g = m.groups()
                    try:
                        if len(g[0]) == 4:
                            y, mn, d = int(g[0]), int(g[1]), int(g[2])
                        else:
                            d, mn, y = int(g[0]), int(g[1]), int(g[2])
                        if 1 <= mn <= 12 and 1 <= d <= 31:
                            return f"{y:04d}{mn:02d}{d:02d}"
                    except Exception:
                        continue
                return ""

            changed: list[dict[str, str]] = []
            for idx, item in enumerate(files, start=1):
                try:
                    suffix = item.suffix.lower()
                    stem_hint = item.stem
                    text_hint = ""
                    if suffix in text_like_exts:
                        try:
                            text_hint = item.read_text(encoding="utf-8", errors="ignore")[:4000]
                        except Exception:
                            text_hint = ""
                    candidate = text_hint or stem_hint
                    first_line = ""
                    for ln in candidate.splitlines():
                        ln = ln.strip()
                        if ln:
                            first_line = ln
                            break
                    low = candidate.casefold()
                    category = ""
                    if any(tok in low for tok in ("invoice", "bill", "receipt", "فاتورة")):
                        category = "invoice"
                    elif any(tok in low for tok in ("report", "summary", "ملخص", "تقرير")):
                        category = "report"
                    elif any(tok in low for tok in ("contract", "agreement", "عقد")):
                        category = "contract"
                    elif any(tok in low for tok in ("resume", "cv", "سيرة")):
                        category = "resume"
                    else:
                        words = re.findall(r"[\w\u0600-\u06FF]+", first_line or "")
                        category = "_".join(words[:6]) if words else "file"
                    category = _slug(category)
                    date_token = _extract_date_token(candidate)
                    if date_token:
                        new_name = f"{category}_{date_token}{suffix}"
                    else:
                        new_name = f"{category}_{idx:03d}{suffix}"
                    target = item.with_name(new_name)
                    if target.exists() and target != item:
                        target = item.with_name(f"{category}_{_timestamp_id()}_{idx:03d}{suffix}")
                    if target != item:
                        item.rename(target)
                        changed.append({"from": str(item), "to": str(target)})
                except Exception:
                    continue
            return _json({"ok": True, "mode": mode_norm, "renamed": len(changed), "items": changed[:200]})
        if mode_norm == "create_folder":
            folder_name = (name or target or "").strip()
            if not folder_name:
                return self._error("name is required")
            out = base_path / folder_name
            out.mkdir(parents=True, exist_ok=True)
            return _json({"ok": True, "mode": "create_folder", "path": str(out)})

        if mode_norm == "delete":
            victim = Path(target or path).expanduser()
            if not victim.exists():
                return self._error(f"target not found: {victim}")
            if permanent:
                if victim.is_dir():
                    import shutil

                    shutil.rmtree(victim)
                else:
                    victim.unlink()
                return _json({"ok": True, "mode": "delete", "permanent": True, "target": str(victim)})
            try:
                from send2trash import send2trash

                send2trash(str(victim))
            except Exception:
                return self._error("send2trash is not available")
            return _json({"ok": True, "mode": "delete", "permanent": False, "target": str(victim)})

        if mode_norm == "rename":
            src = Path(path).expanduser()
            new_name = (name or "").strip()
            if not src.exists():
                return self._error(f"path not found: {src}")
            if not new_name:
                return self._error("name is required")
            dst = src.with_name(new_name)
            src.rename(dst)
            return _json({"ok": True, "mode": "rename", "from": str(src), "to": str(dst)})
        if mode_norm == "copy":
            src = Path(path).expanduser()
            dst = Path(target).expanduser()
            if not src.exists():
                return self._error(f"path not found: {src}")
            if not target:
                return self._error("target is required")
            import shutil

            if src.is_dir():
                shutil.copytree(src, dst, dirs_exist_ok=True)
            else:
                if dst.is_dir():
                    dst = dst / src.name
                shutil.copy2(src, dst)
            return _json({"ok": True, "mode": "copy", "from": str(src), "to": str(dst)})
        if mode_norm == "move":
            src = Path(path).expanduser()
            dst = Path(target).expanduser()
            if not src.exists():
                return self._error(f"path not found: {src}")
            if not target:
                return self._error("target is required")
            import shutil

            moved = shutil.move(str(src), str(dst))
            return _json({"ok": True, "mode": "move", "from": str(src), "to": str(moved)})

        if mode_norm == "zip":
            src = Path(path).expanduser()
            if not src.exists():
                return self._error(f"path not found: {src}")
            import shutil

            archive = shutil.make_archive(str(src), "zip", root_dir=str(src.parent), base_dir=src.name)
            return _json({"ok": True, "mode": "zip", "archive": archive})

        if mode_norm == "unzip":
            zpath = Path(path).expanduser()
            if not zpath.exists():
                return self._error(f"zip not found: {zpath}")
            import zipfile

            out_dir = zpath.with_suffix("")
            with zipfile.ZipFile(zpath, "r") as zf:
                zf.extractall(out_dir)
            return _json({"ok": True, "mode": "unzip", "path": str(zpath), "output": str(out_dir)})

        if mode_norm == "search_ext":
            wanted_ext = (ext or pattern or "").strip()
            if not wanted_ext:
                return self._error("ext is required")
            if not wanted_ext.startswith("."):
                wanted_ext = f".{wanted_ext}"
            matches: list[str] = []
            for root in _iter_search_roots():
                for hit in root.rglob(f"*{wanted_ext}"):
                    if hit.is_file():
                        matches.append(str(hit))
                        if len(matches) >= 100:
                            break
                if len(matches) >= 100:
                    break
            return _json({"ok": True, "mode": "search_ext", "ext": wanted_ext, "count": len(matches), "items": matches})

        if mode_norm == "folder_size":
            folder = Path(path).expanduser()
            if not folder.exists() or not folder.is_dir():
                return self._error(f"folder not found: {folder}")
            total = 0
            for child in folder.rglob("*"):
                if child.is_file():
                    try:
                        total += child.stat().st_size
                    except Exception:
                        pass
            return _json({"ok": True, "mode": "folder_size", "path": str(folder), "bytes": total})

        if mode_norm == "open_cmd_here":
            folder = Path(path).expanduser() if path else Path.cwd()
            if not folder.exists():
                return self._error(f"path not found: {folder}")
            subprocess.Popen(["cmd", "/k", "cd", "/d", str(folder)])
            return _json({"ok": True, "mode": "open_cmd_here", "path": str(folder)})

        if mode_norm == "open_powershell_here":
            folder = Path(path).expanduser() if path else Path.cwd()
            if not folder.exists():
                return self._error(f"path not found: {folder}")
            subprocess.Popen(["powershell", "-NoExit", "-Command", f"Set-Location -Path '{folder}'"])
            return _json({"ok": True, "mode": "open_powershell_here", "path": str(folder)})

        if mode_norm == "empty_recycle_bin":
            ok, out = _run_powershell("Clear-RecycleBin -Force -ErrorAction Stop", timeout=15)
            if ok:
                return _json({"ok": True, "mode": "empty_recycle_bin"})
            return self._error(out or "empty recycle bin failed")
        if mode_norm in {"show_hidden", "hide_hidden"}:
            hidden_value = 1 if mode_norm == "show_hidden" else 2
            super_hidden = 1 if mode_norm == "show_hidden" else 0
            ps = (
                "$k='HKCU:\\Software\\Microsoft\\Windows\\CurrentVersion\\Explorer\\Advanced'; "
                f"Set-ItemProperty -Path $k -Name Hidden -Value {hidden_value}; "
                f"Set-ItemProperty -Path $k -Name ShowSuperHidden -Value {super_hidden}; "
                "Stop-Process -Name explorer -Force -ErrorAction SilentlyContinue; "
                "Start-Process explorer.exe; "
                f"@{{ok=$true; mode='{mode_norm}'}} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=20)
            return out if ok and out else self._error(out or f"{mode_norm} failed")

        return self._error(f"unsupported file_tools mode: {mode_norm}")

    def _window_control(
        self,
        mode: str,
        app: str = "",
        x: Any = None,
        y: Any = None,
        width: Any = None,
        height: Any = None,
        opacity: Any = None,
        name: str = "",
        text: str = "",
    ) -> str:
        mode_norm = (mode or "").strip().lower()
        try:
            import pyautogui
        except Exception as exc:
            return self._error(f"pyautogui unavailable: {exc}")
        pyautogui.FAILSAFE = False
        try:
            if mode_norm in {"show_desktop", "show_desktop_verified"}:
                before_count = None
                try:
                    import pygetwindow as gw

                    before_count = len(
                        [
                            w
                            for w in gw.getAllWindows()
                            if str(getattr(w, "title", "") or "").strip() and bool(getattr(w, "isVisible", True))
                        ]
                    )
                except Exception:
                    before_count = None
                pyautogui.hotkey("win", "d")
                time.sleep(0.22)
                after_count = None
                try:
                    import pygetwindow as gw

                    after_count = len(
                        [
                            w
                            for w in gw.getAllWindows()
                            if str(getattr(w, "title", "") or "").strip() and bool(getattr(w, "isVisible", True))
                        ]
                    )
                except Exception:
                    after_count = None
                verified = True
                fallback_used = False
                if before_count is not None and after_count is not None and after_count >= before_count:
                    ok, _out = _run_powershell(
                        "(New-Object -ComObject Shell.Application).MinimizeAll(); @{ok=$true; mode='show_desktop_verified'; fallback='shell.minimize_all'} | ConvertTo-Json -Compress",
                        timeout=10,
                    )
                    fallback_used = bool(ok)
                    time.sleep(0.18)
                    try:
                        import pygetwindow as gw

                        final_count = len(
                            [
                                w
                                for w in gw.getAllWindows()
                                if str(getattr(w, "title", "") or "").strip() and bool(getattr(w, "isVisible", True))
                            ]
                        )
                        verified = final_count < before_count if before_count is not None else bool(ok)
                        after_count = final_count
                    except Exception:
                        verified = bool(ok)
                return _json(
                    {
                        "ok": bool(verified),
                        "mode": "show_desktop_verified",
                        "before_visible_windows": before_count,
                        "after_visible_windows": after_count,
                        "fallback_used": fallback_used,
                        "error": None if verified else "show desktop verification failed",
                    }
                )
            if mode_norm == "undo_show_desktop":
                pyautogui.hotkey("win", "d")
                return _json({"ok": True, "mode": mode_norm})
            if mode_norm == "close_current":
                pyautogui.hotkey("alt", "f4")
                return _json({"ok": True, "mode": mode_norm})
            if mode_norm == "alt_tab":
                pyautogui.hotkey("alt", "tab")
                return _json({"ok": True, "mode": mode_norm})
            if mode_norm == "task_view":
                pyautogui.hotkey("win", "tab")
                return _json({"ok": True, "mode": mode_norm})
            if mode_norm == "move_next_monitor_right":
                pyautogui.hotkey("win", "shift", "right")
                return _json({"ok": True, "mode": mode_norm})
            if mode_norm == "move_next_monitor_left":
                pyautogui.hotkey("win", "shift", "left")
                return _json({"ok": True, "mode": mode_norm})
            if mode_norm == "split_left":
                pyautogui.hotkey("win", "left")
                return _json({"ok": True, "mode": mode_norm})
            if mode_norm == "split_right":
                pyautogui.hotkey("win", "right")
                return _json({"ok": True, "mode": mode_norm})
            if mode_norm == "project_panel":
                pyautogui.hotkey("win", "p")
                return _json({"ok": True, "mode": mode_norm})
            if mode_norm in {"display_duplicate", "display_extend", "display_internal", "display_external"}:
                arg_map = {
                    "display_duplicate": "/clone",
                    "display_extend": "/extend",
                    "display_internal": "/internal",
                    "display_external": "/external",
                }
                arg = arg_map[mode_norm]
                ok, out = _run_powershell(f"Start-Process DisplaySwitch.exe -ArgumentList '{arg}'", timeout=10)
                return _json({"ok": True, "mode": mode_norm, "arg": arg}) if ok else self._error(out or f"{mode_norm} failed")
            if mode_norm == "aero_shake":
                pyautogui.hotkey("win", "home")
                return _json({"ok": True, "mode": mode_norm})
            if mode_norm in {"always_on_top_on", "always_on_top_off"}:
                flag = "-1" if mode_norm.endswith("_on") else "-2"
                ps = (
                    "$sig='[DllImport(\"user32.dll\")]public static extern IntPtr GetForegroundWindow();"
                    "[DllImport(\"user32.dll\")]public static extern bool SetWindowPos(IntPtr hWnd,IntPtr hWndInsertAfter,int X,int Y,int cx,int cy,uint uFlags);'; "
                    "Add-Type -Name NativeMethods -Namespace Win32 -MemberDefinition $sig; "
                    "$h=[Win32.NativeMethods]::GetForegroundWindow(); "
                    f"[Win32.NativeMethods]::SetWindowPos($h,[intptr]{flag},0,0,0,0,0x0003) | Out-Null; "
                    f"@{{ok=$true; mode='{mode_norm}'}} | ConvertTo-Json -Compress"
                )
                ok, out = _run_powershell(ps, timeout=12)
                return out if ok and out else self._error(out or "always on top toggle failed")
            if mode_norm in {
                "minimize",
                "maximize",
                "restore",
                "bring_to_front",
                "set_focus",
                "hide",
                "show",
                "coords",
                "move_resize",
                "transparency",
                "borderless_on",
                "borderless_off",
                "disable_close_on",
                "disable_close_off",
                "span_all_screens",
                "minimize_to_tray",
                "restore_from_tray",
                "rename_title",
            }:
                import pygetwindow as gw

                wins = gw.getAllWindows()
                target = None
                if app:
                    app_norm = app.casefold()
                    for w in wins:
                        if app_norm in str(getattr(w, "title", "") or "").casefold():
                            target = w
                            break
                if target is None:
                    target = gw.getActiveWindow()
                if target is None:
                    return self._error("no target window found")
                if mode_norm == "minimize":
                    target.minimize()
                elif mode_norm == "maximize":
                    target.maximize()
                elif mode_norm == "restore":
                    target.restore()
                elif mode_norm in {"bring_to_front", "set_focus"}:
                    target.activate()
                elif mode_norm in {"hide", "minimize_to_tray"}:
                    target.minimize()
                    if mode_norm == "minimize_to_tray":
                        try:
                            hwnd_val = int(getattr(target, "_hWnd", 0) or 0)
                            if hwnd_val > 0:
                                ps = (
                                    "$sig='[DllImport(\"user32.dll\")]public static extern bool ShowWindow(IntPtr hWnd,int nCmdShow);'; "
                                    "Add-Type -Name NativeMethods -Namespace Win32 -MemberDefinition $sig; "
                                    f"$h=[intptr]{hwnd_val}; "
                                    "[void][Win32.NativeMethods]::ShowWindow($h,0); "
                                    "@{ok=$true; mode='minimize_to_tray'} | ConvertTo-Json -Compress"
                                )
                                ok, out = _run_powershell(ps, timeout=10)
                                if ok and out:
                                    return out
                        except Exception:
                            pass
                elif mode_norm in {"show", "restore_from_tray"}:
                    target.restore()
                    target.activate()
                elif mode_norm == "coords":
                    return _json(
                        {
                            "ok": True,
                            "mode": mode_norm,
                            "title": str(getattr(target, "title", "") or ""),
                            "x": int(getattr(target, "left", 0) or 0),
                            "y": int(getattr(target, "top", 0) or 0),
                            "width": int(getattr(target, "width", 0) or 0),
                            "height": int(getattr(target, "height", 0) or 0),
                        }
                    )
                elif mode_norm == "rename_title":
                    new_title = (text or name or "").strip()
                    if not new_title:
                        return self._error("text or name is required for rename_title")
                    hwnd_val = int(getattr(target, "_hWnd", 0) or 0)
                    if hwnd_val <= 0:
                        return self._error("window handle unavailable for rename_title")
                    esc_title = new_title.replace("'", "''")
                    ps = (
                        "$sig='[DllImport(\"user32.dll\", CharSet = CharSet.Unicode)]public static extern bool SetWindowText(IntPtr hWnd,string lpString);'; "
                        "Add-Type -Name NativeMethods -Namespace Win32 -MemberDefinition $sig; "
                        f"$h=[intptr]{hwnd_val}; "
                        f"[void][Win32.NativeMethods]::SetWindowText($h,'{esc_title}'); "
                        f"@{{ok=$true; mode='rename_title'; title='{esc_title}'}} | ConvertTo-Json -Compress"
                    )
                    ok, out = _run_powershell(ps, timeout=10)
                    return out if ok and out else self._error(out or "rename title failed")
                elif mode_norm == "move_resize":
                    if x is None or y is None:
                        return self._error("x and y are required for move_resize")
                    tx = int(float(x))
                    ty = int(float(y))
                    tw = int(float(width)) if width is not None else int(getattr(target, "width", 0) or 0)
                    th = int(float(height)) if height is not None else int(getattr(target, "height", 0) or 0)
                    tw = max(160, tw)
                    th = max(90, th)
                    target.moveTo(tx, ty)
                    target.resizeTo(tw, th)
                    return _json(
                        {
                            "ok": True,
                            "mode": mode_norm,
                            "title": str(getattr(target, "title", "") or ""),
                            "x": tx,
                            "y": ty,
                            "width": tw,
                            "height": th,
                        }
                    )
                elif mode_norm == "transparency":
                    if opacity is None:
                        return self._error("opacity is required (10..100)")
                    alpha_percent = _clamp(int(float(opacity)), 10, 100)
                    hwnd_val = int(getattr(target, "_hWnd", 0) or 0)
                    if hwnd_val <= 0:
                        return self._error("window handle unavailable for transparency")
                    ps = (
                        "$sig='[DllImport(\"user32.dll\")]public static extern int GetWindowLong(IntPtr hWnd,int nIndex);"
                        "[DllImport(\"user32.dll\")]public static extern int SetWindowLong(IntPtr hWnd,int nIndex,int dwNewLong);"
                        "[DllImport(\"user32.dll\")]public static extern bool SetLayeredWindowAttributes(IntPtr hwnd,uint crKey,byte bAlpha,uint dwFlags);'; "
                        "Add-Type -Name NativeMethods -Namespace Win32 -MemberDefinition $sig; "
                        f"$h=[intptr]{hwnd_val}; "
                        "$ex=[Win32.NativeMethods]::GetWindowLong($h,-20); "
                        "[void][Win32.NativeMethods]::SetWindowLong($h,-20,($ex -bor 0x80000)); "
                        f"$alpha=[byte][math]::Round({alpha_percent}*2.55); "
                        "[void][Win32.NativeMethods]::SetLayeredWindowAttributes($h,0,$alpha,2); "
                        f"@{{ok=$true; mode='transparency'; opacity={alpha_percent}}} | ConvertTo-Json -Compress"
                    )
                    ok, out = _run_powershell(ps, timeout=12)
                    return out if ok and out else self._error(out or "set transparency failed")
                elif mode_norm in {"borderless_on", "borderless_off"}:
                    hwnd_val = int(getattr(target, "_hWnd", 0) or 0)
                    if hwnd_val <= 0:
                        return self._error("window handle unavailable for borderless mode")
                    if mode_norm.endswith("_on"):
                        ps = (
                            "$sig='[DllImport(\"user32.dll\")]public static extern int GetWindowLong(IntPtr hWnd,int nIndex);"
                            "[DllImport(\"user32.dll\")]public static extern int SetWindowLong(IntPtr hWnd,int nIndex,int dwNewLong);'; "
                            "Add-Type -Name NativeMethods -Namespace Win32 -MemberDefinition $sig; "
                            f"$h=[intptr]{hwnd_val}; "
                            "$style=[Win32.NativeMethods]::GetWindowLong($h,-16); "
                            "$style=$style -band (-bnot 0x00C00000) -band (-bnot 0x00040000); "
                            "[void][Win32.NativeMethods]::SetWindowLong($h,-16,$style); "
                            "@{ok=$true; mode='borderless_on'} | ConvertTo-Json -Compress"
                        )
                    else:
                        ps = (
                            "$sig='[DllImport(\"user32.dll\")]public static extern int GetWindowLong(IntPtr hWnd,int nIndex);"
                            "[DllImport(\"user32.dll\")]public static extern int SetWindowLong(IntPtr hWnd,int nIndex,int dwNewLong);'; "
                            "Add-Type -Name NativeMethods -Namespace Win32 -MemberDefinition $sig; "
                            f"$h=[intptr]{hwnd_val}; "
                            "$style=[Win32.NativeMethods]::GetWindowLong($h,-16); "
                            "$style=$style -bor 0x00C00000 -bor 0x00040000; "
                            "[void][Win32.NativeMethods]::SetWindowLong($h,-16,$style); "
                            "@{ok=$true; mode='borderless_off'} | ConvertTo-Json -Compress"
                        )
                    ok, out = _run_powershell(ps, timeout=12)
                    return out if ok and out else self._error(out or "borderless mode failed")
                elif mode_norm in {"disable_close_on", "disable_close_off"}:
                    hwnd_val = int(getattr(target, "_hWnd", 0) or 0)
                    if hwnd_val <= 0:
                        return self._error("window handle unavailable for close-button control")
                    ps = (
                        "$sig='[DllImport(\"user32.dll\")]public static extern IntPtr GetSystemMenu(IntPtr hWnd,bool bRevert);"
                        "[DllImport(\"user32.dll\")]public static extern bool DeleteMenu(IntPtr hMenu,uint uPosition,uint uFlags);"
                        "[DllImport(\"user32.dll\")]public static extern bool EnableMenuItem(IntPtr hMenu,uint uIDEnableItem,uint uEnable);'; "
                        "Add-Type -Name NativeMethods -Namespace Win32 -MemberDefinition $sig; "
                        f"$h=[intptr]{hwnd_val}; "
                        "$m=[Win32.NativeMethods]::GetSystemMenu($h,$false); "
                    )
                    if mode_norm.endswith("_on"):
                        ps += (
                            "[void][Win32.NativeMethods]::DeleteMenu($m,0xF060,0x0); "
                            "@{ok=$true; mode='disable_close_on'} | ConvertTo-Json -Compress"
                        )
                    else:
                        ps += (
                            "$m=[Win32.NativeMethods]::GetSystemMenu($h,$true); "
                            "@{ok=$true; mode='disable_close_off'} | ConvertTo-Json -Compress"
                        )
                    ok, out = _run_powershell(ps, timeout=12)
                    return out if ok and out else self._error(out or "close button toggle failed")
                elif mode_norm == "span_all_screens":
                    ps = (
                        "$sig='[DllImport(\"user32.dll\")]public static extern int GetSystemMetrics(int nIndex);'; "
                        "Add-Type -Name NativeMethods -Namespace Win32 -MemberDefinition $sig; "
                        "$x=[Win32.NativeMethods]::GetSystemMetrics(76); "
                        "$y=[Win32.NativeMethods]::GetSystemMetrics(77); "
                        "$w=[Win32.NativeMethods]::GetSystemMetrics(78); "
                        "$h=[Win32.NativeMethods]::GetSystemMetrics(79); "
                        "@{x=$x;y=$y;width=$w;height=$h} | ConvertTo-Json -Compress"
                    )
                    ok, out = _run_powershell(ps, timeout=10)
                    if not ok or not out:
                        return self._error(out or "failed to get virtual screen bounds")
                    try:
                        b = json.loads(out)
                        tx = int(b.get("x", 0))
                        ty = int(b.get("y", 0))
                        tw = max(100, int(b.get("width", 0)))
                        th = max(100, int(b.get("height", 0)))
                        target.restore()
                        target.moveTo(tx, ty)
                        target.resizeTo(tw, th)
                        return _json(
                            {
                                "ok": True,
                                "mode": mode_norm,
                                "title": str(getattr(target, "title", "") or ""),
                                "x": tx,
                                "y": ty,
                                "width": tw,
                                "height": th,
                            }
                        )
                    except Exception as exc:
                        return self._error(f"span all screens parse failed: {exc}")
                return _json({"ok": True, "mode": mode_norm, "title": str(getattr(target, "title", "") or "")})
        except Exception as exc:
            return self._error(f"window control failed: {exc}")
        return self._error(f"unsupported window_control mode: {mode_norm}")

    def _process_tools(
        self,
        mode: str,
        pid: Any = None,
        name: str = "",
        other_name: str = "",
        dry_run: bool = False,
        max_kill: Any = None,
        resource: str = "",
        stage: str = "",
        priority: str = "",
        threshold: Any = None,
    ) -> str:
        mode_norm = (mode or "").strip().lower()
        resource_norm = (resource or "").strip().lower()
        stage_norm = (stage or "").strip().lower()
        try:
            import psutil
        except Exception as exc:
            return self._error(f"psutil unavailable: {exc}")

        if mode_norm == "app_reduce":
            if resource_norm not in {"ram", "cpu", "disk", "network"}:
                return self._error("resource must be ram|cpu|disk|network")
            if stage_norm not in {"plan", "execute"}:
                stage_norm = "execute" if not bool(dry_run) else "plan"
            routed_mode = f"app_reduce_{resource_norm}_{stage_norm}"
            return self._process_tools(
                mode=routed_mode,
                pid=pid,
                name=name,
                other_name=other_name,
                dry_run=dry_run,
                max_kill=max_kill,
                resource=resource,
                stage=stage,
                priority=priority,
                threshold=threshold,
            )

        if mode_norm == "list":
            items = []
            for p in psutil.process_iter(["pid", "name", "cpu_percent", "memory_info"]):
                try:
                    info = p.info
                    items.append(
                        {
                            "pid": info.get("pid"),
                            "name": info.get("name"),
                            "cpu": info.get("cpu_percent", 0.0),
                            "ram_mb": round((info.get("memory_info").rss or 0) / (1024 * 1024), 2)
                            if info.get("memory_info")
                            else 0.0,
                        }
                    )
                except Exception:
                    continue
            return _json({"ok": True, "mode": "list", "count": len(items), "items": items[:120]})

        if mode_norm in {"top_cpu", "top_ram"}:
            items = []
            for p in psutil.process_iter(["pid", "name", "cpu_percent", "memory_info"]):
                try:
                    info = p.info
                    items.append(
                        {
                            "pid": info.get("pid"),
                            "name": info.get("name"),
                            "cpu": info.get("cpu_percent", 0.0),
                            "ram_mb": round((info.get("memory_info").rss or 0) / (1024 * 1024), 2)
                            if info.get("memory_info")
                            else 0.0,
                        }
                    )
                except Exception:
                    continue
            key = "cpu" if mode_norm == "top_cpu" else "ram_mb"
            items.sort(key=lambda x: float(x.get(key, 0.0) or 0.0), reverse=True)
            return _json({"ok": True, "mode": mode_norm, "items": items[:10]})
        if mode_norm == "app_memory_total":
            query_raw = (name or "").strip()
            if not query_raw:
                return self._error("name is required")
            query_norm = query_raw.casefold()
            if query_norm.endswith(".exe"):
                query_norm = query_norm[:-4]
            items = []
            total_ram_mb = 0.0
            for p in psutil.process_iter(["pid", "name", "memory_info"]):
                try:
                    info = p.info
                    pname = str(info.get("name") or "").strip()
                    if not pname:
                        continue
                    pname_norm = pname.casefold()
                    pname_no_ext = pname_norm[:-4] if pname_norm.endswith(".exe") else pname_norm
                    if query_norm not in pname_norm and query_norm not in pname_no_ext:
                        continue
                    ram_mb = round((info.get("memory_info").rss or 0) / (1024 * 1024), 2) if info.get("memory_info") else 0.0
                    total_ram_mb += ram_mb
                    items.append(
                        {
                            "pid": info.get("pid"),
                            "name": pname,
                            "ram_mb": ram_mb,
                        }
                    )
                except Exception:
                    continue
            items.sort(key=lambda x: float(x.get("ram_mb", 0.0) or 0.0), reverse=True)
            return _json(
                {
                    "ok": True,
                    "mode": mode_norm,
                    "query": query_raw,
                    "process_count": len(items),
                    "total_ram_mb": round(total_ram_mb, 2),
                    "items": items[:30],
                }
            )
        if mode_norm == "app_process_count_total":
            query_raw = (name or "").strip()
            if not query_raw:
                return self._error("name is required")
            query_norm = query_raw.casefold()
            if query_norm.endswith(".exe"):
                query_norm = query_norm[:-4]
            items: list[dict[str, Any]] = []
            total_ram_mb = 0.0
            for p in psutil.process_iter(["pid", "name", "memory_info"]):
                try:
                    info = p.info
                    pname = str(info.get("name") or "").strip()
                    if not pname:
                        continue
                    pname_norm = pname.casefold()
                    pname_no_ext = pname_norm[:-4] if pname_norm.endswith(".exe") else pname_norm
                    if query_norm not in pname_norm and query_norm not in pname_no_ext:
                        continue
                    ram_mb = round((info.get("memory_info").rss or 0) / (1024 * 1024), 2) if info.get("memory_info") else 0.0
                    total_ram_mb += ram_mb
                    items.append({"pid": int(info.get("pid") or 0), "name": pname, "ram_mb": ram_mb})
                except Exception:
                    continue
            items.sort(key=lambda x: float(x.get("ram_mb", 0.0) or 0.0), reverse=True)
            return _json(
                {
                    "ok": True,
                    "mode": mode_norm,
                    "query": query_raw,
                    "process_count": len(items),
                    "total_ram_mb": round(total_ram_mb, 2),
                    "top_processes": items[:3],
                }
            )
        if mode_norm == "app_cpu_total":
            query_raw = (name or "").strip()
            if not query_raw:
                return self._error("name is required")
            query_norm = query_raw.casefold()
            if query_norm.endswith(".exe"):
                query_norm = query_norm[:-4]
            matched: list[dict[str, Any]] = []
            for p in psutil.process_iter(["pid", "name"]):
                try:
                    info = p.info
                    pname = str(info.get("name") or "").strip()
                    if not pname:
                        continue
                    pname_norm = pname.casefold()
                    pname_no_ext = pname_norm[:-4] if pname_norm.endswith(".exe") else pname_norm
                    if query_norm not in pname_norm and query_norm not in pname_no_ext:
                        continue
                    matched.append({"pid": int(info.get("pid") or 0), "name": pname, "proc": p})
                except Exception:
                    continue
            for item in matched:
                try:
                    item["proc"].cpu_percent(interval=None)
                except Exception:
                    item["cpu"] = 0.0
            time.sleep(0.6)
            total_cpu = 0.0
            items: list[dict[str, Any]] = []
            for item in matched:
                p = item.get("proc")
                try:
                    cpu_val = round(float(p.cpu_percent(interval=None)), 2)
                except Exception:
                    cpu_val = 0.0
                total_cpu += cpu_val
                items.append({"pid": item.get("pid"), "name": item.get("name"), "cpu": cpu_val})
            items.sort(key=lambda x: float(x.get("cpu", 0.0) or 0.0), reverse=True)
            return _json(
                {
                    "ok": True,
                    "mode": mode_norm,
                    "query": query_raw,
                    "process_count": len(items),
                    "total_cpu_percent": round(total_cpu, 2),
                    "items": items[:30],
                }
            )
        if mode_norm == "app_disk_total":
            query_raw = (name or "").strip()
            if not query_raw:
                return self._error("name is required")
            query_norm = query_raw.casefold()
            if query_norm.endswith(".exe"):
                query_norm = query_norm[:-4]
            items: list[dict[str, Any]] = []
            total_read_mb = 0.0
            total_write_mb = 0.0
            for p in psutil.process_iter(["pid", "name", "io_counters"]):
                try:
                    info = p.info
                    pname = str(info.get("name") or "").strip()
                    if not pname:
                        continue
                    pname_norm = pname.casefold()
                    pname_no_ext = pname_norm[:-4] if pname_norm.endswith(".exe") else pname_norm
                    if query_norm not in pname_norm and query_norm not in pname_no_ext:
                        continue
                    io = info.get("io_counters")
                    if not io:
                        continue
                    rb = float(getattr(io, "read_bytes", 0.0) or 0.0)
                    wb = float(getattr(io, "write_bytes", 0.0) or 0.0)
                    read_mb = round(rb / (1024 * 1024), 2)
                    write_mb = round(wb / (1024 * 1024), 2)
                    total_read_mb += read_mb
                    total_write_mb += write_mb
                    items.append(
                        {
                            "pid": info.get("pid"),
                            "name": pname,
                            "read_mb": read_mb,
                            "write_mb": write_mb,
                            "total_mb": round(read_mb + write_mb, 2),
                        }
                    )
                except Exception:
                    continue
            items.sort(key=lambda x: float(x.get("total_mb", 0.0) or 0.0), reverse=True)
            return _json(
                {
                    "ok": True,
                    "mode": mode_norm,
                    "query": query_raw,
                    "process_count": len(items),
                    "total_read_mb": round(total_read_mb, 2),
                    "total_write_mb": round(total_write_mb, 2),
                    "total_disk_mb": round(total_read_mb + total_write_mb, 2),
                    "items": items[:30],
                }
            )
        if mode_norm == "app_network_total":
            query_raw = (name or "").strip()
            if not query_raw:
                return self._error("name is required")
            query_norm = query_raw.casefold()
            if query_norm.endswith(".exe"):
                query_norm = query_norm[:-4]
            matched: list[dict[str, Any]] = []
            matched_pids: set[int] = set()
            for p in psutil.process_iter(["pid", "name"]):
                try:
                    info = p.info
                    pname = str(info.get("name") or "").strip()
                    if not pname:
                        continue
                    pname_norm = pname.casefold()
                    pname_no_ext = pname_norm[:-4] if pname_norm.endswith(".exe") else pname_norm
                    if query_norm not in pname_norm and query_norm not in pname_no_ext:
                        continue
                    pid_val = int(info.get("pid") or 0)
                    if pid_val <= 0:
                        continue
                    matched_pids.add(pid_val)
                    matched.append({"pid": pid_val, "name": pname})
                except Exception:
                    continue
            conn_count_by_pid: dict[int, int] = {}
            established_by_pid: dict[int, int] = {}
            remotes_by_pid: dict[int, set[str]] = {}
            try:
                for conn in psutil.net_connections(kind="inet"):
                    try:
                        pid_val = int(getattr(conn, "pid", 0) or 0)
                        if pid_val not in matched_pids:
                            continue
                        conn_count_by_pid[pid_val] = int(conn_count_by_pid.get(pid_val, 0)) + 1
                        status = str(getattr(conn, "status", "") or "").upper()
                        if status == "ESTABLISHED":
                            established_by_pid[pid_val] = int(established_by_pid.get(pid_val, 0)) + 1
                        raddr = getattr(conn, "raddr", None)
                        remote_ip = ""
                        if raddr and isinstance(raddr, tuple) and len(raddr) >= 1:
                            remote_ip = str(raddr[0] or "")
                        if remote_ip:
                            remotes_by_pid.setdefault(pid_val, set()).add(remote_ip)
                    except Exception:
                        continue
            except Exception as exc:
                return self._error(f"network connection enumeration failed: {exc}")
            items: list[dict[str, Any]] = []
            total_connections = 0
            total_established = 0
            all_remote_ips: set[str] = set()
            for proc in matched:
                pid_val = int(proc.get("pid") or 0)
                conn_count = int(conn_count_by_pid.get(pid_val, 0))
                est_count = int(established_by_pid.get(pid_val, 0))
                remotes = remotes_by_pid.get(pid_val, set())
                total_connections += conn_count
                total_established += est_count
                all_remote_ips.update(remotes)
                items.append(
                    {
                        "pid": pid_val,
                        "name": str(proc.get("name") or ""),
                        "connections": conn_count,
                        "established": est_count,
                        "remote_ips": len(remotes),
                    }
                )
            items.sort(key=lambda x: int(x.get("connections", 0)), reverse=True)
            return _json(
                {
                    "ok": True,
                    "mode": mode_norm,
                    "query": query_raw,
                    "process_count": len(matched),
                    "total_connections": total_connections,
                    "established_connections": total_established,
                    "unique_remote_ips": len(all_remote_ips),
                    "items": items[:30],
                }
            )
        if mode_norm == "app_resource_summary":
            query_raw = (name or "").strip()
            if not query_raw:
                return self._error("name is required")
            query_norm = query_raw.casefold()
            if query_norm.endswith(".exe"):
                query_norm = query_norm[:-4]
            matched: list[dict[str, Any]] = []
            matched_pids: set[int] = set()
            total_ram_mb = 0.0
            total_read_mb = 0.0
            total_write_mb = 0.0
            for p in psutil.process_iter(["pid", "name", "memory_info", "io_counters"]):
                try:
                    info = p.info
                    pname = str(info.get("name") or "").strip()
                    if not pname:
                        continue
                    pname_norm = pname.casefold()
                    pname_no_ext = pname_norm[:-4] if pname_norm.endswith(".exe") else pname_norm
                    if query_norm not in pname_norm and query_norm not in pname_no_ext:
                        continue
                    pid_val = int(info.get("pid") or 0)
                    if pid_val <= 0:
                        continue
                    matched_pids.add(pid_val)
                    ram_mb = round((info.get("memory_info").rss or 0) / (1024 * 1024), 2) if info.get("memory_info") else 0.0
                    total_ram_mb += ram_mb
                    io = info.get("io_counters")
                    rb = float(getattr(io, "read_bytes", 0.0) or 0.0) if io else 0.0
                    wb = float(getattr(io, "write_bytes", 0.0) or 0.0) if io else 0.0
                    read_mb = round(rb / (1024 * 1024), 2)
                    write_mb = round(wb / (1024 * 1024), 2)
                    total_read_mb += read_mb
                    total_write_mb += write_mb
                    matched.append({"pid": pid_val, "name": pname, "proc": p, "ram_mb": ram_mb, "read_mb": read_mb, "write_mb": write_mb})
                except Exception:
                    continue
            for item in matched:
                try:
                    item["proc"].cpu_percent(interval=None)
                except Exception:
                    item["cpu"] = 0.0
            time.sleep(0.6)
            total_cpu = 0.0
            items: list[dict[str, Any]] = []
            for item in matched:
                p = item.get("proc")
                try:
                    cpu_val = round(float(p.cpu_percent(interval=None)), 2)
                except Exception:
                    cpu_val = 0.0
                total_cpu += cpu_val
                items.append(
                    {
                        "pid": int(item.get("pid") or 0),
                        "name": str(item.get("name") or ""),
                        "cpu": cpu_val,
                        "ram_mb": float(item.get("ram_mb") or 0.0),
                        "read_mb": float(item.get("read_mb") or 0.0),
                        "write_mb": float(item.get("write_mb") or 0.0),
                    }
                )
            conn_count_by_pid: dict[int, int] = {}
            established_by_pid: dict[int, int] = {}
            all_remote_ips: set[str] = set()
            try:
                for conn in psutil.net_connections(kind="inet"):
                    try:
                        pid_val = int(getattr(conn, "pid", 0) or 0)
                        if pid_val not in matched_pids:
                            continue
                        conn_count_by_pid[pid_val] = int(conn_count_by_pid.get(pid_val, 0)) + 1
                        status = str(getattr(conn, "status", "") or "").upper()
                        if status == "ESTABLISHED":
                            established_by_pid[pid_val] = int(established_by_pid.get(pid_val, 0)) + 1
                        raddr = getattr(conn, "raddr", None)
                        remote_ip = ""
                        if raddr and isinstance(raddr, tuple) and len(raddr) >= 1:
                            remote_ip = str(raddr[0] or "")
                        if remote_ip:
                            all_remote_ips.add(remote_ip)
                    except Exception:
                        continue
            except Exception:
                pass
            total_connections = 0
            total_established = 0
            for row in items:
                pid_val = int(row.get("pid") or 0)
                c = int(conn_count_by_pid.get(pid_val, 0))
                e = int(established_by_pid.get(pid_val, 0))
                row["connections"] = c
                row["established"] = e
                total_connections += c
                total_established += e
            items.sort(key=lambda x: float(x.get("ram_mb", 0.0) or 0.0), reverse=True)
            return _json(
                {
                    "ok": True,
                    "mode": mode_norm,
                    "query": query_raw,
                    "process_count": len(items),
                    "total_ram_mb": round(total_ram_mb, 2),
                    "total_cpu_percent": round(total_cpu, 2),
                    "total_read_mb": round(total_read_mb, 2),
                    "total_write_mb": round(total_write_mb, 2),
                    "total_disk_mb": round(total_read_mb + total_write_mb, 2),
                    "total_connections": total_connections,
                    "established_connections": total_established,
                    "unique_remote_ips": len(all_remote_ips),
                    "items": items[:30],
                }
            )
        if mode_norm == "app_compare":
            left_name = (name or "").strip()
            right_name = (other_name or "").strip()
            if not left_name or not right_name:
                return self._error("name and other_name are required")
            try:
                left_raw = self._process_tools("app_resource_summary", name=left_name)
                right_raw = self._process_tools("app_resource_summary", name=right_name)
                left = json.loads(left_raw) if isinstance(left_raw, str) else {}
                right = json.loads(right_raw) if isinstance(right_raw, str) else {}
            except Exception as exc:
                return self._error(f"app_compare failed: {exc}")
            if not isinstance(left, dict) or not left.get("ok"):
                return self._error("failed to summarize first app")
            if not isinstance(right, dict) or not right.get("ok"):
                return self._error("failed to summarize second app")
            winners: dict[str, str] = {}
            metrics = (
                ("ram", "total_ram_mb"),
                ("cpu", "total_cpu_percent"),
                ("disk", "total_disk_mb"),
                ("network", "total_connections"),
            )
            left_query = str(left.get("query") or left_name)
            right_query = str(right.get("query") or right_name)
            for label, key in metrics:
                lv = float(left.get(key) or 0.0)
                rv = float(right.get(key) or 0.0)
                if lv > rv:
                    winners[label] = left_query
                elif rv > lv:
                    winners[label] = right_query
                else:
                    winners[label] = "equal"
            recommendations: list[str] = []
            if winners.get("ram") not in {"", "equal"}:
                recommendations.append(f"ram_hotspot={winners['ram']}")
            if winners.get("cpu") not in {"", "equal"}:
                recommendations.append(f"cpu_hotspot={winners['cpu']}")
            if winners.get("disk") not in {"", "equal"}:
                recommendations.append(f"disk_hotspot={winners['disk']}")
            if winners.get("network") not in {"", "equal"}:
                recommendations.append(f"network_hotspot={winners['network']}")
            if not recommendations:
                recommendations.append("balanced_usage")
            return _json(
                {
                    "ok": True,
                    "mode": mode_norm,
                    "left": left,
                    "right": right,
                    "winners": winners,
                    "recommendations": recommendations,
                }
            )
        if mode_norm == "app_reduce_ram_plan":
            query_raw = (name or "").strip()
            if not query_raw:
                return self._error("name is required")
            query_norm = query_raw.casefold()
            if query_norm.endswith(".exe"):
                query_norm = query_norm[:-4]
            items: list[dict[str, Any]] = []
            total_ram_mb = 0.0
            for p in psutil.process_iter(["pid", "name", "memory_info"]):
                try:
                    info = p.info
                    pname = str(info.get("name") or "").strip()
                    if not pname:
                        continue
                    pname_norm = pname.casefold()
                    pname_no_ext = pname_norm[:-4] if pname_norm.endswith(".exe") else pname_norm
                    if query_norm not in pname_norm and query_norm not in pname_no_ext:
                        continue
                    ram_mb = round((info.get("memory_info").rss or 0) / (1024 * 1024), 2) if info.get("memory_info") else 0.0
                    total_ram_mb += ram_mb
                    items.append({"pid": int(info.get("pid") or 0), "name": pname, "ram_mb": ram_mb})
                except Exception:
                    continue
            items.sort(key=lambda x: float(x.get("ram_mb", 0.0) or 0.0), reverse=True)
            top = items[:5]
            reclaimable_mb = round(sum(float(x.get("ram_mb", 0.0) or 0.0) for x in top[1:]), 2) if len(top) > 1 else 0.0
            plan = [
                "close_heaviest_secondary_processes",
                "close_extra_windows_or_tabs",
                "restart_app_if_needed",
            ]
            return _json(
                {
                    "ok": True,
                    "mode": mode_norm,
                    "query": query_raw,
                    "process_count": len(items),
                    "total_ram_mb": round(total_ram_mb, 2),
                    "reclaimable_mb_estimate": reclaimable_mb,
                    "top_processes": top,
                    "plan": plan,
                }
            )
        if mode_norm == "app_reduce_ram_execute":
            query_raw = (name or "").strip()
            if not query_raw:
                return self._error("name is required")
            kill_limit = 9999
            if max_kill is not None:
                try:
                    kill_limit = max(0, int(max_kill))
                except Exception:
                    kill_limit = 9999
            query_norm = query_raw.casefold()
            if query_norm.endswith(".exe"):
                query_norm = query_norm[:-4]
            items: list[dict[str, Any]] = []
            for p in psutil.process_iter(["pid", "name", "memory_info"]):
                try:
                    info = p.info
                    pname = str(info.get("name") or "").strip()
                    if not pname:
                        continue
                    pname_norm = pname.casefold()
                    pname_no_ext = pname_norm[:-4] if pname_norm.endswith(".exe") else pname_norm
                    if query_norm not in pname_norm and query_norm not in pname_no_ext:
                        continue
                    ram_mb = round((info.get("memory_info").rss or 0) / (1024 * 1024), 2) if info.get("memory_info") else 0.0
                    items.append({"pid": int(info.get("pid") or 0), "name": pname, "ram_mb": ram_mb, "proc": p})
                except Exception:
                    continue
            items.sort(key=lambda x: float(x.get("ram_mb", 0.0) or 0.0), reverse=True)
            if not items:
                return _json(
                    {
                        "ok": True,
                        "mode": mode_norm,
                        "query": query_raw,
                        "process_count": 0,
                        "killed_count": 0,
                        "killed": [],
                    }
                )
            protected = items[0]
            killed: list[dict[str, Any]] = []
            for item in items[1:]:
                if len(killed) >= kill_limit:
                    break
                p = item.get("proc")
                try:
                    if not dry_run:
                        p.kill()
                    killed.append(
                        {
                            "pid": int(item.get("pid") or 0),
                            "name": str(item.get("name") or ""),
                            "ram_mb": float(item.get("ram_mb") or 0.0),
                        }
                    )
                except Exception:
                    continue
            return _json(
                {
                    "ok": True,
                    "mode": mode_norm,
                    "query": query_raw,
                    "process_count": len(items),
                    "dry_run": bool(dry_run),
                    "max_kill": int(kill_limit),
                    "protected_pid": int(protected.get("pid") or 0),
                    "protected_name": str(protected.get("name") or ""),
                    "protected_ram_mb": float(protected.get("ram_mb") or 0.0),
                    "killed_count": len(killed),
                    "killed": killed[:30],
                }
            )
        if mode_norm == "app_reduce_cpu_plan":
            query_raw = (name or "").strip()
            if not query_raw:
                return self._error("name is required")
            query_norm = query_raw.casefold()
            if query_norm.endswith(".exe"):
                query_norm = query_norm[:-4]
            matched: list[dict[str, Any]] = []
            for p in psutil.process_iter(["pid", "name"]):
                try:
                    info = p.info
                    pname = str(info.get("name") or "").strip()
                    if not pname:
                        continue
                    pname_norm = pname.casefold()
                    pname_no_ext = pname_norm[:-4] if pname_norm.endswith(".exe") else pname_norm
                    if query_norm not in pname_norm and query_norm not in pname_no_ext:
                        continue
                    matched.append({"pid": int(info.get("pid") or 0), "name": pname, "proc": p})
                except Exception:
                    continue
            for item in matched:
                try:
                    item["proc"].cpu_percent(interval=None)
                except Exception:
                    item["cpu"] = 0.0
            time.sleep(0.6)
            items: list[dict[str, Any]] = []
            total_cpu = 0.0
            for item in matched:
                p = item.get("proc")
                try:
                    cpu_val = round(float(p.cpu_percent(interval=None)), 2)
                except Exception:
                    cpu_val = 0.0
                total_cpu += cpu_val
                items.append({"pid": int(item.get("pid") or 0), "name": str(item.get("name") or ""), "cpu": cpu_val})
            items.sort(key=lambda x: float(x.get("cpu", 0.0) or 0.0), reverse=True)
            reclaimable = round(sum(float(x.get("cpu", 0.0) or 0.0) for x in items[1:5]), 2) if len(items) > 1 else 0.0
            return _json(
                {
                    "ok": True,
                    "mode": mode_norm,
                    "query": query_raw,
                    "process_count": len(items),
                    "total_cpu_percent": round(total_cpu, 2),
                    "reclaimable_cpu_estimate": reclaimable,
                    "top_processes": items[:5],
                }
            )
        if mode_norm == "app_reduce_cpu_execute":
            query_raw = (name or "").strip()
            if not query_raw:
                return self._error("name is required")
            kill_threshold = float(threshold) if threshold is not None else 5.0
            kill_threshold = max(1.0, min(100.0, kill_threshold))
            kill_limit = 9999
            if max_kill is not None:
                try:
                    kill_limit = max(0, int(max_kill))
                except Exception:
                    kill_limit = 9999
            query_norm = query_raw.casefold()
            if query_norm.endswith(".exe"):
                query_norm = query_norm[:-4]
            matched: list[dict[str, Any]] = []
            for p in psutil.process_iter(["pid", "name"]):
                try:
                    info = p.info
                    pname = str(info.get("name") or "").strip()
                    if not pname:
                        continue
                    pname_norm = pname.casefold()
                    pname_no_ext = pname_norm[:-4] if pname_norm.endswith(".exe") else pname_norm
                    if query_norm not in pname_norm and query_norm not in pname_no_ext:
                        continue
                    matched.append({"pid": int(info.get("pid") or 0), "name": pname, "proc": p})
                except Exception:
                    continue
            for item in matched:
                try:
                    item["proc"].cpu_percent(interval=None)
                except Exception:
                    item["cpu"] = 0.0
            time.sleep(0.6)
            sampled: list[dict[str, Any]] = []
            for item in matched:
                p = item.get("proc")
                try:
                    cpu_val = round(float(p.cpu_percent(interval=None)), 2)
                except Exception:
                    cpu_val = 0.0
                sampled.append({"pid": int(item.get("pid") or 0), "name": str(item.get("name") or ""), "cpu": cpu_val, "proc": p})
            sampled.sort(key=lambda x: float(x.get("cpu", 0.0) or 0.0), reverse=True)
            protected = sampled[0] if sampled else None
            killed: list[dict[str, Any]] = []
            for item in sampled[1:]:
                if len(killed) >= kill_limit:
                    break
                cpu_val = float(item.get("cpu") or 0.0)
                if cpu_val < kill_threshold:
                    continue
                p = item.get("proc")
                try:
                    if not dry_run:
                        p.kill()
                    killed.append({"pid": int(item.get("pid") or 0), "name": str(item.get("name") or ""), "cpu": round(cpu_val, 2)})
                except Exception:
                    continue
            return _json(
                {
                    "ok": True,
                    "mode": mode_norm,
                    "query": query_raw,
                    "dry_run": bool(dry_run),
                    "max_kill": int(kill_limit),
                    "threshold": kill_threshold,
                    "process_count": len(sampled),
                    "protected_pid": int((protected or {}).get("pid") or 0),
                    "protected_name": str((protected or {}).get("name") or ""),
                    "protected_cpu": float((protected or {}).get("cpu") or 0.0),
                    "killed_count": len(killed),
                    "killed": killed[:30],
                }
            )
        if mode_norm == "app_reduce_disk_plan":
            query_raw = (name or "").strip()
            if not query_raw:
                return self._error("name is required")
            query_norm = query_raw.casefold()
            if query_norm.endswith(".exe"):
                query_norm = query_norm[:-4]
            items: list[dict[str, Any]] = []
            total_disk_mb = 0.0
            for p in psutil.process_iter(["pid", "name", "io_counters"]):
                try:
                    info = p.info
                    pname = str(info.get("name") or "").strip()
                    if not pname:
                        continue
                    pname_norm = pname.casefold()
                    pname_no_ext = pname_norm[:-4] if pname_norm.endswith(".exe") else pname_norm
                    if query_norm not in pname_norm and query_norm not in pname_no_ext:
                        continue
                    io = info.get("io_counters")
                    if not io:
                        continue
                    rb = float(getattr(io, "read_bytes", 0.0) or 0.0)
                    wb = float(getattr(io, "write_bytes", 0.0) or 0.0)
                    total_mb = round((rb + wb) / (1024 * 1024), 2)
                    total_disk_mb += total_mb
                    items.append({"pid": int(info.get("pid") or 0), "name": pname, "disk_mb": total_mb})
                except Exception:
                    continue
            items.sort(key=lambda x: float(x.get("disk_mb", 0.0) or 0.0), reverse=True)
            reclaimable = round(sum(float(x.get("disk_mb", 0.0) or 0.0) for x in items[1:5]), 2) if len(items) > 1 else 0.0
            return _json(
                {
                    "ok": True,
                    "mode": mode_norm,
                    "query": query_raw,
                    "process_count": len(items),
                    "total_disk_mb": round(total_disk_mb, 2),
                    "reclaimable_disk_estimate": reclaimable,
                    "top_processes": items[:5],
                }
            )
        if mode_norm == "app_reduce_disk_execute":
            query_raw = (name or "").strip()
            if not query_raw:
                return self._error("name is required")
            kill_threshold = float(threshold) if threshold is not None else 50.0
            kill_threshold = max(1.0, kill_threshold)
            kill_limit = 9999
            if max_kill is not None:
                try:
                    kill_limit = max(0, int(max_kill))
                except Exception:
                    kill_limit = 9999
            query_norm = query_raw.casefold()
            if query_norm.endswith(".exe"):
                query_norm = query_norm[:-4]
            items: list[dict[str, Any]] = []
            for p in psutil.process_iter(["pid", "name", "io_counters"]):
                try:
                    info = p.info
                    pname = str(info.get("name") or "").strip()
                    if not pname:
                        continue
                    pname_norm = pname.casefold()
                    pname_no_ext = pname_norm[:-4] if pname_norm.endswith(".exe") else pname_norm
                    if query_norm not in pname_norm and query_norm not in pname_no_ext:
                        continue
                    io = info.get("io_counters")
                    if not io:
                        continue
                    rb = float(getattr(io, "read_bytes", 0.0) or 0.0)
                    wb = float(getattr(io, "write_bytes", 0.0) or 0.0)
                    total_mb = round((rb + wb) / (1024 * 1024), 2)
                    items.append({"pid": int(info.get("pid") or 0), "name": pname, "disk_mb": total_mb, "proc": p})
                except Exception:
                    continue
            items.sort(key=lambda x: float(x.get("disk_mb", 0.0) or 0.0), reverse=True)
            protected = items[0] if items else None
            killed: list[dict[str, Any]] = []
            for item in items[1:]:
                if len(killed) >= kill_limit:
                    break
                if float(item.get("disk_mb") or 0.0) < kill_threshold:
                    continue
                try:
                    if not dry_run:
                        item["proc"].kill()
                    killed.append({"pid": int(item.get("pid") or 0), "name": str(item.get("name") or ""), "disk_mb": float(item.get("disk_mb") or 0.0)})
                except Exception:
                    continue
            return _json(
                {
                    "ok": True,
                    "mode": mode_norm,
                    "query": query_raw,
                    "dry_run": bool(dry_run),
                    "max_kill": int(kill_limit),
                    "threshold": kill_threshold,
                    "process_count": len(items),
                    "protected_pid": int((protected or {}).get("pid") or 0),
                    "protected_name": str((protected or {}).get("name") or ""),
                    "protected_disk_mb": float((protected or {}).get("disk_mb") or 0.0),
                    "killed_count": len(killed),
                    "killed": killed[:30],
                }
            )
        if mode_norm == "app_reduce_network_plan":
            query_raw = (name or "").strip()
            if not query_raw:
                return self._error("name is required")
            query_norm = query_raw.casefold()
            if query_norm.endswith(".exe"):
                query_norm = query_norm[:-4]
            matched: list[dict[str, Any]] = []
            pids: set[int] = set()
            for p in psutil.process_iter(["pid", "name"]):
                try:
                    info = p.info
                    pname = str(info.get("name") or "").strip()
                    if not pname:
                        continue
                    pname_norm = pname.casefold()
                    pname_no_ext = pname_norm[:-4] if pname_norm.endswith(".exe") else pname_norm
                    if query_norm not in pname_norm and query_norm not in pname_no_ext:
                        continue
                    pid_val = int(info.get("pid") or 0)
                    if pid_val <= 0:
                        continue
                    pids.add(pid_val)
                    matched.append({"pid": pid_val, "name": pname})
                except Exception:
                    continue
            counts: dict[int, int] = {}
            for conn in psutil.net_connections(kind="inet"):
                try:
                    pid_val = int(getattr(conn, "pid", 0) or 0)
                    if pid_val in pids:
                        counts[pid_val] = int(counts.get(pid_val, 0)) + 1
                except Exception:
                    continue
            items = [{"pid": m["pid"], "name": m["name"], "connections": int(counts.get(m["pid"], 0))} for m in matched]
            items.sort(key=lambda x: int(x.get("connections", 0)), reverse=True)
            total_connections = sum(int(x.get("connections", 0)) for x in items)
            reclaimable = sum(int(x.get("connections", 0)) for x in items[1:5]) if len(items) > 1 else 0
            return _json(
                {
                    "ok": True,
                    "mode": mode_norm,
                    "query": query_raw,
                    "process_count": len(items),
                    "total_connections": total_connections,
                    "reclaimable_network_estimate": reclaimable,
                    "top_processes": items[:5],
                }
            )
        if mode_norm == "app_reduce_network_execute":
            query_raw = (name or "").strip()
            if not query_raw:
                return self._error("name is required")
            kill_threshold = float(threshold) if threshold is not None else 3.0
            kill_threshold = max(1.0, kill_threshold)
            kill_limit = 9999
            if max_kill is not None:
                try:
                    kill_limit = max(0, int(max_kill))
                except Exception:
                    kill_limit = 9999
            query_norm = query_raw.casefold()
            if query_norm.endswith(".exe"):
                query_norm = query_norm[:-4]
            matched: list[dict[str, Any]] = []
            pids: set[int] = set()
            procs: dict[int, Any] = {}
            for p in psutil.process_iter(["pid", "name"]):
                try:
                    info = p.info
                    pname = str(info.get("name") or "").strip()
                    if not pname:
                        continue
                    pname_norm = pname.casefold()
                    pname_no_ext = pname_norm[:-4] if pname_norm.endswith(".exe") else pname_norm
                    if query_norm not in pname_norm and query_norm not in pname_no_ext:
                        continue
                    pid_val = int(info.get("pid") or 0)
                    if pid_val <= 0:
                        continue
                    pids.add(pid_val)
                    procs[pid_val] = p
                    matched.append({"pid": pid_val, "name": pname})
                except Exception:
                    continue
            counts: dict[int, int] = {}
            for conn in psutil.net_connections(kind="inet"):
                try:
                    pid_val = int(getattr(conn, "pid", 0) or 0)
                    if pid_val in pids:
                        counts[pid_val] = int(counts.get(pid_val, 0)) + 1
                except Exception:
                    continue
            items = [{"pid": m["pid"], "name": m["name"], "connections": int(counts.get(m["pid"], 0))} for m in matched]
            items.sort(key=lambda x: int(x.get("connections", 0)), reverse=True)
            protected = items[0] if items else None
            killed: list[dict[str, Any]] = []
            for item in items[1:]:
                if len(killed) >= kill_limit:
                    break
                if float(item.get("connections") or 0.0) < kill_threshold:
                    continue
                try:
                    if not dry_run:
                        procs[int(item["pid"])].kill()
                    killed.append({"pid": int(item.get("pid") or 0), "name": str(item.get("name") or ""), "connections": int(item.get("connections") or 0)})
                except Exception:
                    continue
            return _json(
                {
                    "ok": True,
                    "mode": mode_norm,
                    "query": query_raw,
                    "dry_run": bool(dry_run),
                    "max_kill": int(kill_limit),
                    "threshold": kill_threshold,
                    "process_count": len(items),
                    "protected_pid": int((protected or {}).get("pid") or 0),
                    "protected_name": str((protected or {}).get("name") or ""),
                    "protected_connections": int((protected or {}).get("connections") or 0),
                    "killed_count": len(killed),
                    "killed": killed[:30],
                }
            )

        if mode_norm == "kill_pid":
            try:
                target_pid = int(pid)
            except Exception:
                return self._error("valid pid is required")
            try:
                p = psutil.Process(target_pid)
                p.kill()
                return _json({"ok": True, "mode": mode_norm, "pid": target_pid})
            except Exception as exc:
                return self._error(f"kill_pid failed: {exc}")

        if mode_norm == "kill_name":
            query = (name or "").strip().casefold()
            if not query:
                return self._error("name is required")
            killed = 0
            for p in psutil.process_iter(["name"]):
                try:
                    pname = str(p.info.get("name") or "").casefold()
                    if query in pname:
                        p.kill()
                        killed += 1
                except Exception:
                    continue
            return _json({"ok": True, "mode": mode_norm, "name": name, "killed": killed})
        if mode_norm == "close_browsers":
            browser_names = {"chrome.exe", "msedge.exe", "firefox.exe", "opera.exe", "brave.exe"}
            killed = 0
            for p in psutil.process_iter(["name"]):
                try:
                    pname = str(p.info.get("name") or "").lower()
                    if pname in browser_names:
                        p.kill()
                        killed += 1
                except Exception:
                    continue
            return _json({"ok": True, "mode": mode_norm, "killed": killed})
        if mode_norm == "close_office":
            office_names = {
                "winword.exe",
                "excel.exe",
                "powerpnt.exe",
                "outlook.exe",
                "onenote.exe",
                "msaccess.exe",
            }
            killed = 0
            for p in psutil.process_iter(["name"]):
                try:
                    pname = str(p.info.get("name") or "").lower()
                    if pname in office_names:
                        p.kill()
                        killed += 1
                except Exception:
                    continue
            return _json({"ok": True, "mode": mode_norm, "killed": killed})

        if mode_norm == "path_by_pid":
            try:
                target_pid = int(pid)
            except Exception:
                return self._error("valid pid is required")
            try:
                p = psutil.Process(target_pid)
                return _json({"ok": True, "mode": mode_norm, "pid": target_pid, "path": p.exe()})
            except Exception as exc:
                return self._error(f"path_by_pid failed: {exc}")
        if mode_norm in {"cpu_by_pid", "ram_by_pid", "threads_by_pid", "start_time_by_pid"}:
            try:
                target_pid = int(pid)
            except Exception:
                return self._error("valid pid is required")
            try:
                proc = psutil.Process(target_pid)
                if mode_norm == "cpu_by_pid":
                    cpu_val = float(proc.cpu_percent(interval=0.2))
                    return _json({"ok": True, "mode": mode_norm, "pid": target_pid, "cpu": round(cpu_val, 2)})
                if mode_norm == "ram_by_pid":
                    mem = proc.memory_info()
                    return _json(
                        {
                            "ok": True,
                            "mode": mode_norm,
                            "pid": target_pid,
                            "ram_mb": round(float(mem.rss) / (1024 * 1024), 2),
                        }
                    )
                if mode_norm == "threads_by_pid":
                    return _json({"ok": True, "mode": mode_norm, "pid": target_pid, "threads": int(proc.num_threads())})
                started = datetime.fromtimestamp(proc.create_time()).isoformat()
                return _json({"ok": True, "mode": mode_norm, "pid": target_pid, "start_time": started})
            except Exception as exc:
                return self._error(f"{mode_norm} failed: {exc}")
        if mode_norm == "app_uptime":
            items = []
            now_ts = time.time()
            for p in psutil.process_iter(["pid", "name", "create_time"]):
                try:
                    started = float(p.info.get("create_time") or 0.0)
                    if started <= 0:
                        continue
                    items.append(
                        {
                            "pid": int(p.info.get("pid") or 0),
                            "name": str(p.info.get("name") or ""),
                            "uptime_seconds": max(0, int(now_ts - started)),
                        }
                    )
                except Exception:
                    continue
            items.sort(key=lambda x: int(x.get("uptime_seconds", 0)), reverse=True)
            return _json({"ok": True, "mode": mode_norm, "items": items[:50]})
        if mode_norm == "set_priority":
            try:
                target_pid = int(pid)
            except Exception:
                return self._error("valid pid is required")
            priority_norm = (priority or name or "").strip().lower()
            if not priority_norm:
                return self._error("priority is required (high|normal|low)")
            priority_map: dict[str, int] = {
                "high": psutil.HIGH_PRIORITY_CLASS,
                "normal": psutil.NORMAL_PRIORITY_CLASS,
                "low": psutil.BELOW_NORMAL_PRIORITY_CLASS,
            }
            if priority_norm not in priority_map:
                return self._error("priority must be one of: high, normal, low")
            try:
                proc = psutil.Process(target_pid)
                proc.nice(priority_map[priority_norm])
                return _json(
                    {
                        "ok": True,
                        "mode": mode_norm,
                        "pid": target_pid,
                        "priority": priority_norm,
                    }
                )
            except Exception as exc:
                return self._error(f"set_priority failed: {exc}")
        if mode_norm == "suspend_pid":
            try:
                target_pid = int(pid)
            except Exception:
                return self._error("valid pid is required")
            try:
                psutil.Process(target_pid).suspend()
                return _json({"ok": True, "mode": mode_norm, "pid": target_pid})
            except Exception as exc:
                return self._error(f"suspend_pid failed: {exc}")
        if mode_norm == "resume_pid":
            try:
                target_pid = int(pid)
            except Exception:
                return self._error("valid pid is required")
            try:
                psutil.Process(target_pid).resume()
                return _json({"ok": True, "mode": mode_norm, "pid": target_pid})
            except Exception as exc:
                return self._error(f"resume_pid failed: {exc}")
        if mode_norm == "unresponsive":
            ps = (
                "Get-Process | Where-Object {$_.Responding -eq $false} | "
                "Select-Object Id,ProcessName,MainWindowTitle | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=12)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "unresponsive process query failed")
        if mode_norm == "kill_unresponsive":
            ps = (
                "$items=Get-Process | Where-Object {$_.Responding -eq $false}; "
                "$killed=@(); "
                "foreach($p in $items){ try { Stop-Process -Id $p.Id -Force -ErrorAction Stop; $killed += [pscustomobject]@{pid=$p.Id;name=$p.ProcessName} } catch {} }; "
                "@{ok=$true; mode='kill_unresponsive'; killed=$killed} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=18)
            if ok and out:
                return out
            return self._error(out or "kill unresponsive failed")
        if mode_norm == "kill_high_cpu":
            kill_threshold = float(threshold) if threshold is not None else 50.0
            kill_threshold = max(5.0, min(100.0, kill_threshold))
            kill_limit = 9999
            if max_kill is not None:
                try:
                    kill_limit = max(0, int(max_kill))
                except Exception:
                    kill_limit = 9999
            warm = {}
            for p in psutil.process_iter(["pid", "name"]):
                try:
                    warm[int(p.info.get("pid") or 0)] = p.cpu_percent(interval=None)
                except Exception:
                    continue
            time.sleep(0.7)
            killed: list[dict[str, Any]] = []
            for p in psutil.process_iter(["pid", "name"]):
                try:
                    if len(killed) >= kill_limit:
                        break
                    cpu_val = float(p.cpu_percent(interval=None))
                    if cpu_val >= kill_threshold:
                        killed.append(
                            {
                                "pid": int(p.info.get("pid") or 0),
                                "name": str(p.info.get("name") or ""),
                                "cpu": round(cpu_val, 2),
                            }
                        )
                        if not dry_run:
                            p.kill()
                except Exception:
                    continue
            return _json(
                {
                    "ok": True,
                    "mode": mode_norm,
                    "dry_run": bool(dry_run),
                    "max_kill": int(kill_limit),
                    "threshold": kill_threshold,
                    "killed": killed,
                    "count": len(killed),
                }
            )
        if mode_norm == "restart_explorer":
            ps = (
                "Get-Process explorer -ErrorAction SilentlyContinue | Stop-Process -Force -ErrorAction SilentlyContinue; "
                "Start-Process explorer.exe; "
                "@{ok=$true; mode='restart_explorer'} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=15)
            if ok and out:
                return out
            return self._error(out or "restart explorer failed")

        return self._error(f"unsupported process_tools mode: {mode_norm}")

    def _service_tools(self, mode: str, name: str = "", startup: str = "") -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm == "list":
            ps = (
                "Get-Service | Select-Object -First 150 Name,DisplayName,Status,StartType | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=15)
            if ok and out:
                return _json({"ok": True, "mode": "list", "data": json.loads(out)})
            return self._error(out or "list services failed")
        if mode_norm in {"describe", "description"}:
            svc = (name or "").strip()
            if not svc:
                return self._error("service name is required")
            ps = (
                f"Get-CimInstance Win32_Service -Filter \"Name='{svc}'\" | "
                "Select-Object Name,DisplayName,State,StartMode,StartName,Description | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=15)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "name": svc, "data": json.loads(out)})
            return self._error(out or "service description failed")
        if mode_norm in {"dependencies", "deps"}:
            svc = (name or "").strip()
            if not svc:
                return self._error("service name is required")
            ps = (
                f"$s=Get-Service -Name '{svc}' -ErrorAction Stop; "
                "$deps=@($s.ServicesDependedOn | Select-Object Name,DisplayName,Status); "
                "$dependents=@($s.DependentServices | Select-Object Name,DisplayName,Status); "
                "@{name=$s.Name; display_name=$s.DisplayName; dependencies=$deps; dependents=$dependents} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=15)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "name": svc, "data": json.loads(out)})
            return self._error(out or "service dependencies failed")
        if mode_norm in {"user_services", "list_user_services"}:
            ps = (
                "Get-CimInstance Win32_Service | "
                "Where-Object {$_.StartName -and $_.StartName -notlike 'LocalSystem' -and $_.StartName -notlike 'NT AUTHORITY*'} | "
                "Select-Object Name,DisplayName,State,StartMode,StartName | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=20)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "user services query failed")
        svc = (name or "").strip()
        if not svc:
            return self._error("service name is required")
        if mode_norm == "start":
            ok, out = _run_powershell(f"Start-Service -Name '{svc}'", timeout=15)
            return _json({"ok": True, "mode": mode_norm, "name": svc}) if ok else self._error(out or "start service failed")
        if mode_norm == "stop":
            ok, out = _run_powershell(f"Stop-Service -Name '{svc}' -Force", timeout=15)
            return _json({"ok": True, "mode": mode_norm, "name": svc}) if ok else self._error(out or "stop service failed")
        if mode_norm == "restart":
            ok, out = _run_powershell(f"Restart-Service -Name '{svc}' -Force", timeout=15)
            return _json({"ok": True, "mode": mode_norm, "name": svc}) if ok else self._error(out or "restart service failed")
        if mode_norm == "startup":
            start_mode = (startup or "").strip().lower()
            mapped = {"auto": "Automatic", "automatic": "Automatic", "manual": "Manual", "disabled": "Disabled"}.get(start_mode)
            if not mapped:
                return self._error("startup must be auto|manual|disabled")
            ok, out = _run_powershell(f"Set-Service -Name '{svc}' -StartupType {mapped}", timeout=15)
            return _json({"ok": True, "mode": mode_norm, "name": svc, "startup": mapped}) if ok else self._error(out or "set startup failed")
        return self._error(f"unsupported service_tools mode: {mode_norm}")

    def _background_tools(self, mode: str, max_results: int = 50) -> str:
        mode_norm = (mode or "").strip().lower()
        max_results = _clamp(max_results, 1, 200)
        try:
            import psutil
        except Exception as exc:
            return self._error(f"psutil unavailable: {exc}")

        windows = _serialize_windows(include_untitled=False, limit=600)
        window_pids = {int(w.get("pid") or 0) for w in windows if int(w.get("pid") or 0) > 0}

        if mode_norm in {"count_background", "background_count"}:
            total = 0
            background = 0
            for p in psutil.process_iter(["pid"]):
                try:
                    pid = int(p.info.get("pid") or 0)
                    if pid <= 0:
                        continue
                    total += 1
                    if pid not in window_pids:
                        background += 1
                except Exception:
                    continue
            return _json({"ok": True, "mode": mode_norm, "total": total, "background": background})

        if mode_norm in {"list_visible_windows", "visible_windows"}:
            return _json({"ok": True, "mode": mode_norm, "count": len(windows), "items": windows[:max_results]})

        if mode_norm in {"list_minimized_windows", "minimized_windows"}:
            try:
                import pygetwindow as gw

                items = []
                for win in gw.getAllWindows():
                    try:
                        if not bool(getattr(win, "isMinimized", False)):
                            continue
                        title = str(getattr(win, "title", "") or "").strip()
                        if not title:
                            continue
                        items.append(
                            {
                                "title": title,
                                "left": int(getattr(win, "left", 0) or 0),
                                "top": int(getattr(win, "top", 0) or 0),
                                "width": int(getattr(win, "width", 0) or 0),
                                "height": int(getattr(win, "height", 0) or 0),
                            }
                        )
                    except Exception:
                        continue
                return _json({"ok": True, "mode": mode_norm, "count": len(items), "items": items[:max_results]})
            except Exception as exc:
                return self._error(f"minimized windows failed: {exc}")

        if mode_norm in {"ghost_apps", "headless_heavy"}:
            items = []
            for p in psutil.process_iter(["pid", "name", "cpu_percent", "memory_info"]):
                try:
                    pid = int(p.info.get("pid") or 0)
                    if pid <= 0 or pid in window_pids:
                        continue
                    cpu = float(p.info.get("cpu_percent") or 0.0)
                    ram_mb = round(float((p.info.get("memory_info").rss or 0.0)) / (1024 * 1024), 2) if p.info.get("memory_info") else 0.0
                    if cpu >= 1.0 or ram_mb >= 120.0:
                        items.append({"pid": pid, "name": str(p.info.get("name") or ""), "cpu": cpu, "ram_mb": ram_mb})
                except Exception:
                    continue
            items.sort(key=lambda x: (float(x.get("cpu", 0.0)), float(x.get("ram_mb", 0.0))), reverse=True)
            return _json({"ok": True, "mode": mode_norm, "count": len(items), "items": items[:max_results]})

        if mode_norm in {"activity_time", "uptime_per_app"}:
            now_ts = time.time()
            items = []
            for p in psutil.process_iter(["pid", "name", "create_time"]):
                try:
                    created = float(p.info.get("create_time") or 0.0)
                    if created <= 0:
                        continue
                    items.append(
                        {
                            "pid": int(p.info.get("pid") or 0),
                            "name": str(p.info.get("name") or ""),
                            "uptime_seconds": max(0, int(now_ts - created)),
                        }
                    )
                except Exception:
                    continue
            items.sort(key=lambda x: int(x.get("uptime_seconds", 0)), reverse=True)
            return _json({"ok": True, "mode": mode_norm, "items": items[:max_results]})

        if mode_norm in {"network_usage_per_app", "net_usage_by_app"}:
            ps = (
                "$conns=Get-NetTCPConnection -State Established -ErrorAction SilentlyContinue | Group-Object OwningProcess; "
                "$rows=@(); "
                "foreach($g in $conns){ "
                "  $pid=[int]$g.Name; "
                "  $p=Get-Process -Id $pid -ErrorAction SilentlyContinue; "
                "  if($p){ "
                "    $rows += [pscustomobject]@{pid=$pid;name=$p.ProcessName;connections=$g.Count;path=$p.Path} "
                "  } "
                "}; "
                "$rows | Sort-Object connections -Descending | Select-Object -First 80 | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=20)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "items": json.loads(out)})
            return self._error(out or "network usage per app failed")

        if mode_norm in {
            "camera_usage",
            "mic_usage",
            "camera_usage_now",
            "mic_usage_now",
            "camera_active_now",
            "mic_active_now",
        }:
            cap = "webcam" if "camera" in mode_norm else "microphone"
            active_only = mode_norm.endswith("_now") or "active_now" in mode_norm
            ps = (
                f"$root='HKCU:\\Software\\Microsoft\\Windows\\CurrentVersion\\CapabilityAccessManager\\ConsentStore\\{cap}\\NonPackaged'; "
                "$rows=@(); "
                "if(Test-Path $root){ "
                "  Get-ChildItem $root -ErrorAction SilentlyContinue | ForEach-Object { "
                "    $p=Get-ItemProperty -Path $_.PSPath -ErrorAction SilentlyContinue; "
                "    $rows += [pscustomobject]@{app=$_.PSChildName;last_start=$p.LastUsedTimeStart;last_stop=$p.LastUsedTimeStop} "
                "  } "
                "}; "
                "$rows | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=15)
            if ok and out:
                try:
                    rows = json.loads(out)
                    if isinstance(rows, dict):
                        rows = [rows]
                    normalized: list[dict[str, Any]] = []
                    for row in rows if isinstance(rows, list) else []:
                        if not isinstance(row, dict):
                            continue
                        start_raw = row.get("last_start")
                        stop_raw = row.get("last_stop")
                        try:
                            start_val = int(str(start_raw or "").strip())
                        except Exception:
                            start_val = 0
                        try:
                            stop_val = int(str(stop_raw or "").strip())
                        except Exception:
                            stop_val = 0
                        active_now = start_val > 0 and (stop_val <= 0 or stop_val < start_val)
                        item = {
                            "app": str(row.get("app") or ""),
                            "last_start": start_raw,
                            "last_stop": stop_raw,
                            "last_start_iso": _filetime_to_iso(start_raw),
                            "last_stop_iso": _filetime_to_iso(stop_raw),
                            "active_now": active_now,
                        }
                        if active_only and not active_now:
                            continue
                        normalized.append(item)
                    return _json(
                        {
                            "ok": True,
                            "mode": mode_norm,
                            "capability": cap,
                            "active_only": active_only,
                            "active_count": sum(1 for it in normalized if bool(it.get("active_now"))),
                            "count": len(normalized),
                            "items": normalized[:max_results],
                        }
                    )
                except Exception as exc:
                    return self._error(f"{cap} usage parse failed: {exc}")
            return self._error(out or f"{cap} usage lookup failed")

        if mode_norm in {"wake_lock_apps", "sleep_blockers"}:
            ok, out = _run_powershell("powercfg /requests", timeout=20)
            if ok:
                return _json({"ok": True, "mode": mode_norm, "output": out[:3000]})
            return self._error(out or "power requests failed")

        if mode_norm in {"process_paths", "running_app_paths"}:
            items = []
            for p in psutil.process_iter(["pid", "name", "exe"]):
                try:
                    path = str(p.info.get("exe") or "")
                    if not path:
                        continue
                    items.append({"pid": int(p.info.get("pid") or 0), "name": str(p.info.get("name") or ""), "path": path})
                except Exception:
                    continue
            return _json({"ok": True, "mode": mode_norm, "count": len(items), "items": items[:max_results]})

        return self._error(f"unsupported background_tools mode: {mode_norm}")

    def _startup_tools(
        self,
        mode: str,
        name: str = "",
        seconds: Any = None,
        monitor_seconds: Any = None,
        notify: bool = False,
    ) -> str:
        mode_norm = (mode or "").strip().lower()
        name_query = (name or "").strip()

        def _startup_snapshot() -> tuple[bool, list[dict[str, Any]], str]:
            ps = (
                "Get-CimInstance Win32_StartupCommand | "
                "Select-Object Name,Command,Location,User | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=20)
            if not ok or not out:
                return False, [], out or "startup snapshot failed"
            try:
                raw = json.loads(out)
                if isinstance(raw, dict):
                    raw = [raw]
                if not isinstance(raw, list):
                    return False, [], "startup snapshot parse failed"
                items = [it for it in raw if isinstance(it, dict)]
                return True, items, ""
            except Exception as exc:
                return False, [], f"startup snapshot parse failed: {exc}"

        if mode_norm in {"list", "startup_list"}:
            ok, items, err = _startup_snapshot()
            if ok:
                return _json({"ok": True, "mode": mode_norm, "items": items})
            return self._error(err)

        if mode_norm in {"impact_time", "startup_impact_time", "impact_breakdown"}:
            ps = (
                "$e=Get-WinEvent -FilterHashtable @{LogName='Microsoft-Windows-Diagnostics-Performance/Operational';ID=100} -MaxEvents 1 -ErrorAction SilentlyContinue; "
                "if(-not $e){ @{ok=$false;error='no startup diagnostics event found'} | ConvertTo-Json -Compress; exit 0 }; "
                "$xml=[xml]$e.ToXml(); "
                "$obj=@{}; foreach($d in $xml.Event.EventData.Data){ $obj[$d.Name]=$d.'#text' }; "
                "$boot=[double]($obj['BootTime']); "
                "$main=[double]($obj['MainPathBootTime']); "
                "$post=[double]($obj['BootPostBootTime']); "
                "@{ok=$true; event_time=$e.TimeCreated; data=$obj; boot_ms=$boot; main_path_ms=$main; post_boot_ms=$post; boot_seconds=([math]::Round($boot/1000,2)); main_path_seconds=([math]::Round($main/1000,2)); post_boot_seconds=([math]::Round($post/1000,2))} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=20)
            if ok and out:
                return out
            return self._error(out or "startup impact query failed")

        if mode_norm in {"registry_startups", "registry_list"}:
            ps = (
                "$paths=@("
                "'HKCU:\\Software\\Microsoft\\Windows\\CurrentVersion\\Run',"
                "'HKLM:\\Software\\Microsoft\\Windows\\CurrentVersion\\Run',"
                "'HKLM:\\Software\\WOW6432Node\\Microsoft\\Windows\\CurrentVersion\\Run'"
                "); "
                "$rows=@(); "
                "foreach($p in $paths){ "
                " if(Test-Path $p){ "
                "  $props=(Get-ItemProperty -Path $p); "
                "  foreach($n in $props.PSObject.Properties.Name){ "
                "   if($n -in 'PSPath','PSParentPath','PSChildName','PSDrive','PSProvider'){continue}; "
                "   $rows += [pscustomobject]@{path=$p;name=$n;command=$props.$n} "
                "  } "
                " } "
                "}; "
                "$rows | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=20)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "items": json.loads(out)})
            return self._error(out or "registry startup list failed")

        if mode_norm in {"folder_startups", "folder_list"}:
            ps = (
                "$paths=@($env:APPDATA+'\\Microsoft\\Windows\\Start Menu\\Programs\\Startup',"
                "$env:ProgramData+'\\Microsoft\\Windows\\Start Menu\\Programs\\Startup'); "
                "$rows=@(); "
                "foreach($p in $paths){ if(Test-Path $p){ Get-ChildItem -Path $p -Force -ErrorAction SilentlyContinue | ForEach-Object { $rows += [pscustomobject]@{folder=$p;name=$_.Name;full_path=$_.FullName} } } }; "
                "$rows | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=20)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "items": json.loads(out)})
            return self._error(out or "startup folder list failed")

        if mode_norm in {"disable", "disable_startup", "enable", "enable_startup"}:
            if not name_query:
                return self._error("name is required")
            escaped = name_query.replace("'", "''")
            enable_mode = mode_norm.startswith("enable")
            ps = (
                f"$q='{escaped}'; "
                "$paths=@("
                "'HKCU:\\Software\\Microsoft\\Windows\\CurrentVersion\\Run',"
                "'HKLM:\\Software\\Microsoft\\Windows\\CurrentVersion\\Run'"
                "); "
                "$changed=0; $rows=@(); "
                "foreach($p in $paths){ "
                " if(!(Test-Path $p)){continue}; "
                " $disabled=($p -replace '\\\\Run$','\\RunDisabled'); "
                " if(!(Test-Path $disabled)){ New-Item -Path $disabled -Force | Out-Null }; "
                " $src=if(" + ("$true" if enable_mode else "$false") + "){$disabled}else{$p}; "
                " $dst=if(" + ("$true" if enable_mode else "$false") + "){$p}else{$disabled}; "
                " if(!(Test-Path $src)){continue}; "
                " $props=(Get-ItemProperty -Path $src); "
                " foreach($n in $props.PSObject.Properties.Name){ "
                "  if($n -in 'PSPath','PSParentPath','PSChildName','PSDrive','PSProvider'){continue}; "
                "  $v=[string]$props.$n; "
                "  if($n -like ('*'+$q+'*') -or $v -like ('*'+$q+'*')){ "
                "    New-ItemProperty -Path $dst -Name $n -Value $v -PropertyType String -Force | Out-Null; "
                "    Remove-ItemProperty -Path $src -Name $n -ErrorAction SilentlyContinue; "
                "    $changed++; $rows += [pscustomobject]@{name=$n;from=$src;to=$dst} "
                "  } "
                " } "
                "}; "
                "$folders=@($env:APPDATA+'\\Microsoft\\Windows\\Start Menu\\Programs\\Startup',$env:ProgramData+'\\Microsoft\\Windows\\Start Menu\\Programs\\Startup'); "
                "foreach($f in $folders){ "
                " if(!(Test-Path $f)){continue}; "
                " Get-ChildItem -Path $f -Force -ErrorAction SilentlyContinue | ForEach-Object { "
                "  $nm=$_.Name; "
                "  if($nm -like ('*'+$q+'*')){ "
                "    if(" + ("$true" if enable_mode else "$false") + " -and $nm.EndsWith('.disabled')){ "
                "      $new=$_.FullName.Substring(0,$_.FullName.Length-9); Rename-Item -Path $_.FullName -NewName ([System.IO.Path]::GetFileName($new)) -Force; $changed++ "
                "    } "
                "    if(" + ("$false" if enable_mode else "$true") + " -and -not $nm.EndsWith('.disabled')){ "
                "      Rename-Item -Path $_.FullName -NewName ($nm+'.disabled') -Force; $changed++ "
                "    } "
                "  } "
                " } "
                "}; "
                "@{ok=$true; changed=$changed; items=$rows} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=30)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "query": name_query, "result": json.loads(out)})
            return self._error(out or f"{mode_norm} failed")

        if mode_norm in {"detect_new", "watch_new"}:
            ok, current, err = _startup_snapshot()
            if not ok:
                return self._error(err)

            state_dir = Path.home() / ".Mudabbir"
            state_dir.mkdir(parents=True, exist_ok=True)
            state_file = state_dir / "startup_watch_snapshot.json"
            previous: list[dict[str, Any]] = []
            if state_file.exists():
                try:
                    previous = json.loads(state_file.read_text(encoding="utf-8"))
                    if isinstance(previous, dict):
                        previous = [previous]
                    if not isinstance(previous, list):
                        previous = []
                except Exception:
                    previous = []
            prev_keys = {
                (
                    str(item.get("Name", "") or ""),
                    str(item.get("Command", "") or ""),
                    str(item.get("Location", "") or ""),
                )
                for item in previous
                if isinstance(item, dict)
            }
            added = []
            for item in current:
                key = (
                    str(item.get("Name", "") or ""),
                    str(item.get("Command", "") or ""),
                    str(item.get("Location", "") or ""),
                )
                if key not in prev_keys:
                    added.append(item)
            try:
                state_file.write_text(json.dumps(current, ensure_ascii=False, indent=2), encoding="utf-8")
            except Exception:
                pass
            return _json({"ok": True, "mode": mode_norm, "added_count": len(added), "added": added[:40]})

        if mode_norm in {"watch_new_live", "watch_live", "monitor_new_startup"}:
            ok, baseline, err = _startup_snapshot()
            if not ok:
                return self._error(err)
            try:
                interval_sec = float(seconds if seconds is not None else 8.0)
            except Exception:
                interval_sec = 8.0
            try:
                duration_sec = float(monitor_seconds if monitor_seconds is not None else 90.0)
            except Exception:
                duration_sec = 90.0
            interval_sec = max(2.0, min(120.0, interval_sec))
            duration_sec = max(interval_sec, min(1800.0, duration_sec))
            end_time = time.time() + duration_sec

            seen = {
                (
                    str(item.get("Name", "") or ""),
                    str(item.get("Command", "") or ""),
                    str(item.get("Location", "") or ""),
                )
                for item in baseline
                if isinstance(item, dict)
            }
            detected: list[dict[str, Any]] = []
            rounds = 0
            while time.time() < end_time:
                rounds += 1
                time.sleep(interval_sec)
                ok_now, current, err_now = _startup_snapshot()
                if not ok_now:
                    return self._error(err_now)
                for item in current:
                    key = (
                        str(item.get("Name", "") or ""),
                        str(item.get("Command", "") or ""),
                        str(item.get("Location", "") or ""),
                    )
                    if key in seen:
                        continue
                    seen.add(key)
                    detected.append(item)
                    if notify:
                        label = str(item.get("Name") or "New startup item")
                        self._automation_tools("popup", text=f"Startup item detected: {label}")
            return _json(
                {
                    "ok": True,
                    "mode": mode_norm,
                    "duration_seconds": duration_sec,
                    "interval_seconds": interval_sec,
                    "poll_rounds": rounds,
                    "added_count": len(detected),
                    "added": detected[:80],
                }
            )

        if mode_norm in {"signature_check", "security_scan"}:
            ps = (
                "$items=Get-CimInstance Win32_StartupCommand | Select-Object Name,Command,Location,User; "
                "$rows=@(); "
                "foreach($i in $items){ "
                " $cmd=[string]$i.Command; "
                " $exe=$null; "
                " if($cmd -match '^[\"'']([^\"'']+\\.(exe|bat|cmd|ps1))'){ $exe=$matches[1] } "
                " elseif($cmd -match '^([^ ]+\\.(exe|bat|cmd|ps1))'){ $exe=$matches[1] }; "
                " $status='unknown'; "
                " if($exe -and (Test-Path $exe)){ "
                "   try { $sig=Get-AuthenticodeSignature -FilePath $exe -ErrorAction Stop; $status=[string]$sig.Status } catch { $status='error' } "
                " }; "
                " $rows += [pscustomobject]@{name=$i.Name;command=$cmd;location=$i.Location;signature_status=$status;path=$exe} "
                "}; "
                "$rows | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=30)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "items": json.loads(out)})
            return self._error(out or "startup signature check failed")
        if mode_norm in {"full_audit", "startup_full_audit"}:
            base_raw = self._startup_tools("list")
            sig_raw = self._startup_tools("signature_check")
            reg_raw = self._startup_tools("registry_startups")
            folder_raw = self._startup_tools("folder_startups")
            def _as_obj(raw: Any) -> dict[str, Any]:
                if not isinstance(raw, str):
                    return {}
                if raw.lower().startswith("error:"):
                    return {"ok": False, "error": raw}
                try:
                    parsed = json.loads(raw)
                    if isinstance(parsed, dict):
                        return parsed
                    if isinstance(parsed, list):
                        return {"ok": True, "items": parsed}
                    return {"ok": True, "value": parsed}
                except Exception:
                    return {"ok": False, "error": "non-json response", "raw": raw[:800]}

            base = _as_obj(base_raw)
            sig = _as_obj(sig_raw)
            reg = _as_obj(reg_raw)
            folder = _as_obj(folder_raw)
            sig_items = list(sig.get("items") or [])
            risky = []
            for item in sig_items:
                try:
                    status = str(item.get("signature_status") or "").strip().lower()
                    if status in {"error", "notsigned", "hashmismatch", "unknownerror"}:
                        risky.append(item)
                except Exception:
                    continue
            return _json(
                {
                    "ok": True,
                    "mode": mode_norm,
                    "startup_count": len(list(base.get("items") or [])),
                    "registry_count": len(list(reg.get("items") or [])),
                    "folder_count": len(list(folder.get("items") or [])),
                    "signature_issues_count": len(risky),
                    "signature_issues": risky[:50],
                    "startup_items": list(base.get("items") or [])[:120],
                    "partial_errors": {
                        "base": base.get("error"),
                        "signature": sig.get("error"),
                        "registry": reg.get("error"),
                        "folder": folder.get("error"),
                    },
                }
            )

        return self._error(f"unsupported startup_tools mode: {mode_norm}")

    def _clipboard_tools(self, mode: str) -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm == "clear":
            ok, out = _run_powershell("Set-Clipboard -Value ''", timeout=8)
            return _json({"ok": True, "mode": "clear"}) if ok else self._error(out or "clear clipboard failed")
        if mode_norm == "history":
            try:
                import pyautogui

                pyautogui.hotkey("win", "v")
                return _json({"ok": True, "mode": "history"})
            except Exception as exc:
                return self._error(f"clipboard history failed: {exc}")
        return self._error(f"unsupported clipboard_tools mode: {mode_norm}")

    def _browser_control(self, mode: str) -> str:
        mode_norm = (mode or "").strip().lower()
        mapping = {
            "new_tab": ("ctrl", "t"),
            "close_tab": ("ctrl", "w"),
            "reopen_tab": ("ctrl", "shift", "t"),
            "next_tab": ("ctrl", "tab"),
            "prev_tab": ("ctrl", "shift", "tab"),
            "reload": ("ctrl", "r"),
            "incognito": ("ctrl", "shift", "n"),
            "history": ("ctrl", "h"),
            "downloads": ("ctrl", "j"),
            "find": ("ctrl", "f"),
            "zoom_in": ("ctrl", "+"),
            "zoom_out": ("ctrl", "-"),
            "zoom_reset": ("ctrl", "0"),
            "save_pdf": ("ctrl", "p"),
            "home": ("alt", "home"),
        }
        hotkey = mapping.get(mode_norm)
        if not hotkey:
            return self._error(f"unsupported browser_control mode: {mode_norm}")
        try:
            import pyautogui

            pyautogui.hotkey(*hotkey)
            return _json({"ok": True, "mode": mode_norm, "keys": list(hotkey)})
        except Exception as exc:
            return self._error(f"browser control failed: {exc}")

    def _user_tools(self, mode: str, username: str = "", password: str = "", group: str = "") -> str:
        mode_norm = (mode or "").strip().lower()
        user = (username or "").strip()
        if mode_norm == "list":
            ok, out = _run_powershell(
                "Get-LocalUser | Select-Object Name,Enabled,LastLogon,PasswordLastSet | ConvertTo-Json -Compress",
                timeout=15,
            )
            if ok and out:
                return _json({"ok": True, "mode": "list", "data": json.loads(out)})
            return self._error(out or "list users failed")
        if mode_norm == "create":
            if not user or not password:
                return self._error("username and password are required")
            ps = (
                "$p=ConvertTo-SecureString '{pwd}' -AsPlainText -Force; "
                "New-LocalUser -Name '{usr}' -Password $p -FullName '{usr}' -ErrorAction Stop; "
                "@{{ok=$true; mode='create'; user='{usr}'}} | ConvertTo-Json -Compress"
            ).format(usr=user, pwd=password.replace("'", "''"))
            ok, out = _run_powershell(ps, timeout=20)
            if ok and out:
                return out
            return self._error(out or "create user failed")
        if mode_norm == "delete":
            if not user:
                return self._error("username is required")
            ok, out = _run_powershell(f"Remove-LocalUser -Name '{user}' -ErrorAction Stop", timeout=20)
            return _json({"ok": True, "mode": "delete", "user": user}) if ok else self._error(out or "delete user failed")
        if mode_norm == "set_password":
            if not user or not password:
                return self._error("username and password are required")
            ps = (
                "$p=ConvertTo-SecureString '{pwd}' -AsPlainText -Force; "
                "Set-LocalUser -Name '{usr}' -Password $p -ErrorAction Stop; "
                "@{{ok=$true; mode='set_password'; user='{usr}'}} | ConvertTo-Json -Compress"
            ).format(usr=user, pwd=password.replace("'", "''"))
            ok, out = _run_powershell(ps, timeout=20)
            return out if ok and out else self._error(out or "set password failed")
        if mode_norm == "set_type":
            if not user:
                return self._error("username is required")
            grp = (group or "").strip().lower()
            local_group = "Administrators" if grp in {"admin", "administrator", "administrators"} else "Users"
            ok, out = _run_powershell(f"Add-LocalGroupMember -Group '{local_group}' -Member '{user}' -ErrorAction Stop", timeout=20)
            return _json({"ok": True, "mode": "set_type", "user": user, "group": local_group}) if ok else self._error(out or "set user type failed")
        return self._error(f"unsupported user_tools mode: {mode_norm}")

    def _task_tools(self, mode: str, name: str = "", command: str = "", trigger: str = "") -> str:
        mode_norm = (mode or "").strip().lower()
        task_name = (name or "").strip()
        if mode_norm == "list":
            ok, out = _run_powershell(
                "Get-ScheduledTask | Select-Object -First 120 TaskName,TaskPath,State | ConvertTo-Json -Compress",
                timeout=20,
            )
            if ok and out:
                return _json({"ok": True, "mode": "list", "data": json.loads(out)})
            return self._error(out or "list tasks failed")
        if mode_norm == "running":
            ok, out = _run_powershell(
                "Get-ScheduledTask | Where-Object {$_.State -eq 'Running'} | "
                "Select-Object -First 120 TaskName,TaskPath,State | ConvertTo-Json -Compress",
                timeout=20,
            )
            if ok and out:
                return _json({"ok": True, "mode": "running", "data": json.loads(out)})
            return self._error(out or "running tasks query failed")
        if mode_norm == "last_run":
            ok, out = _run_powershell(
                "Get-ScheduledTask | Select-Object -First 120 TaskName,TaskPath,@{Name='LastRunTime';Expression={"
                "(Get-ScheduledTaskInfo -TaskName $_.TaskName -TaskPath $_.TaskPath -ErrorAction SilentlyContinue).LastRunTime"
                "}} | ConvertTo-Json -Compress",
                timeout=25,
            )
            if ok and out:
                return _json({"ok": True, "mode": "last_run", "data": json.loads(out)})
            return self._error(out or "last run tasks query failed")
        if not task_name:
            return self._error("task name is required")
        if mode_norm == "run":
            ok, out = _run_powershell(f"Start-ScheduledTask -TaskName '{task_name}'", timeout=15)
            return _json({"ok": True, "mode": "run", "name": task_name}) if ok else self._error(out or "run task failed")
        if mode_norm == "end":
            ok, out = _run_powershell(f"Stop-ScheduledTask -TaskName '{task_name}'", timeout=15)
            return _json({"ok": True, "mode": "end", "name": task_name}) if ok else self._error(out or "end task failed")
        if mode_norm == "enable":
            ok, out = _run_powershell(f"Enable-ScheduledTask -TaskName '{task_name}'", timeout=15)
            return _json({"ok": True, "mode": "enable", "name": task_name}) if ok else self._error(out or "enable task failed")
        if mode_norm == "disable":
            ok, out = _run_powershell(f"Disable-ScheduledTask -TaskName '{task_name}'", timeout=15)
            return _json({"ok": True, "mode": "disable", "name": task_name}) if ok else self._error(out or "disable task failed")
        if mode_norm == "delete":
            ok, out = _run_powershell(f"Unregister-ScheduledTask -TaskName '{task_name}' -Confirm:$false", timeout=20)
            return _json({"ok": True, "mode": "delete", "name": task_name}) if ok else self._error(out or "delete task failed")
        if mode_norm == "create":
            cmd = (command or "").strip()
            trig = (trigger or "").strip() or "DAILY"
            if not cmd:
                return self._error("command is required")
            ps = (
                "$act=New-ScheduledTaskAction -Execute 'cmd.exe' -Argument '/c {cmd}'; "
                "$trg=New-ScheduledTaskTrigger -Daily -At 09:00; "
                "Register-ScheduledTask -TaskName '{name}' -Action $act -Trigger $trg -Force | Out-Null; "
                "@{{ok=$true; mode='create'; name='{name}'; trigger='{trig}'}} | ConvertTo-Json -Compress"
            ).format(name=task_name.replace("'", "''"), cmd=cmd.replace("'", "''"), trig=trig)
            ok, out = _run_powershell(ps, timeout=25)
            return out if ok and out else self._error(out or "create task failed")
        return self._error(f"unsupported task_tools mode: {mode_norm}")

    def _registry_tools(
        self,
        mode: str,
        key: str = "",
        value_name: str = "",
        value_data: str = "",
        value_type: str = "REG_SZ",
    ) -> str:
        mode_norm = (mode or "").strip().lower()
        reg_key = (key or "").strip()
        if not reg_key:
            return self._error("key is required")
        if mode_norm == "query":
            ok, out = _run_powershell(f"Get-ItemProperty -Path '{reg_key}' | ConvertTo-Json -Compress", timeout=15)
            if ok and out:
                return _json({"ok": True, "mode": "query", "key": reg_key, "data": json.loads(out)})
            return self._error(out or "registry query failed")
        if mode_norm == "add_key":
            ok, out = _run_powershell(f"New-Item -Path '{reg_key}' -Force | Out-Null; 'ok'", timeout=15)
            return _json({"ok": True, "mode": "add_key", "key": reg_key}) if ok else self._error(out or "registry add key failed")
        if mode_norm == "delete_key":
            ok, out = _run_powershell(f"Remove-Item -Path '{reg_key}' -Recurse -Force", timeout=15)
            return _json({"ok": True, "mode": "delete_key", "key": reg_key}) if ok else self._error(out or "registry delete key failed")
        if mode_norm == "set_value":
            if not value_name:
                return self._error("value_name is required")
            ps = (
                f"New-ItemProperty -Path '{reg_key}' -Name '{value_name}' -Value '{value_data}' "
                f"-PropertyType {value_type} -Force | Out-Null; 'ok'"
            )
            ok, out = _run_powershell(ps, timeout=15)
            return _json({"ok": True, "mode": "set_value", "key": reg_key, "name": value_name}) if ok else self._error(out or "registry set value failed")
        if mode_norm == "backup":
            export_path = Path.home() / f"registry_backup_{_timestamp_id()}.reg"
            ok, out = _run_powershell(f"reg export \"{reg_key}\" \"{export_path}\" /y", timeout=20)
            return _json({"ok": True, "mode": mode_norm, "key": reg_key, "path": str(export_path)}) if ok else self._error(out or "registry backup failed")
        if mode_norm == "restore":
            import_path = Path(value_data).expanduser() if value_data else Path.home() / "registry_backup.reg"
            if not import_path.exists():
                return self._error(f"backup file not found: {import_path}")
            ok, out = _run_powershell(f"reg import \"{import_path}\"", timeout=20)
            return _json({"ok": True, "mode": mode_norm, "path": str(import_path)}) if ok else self._error(out or "registry restore failed")
        return self._error(f"unsupported registry_tools mode: {mode_norm}")

    def _disk_tools(self, mode: str, drive: str = "") -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm == "smart_status":
            ps = (
                "Get-PhysicalDisk | Select-Object FriendlyName,HealthStatus,OperationalStatus,MediaType,Size "
                "| ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=15)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "smart status failed")
        if mode_norm == "temp_files_clean":
            ps = (
                "$paths=@($env:TEMP,'C:\\Windows\\Temp'); "
                "foreach($p in $paths){ if(Test-Path $p){ Get-ChildItem $p -Recurse -Force -ErrorAction SilentlyContinue | "
                "Remove-Item -Recurse -Force -ErrorAction SilentlyContinue }}; "
                "@{ok=$true; mode='temp_files_clean'} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=30)
            return out if ok and out else self._error(out or "temp clean failed")
        if mode_norm == "disk_usage":
            ps = (
                "Get-PSDrive -PSProvider FileSystem | Select-Object Name,Used,Free "
                "| ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=12)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "disk usage failed")
        if mode_norm == "chkdsk_scan":
            drv = (drive or "C:").strip()
            ok, out = _run_powershell(f"chkdsk {drv}", timeout=40)
            return _json({"ok": True, "mode": mode_norm, "drive": drv, "output": out[:1400]}) if ok else self._error(out or "chkdsk failed")
        if mode_norm == "prefetch_clean":
            ps = (
                "$p='C:\\Windows\\Prefetch'; "
                "if(Test-Path $p){Get-ChildItem $p -Force -ErrorAction SilentlyContinue | Remove-Item -Force -ErrorAction SilentlyContinue}; "
                "@{ok=$true; mode='prefetch_clean'} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=25)
            return out if ok and out else self._error(out or "prefetch clean failed")
        if mode_norm == "logs_clean":
            ps = (
                "wevtutil el | Foreach-Object {wevtutil cl $_}; "
                "@{ok=$true; mode='logs_clean'} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=45)
            return out if ok and out else self._error(out or "logs clean failed")
        if mode_norm == "defrag":
            drv = (drive or "C:").strip()
            ok, out = _run_powershell(f"defrag {drv} /O", timeout=180)
            return _json({"ok": True, "mode": mode_norm, "drive": drv, "output": out[:1800]}) if ok else self._error(out or "defrag failed")
        return self._error(f"unsupported disk_tools mode: {mode_norm}")

    def _security_tools(self, mode: str, target: str = "", port: Any = None, rule_name: str = "") -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm == "firewall_status":
            ok, out = _run_powershell("Get-NetFirewallProfile | Select-Object Name,Enabled | ConvertTo-Json -Compress", timeout=12)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "firewall status failed")
        if mode_norm == "firewall_enable":
            ok, out = _run_powershell("Set-NetFirewallProfile -Profile Domain,Public,Private -Enabled True", timeout=12)
            return _json({"ok": True, "mode": mode_norm}) if ok else self._error(out or "firewall enable failed")
        if mode_norm == "firewall_disable":
            ok, out = _run_powershell("Set-NetFirewallProfile -Profile Domain,Public,Private -Enabled False", timeout=12)
            return _json({"ok": True, "mode": mode_norm}) if ok else self._error(out or "firewall disable failed")
        if mode_norm == "block_port":
            try:
                p = int(port)
            except Exception:
                return self._error("valid port is required")
            rn = (rule_name or f"Mudabbir_Block_{p}").strip()
            ps = (
                f"New-NetFirewallRule -DisplayName '{rn}' -Direction Inbound -Action Block "
                f"-Protocol TCP -LocalPort {p} | Out-Null; 'ok'"
            )
            ok, out = _run_powershell(ps, timeout=15)
            return _json({"ok": True, "mode": mode_norm, "port": p, "rule_name": rn}) if ok else self._error(out or "block port failed")
        if mode_norm == "unblock_rule":
            rn = (rule_name or target).strip()
            if not rn:
                return self._error("rule_name is required")
            ok, out = _run_powershell(f"Remove-NetFirewallRule -DisplayName '{rn}'", timeout=15)
            return _json({"ok": True, "mode": mode_norm, "rule_name": rn}) if ok else self._error(out or "remove firewall rule failed")
        if mode_norm == "disable_usb":
            ps = (
                "Set-ItemProperty -Path 'HKLM:\\SYSTEM\\CurrentControlSet\\Services\\USBSTOR' -Name Start -Value 4; "
                "@{ok=$true; mode='disable_usb'} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=12)
            return out if ok and out else self._error(out or "disable usb failed")
        if mode_norm == "enable_usb":
            ps = (
                "Set-ItemProperty -Path 'HKLM:\\SYSTEM\\CurrentControlSet\\Services\\USBSTOR' -Name Start -Value 3; "
                "@{ok=$true; mode='enable_usb'} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=12)
            return out if ok and out else self._error(out or "enable usb failed")
        if mode_norm == "disable_camera":
            ps = (
                "Get-PnpDevice -Class Camera | Disable-PnpDevice -Confirm:$false; "
                "@{ok=$true; mode='disable_camera'} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=20)
            return out if ok and out else self._error(out or "disable camera failed")
        if mode_norm == "enable_camera":
            ps = (
                "Get-PnpDevice -Class Camera | Enable-PnpDevice -Confirm:$false; "
                "@{ok=$true; mode='enable_camera'} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=20)
            return out if ok and out else self._error(out or "enable camera failed")
        if mode_norm in {"logged_in_users", "sessions"}:
            ok, out = _run_powershell("quser", timeout=12)
            if ok:
                lines = [ln.rstrip() for ln in (out or "").splitlines() if ln.strip()]
                return _json({"ok": True, "mode": mode_norm, "lines": lines})
            return self._error(out or "logged-in users query failed")
        if mode_norm in {"remote_sessions_list", "list_remote_sessions"}:
            ps = (
                "$raw=quser 2>$null; "
                "if(-not $raw){ @{ok=$true; mode='remote_sessions_list'; count=0; sessions=@()} | ConvertTo-Json -Compress; exit 0 }; "
                "$rows=@(); "
                "foreach($line in $raw | Select-Object -Skip 1){ "
                "  $parts=($line -replace '^\\s+','' -replace '\\s+',' ').Split(' '); "
                "  if($parts.Length -ge 5){ "
                "    $rows += [pscustomobject]@{user=$parts[0]; session_name=$parts[1]; session_id=$parts[2]; state=$parts[3]; idle=$parts[4]} "
                "  } "
                "}; "
                "$remote=$rows | Where-Object {($_.session_name -like 'rdp*') -or ($_.session_name -like 'console') -or ($_.state -eq 'Active')}; "
                "@{ok=$true; mode='remote_sessions_list'; count=$remote.Count; sessions=$remote} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=18)
            if ok and out:
                try:
                    data = json.loads(out)
                    sessions = data.get("sessions")
                    if isinstance(sessions, dict):
                        sessions = [sessions]
                    if not isinstance(sessions, list):
                        sessions = []
                    data["sessions"] = sessions
                    data["count"] = len(sessions)
                    return _json(data)
                except Exception:
                    return out
            return self._error(out or "remote sessions list failed")
        if mode_norm in {"recent_files", "recent_files_list"}:
            recent_dir = Path.home() / "AppData/Roaming/Microsoft/Windows/Recent"
            if not recent_dir.exists():
                return _json({"ok": True, "mode": mode_norm, "count": 0, "items": []})
            items = []
            for p in recent_dir.glob("*"):
                try:
                    if not p.is_file():
                        continue
                    items.append(
                        {
                            "name": p.name,
                            "path": str(p),
                            "modified": datetime.fromtimestamp(p.stat().st_mtime).isoformat(),
                        }
                    )
                except Exception:
                    continue
            items.sort(key=lambda x: str(x.get("modified", "")), reverse=True)
            return _json({"ok": True, "mode": mode_norm, "count": len(items), "items": items[:150]})
        if mode_norm in {"recent_files_clear", "clear_recent_files"}:
            recent_dir = Path.home() / "AppData/Roaming/Microsoft/Windows/Recent"
            removed = 0
            if recent_dir.exists():
                for p in recent_dir.glob("*"):
                    try:
                        if p.is_file():
                            p.unlink(missing_ok=True)
                            removed += 1
                    except Exception:
                        continue
            return _json({"ok": True, "mode": mode_norm, "removed": removed})
        if mode_norm in {"current_connections_ips", "connections"}:
            ps = (
                "Get-NetTCPConnection -State Established -ErrorAction SilentlyContinue | "
                "Select-Object LocalAddress,LocalPort,RemoteAddress,RemotePort,OwningProcess | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=15)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "items": json.loads(out)})
            return self._error(out or "connections query failed")
        if mode_norm in {"admin_processes", "elevated_processes"}:
            ps = (
                "$rows=@(); "
                "Get-CimInstance Win32_Process | ForEach-Object { "
                "  $o=$_.GetOwner(); "
                "  if($o -and ($o.User -in @('SYSTEM','Administrator'))){ "
                "    $rows += [pscustomobject]@{pid=$_.ProcessId;name=$_.Name;user=$o.User;domain=$o.Domain} "
                "  } "
                "}; "
                "$rows | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=20)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "items": json.loads(out)})
            return self._error(out or "admin process query failed")
        if mode_norm in {"failed_audit_logins", "failed_logins"}:
            ps = (
                "Get-WinEvent -FilterHashtable @{LogName='Security';ID=4625} -MaxEvents 50 | "
                "Select-Object TimeCreated,Id,LevelDisplayName,Message | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=25)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "items": json.loads(out)})
            return self._error(out or "failed login audit query failed")
        if mode_norm in {"close_remote_sessions", "logoff_remote"}:
            ps = (
                "$raw=quser 2>$null; "
                "if(-not $raw){ @{ok=$false;error='no sessions'} | ConvertTo-Json -Compress; exit 0 }; "
                "$killed=@(); "
                "foreach($line in $raw | Select-Object -Skip 1){ "
                "  $parts=($line -replace '^\\s+','' -replace '\\s+',' ').Split(' '); "
                "  if($parts.Length -ge 3){ "
                "    $id=$parts[2]; "
                "    if($id -match '^[0-9]+$'){ "
                "      try { logoff $id; $killed += $id } catch {} "
                "    } "
                "  } "
                "}; "
                "@{ok=$true; mode='close_remote_sessions'; closed=$killed} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=20)
            if ok and out:
                return out
            return self._error(out or "close remote sessions failed")
        if mode_norm in {"intrusion_summary", "audit_intrusion_summary", "failed_login_summary"}:
            ps = (
                "$events=Get-WinEvent -FilterHashtable @{LogName='Security';ID=4625} -MaxEvents 200 -ErrorAction SilentlyContinue; "
                "$users=@{}; $ips=@{}; $statuses=@{}; $recent=@(); "
                "foreach($ev in $events){ "
                "  try{ "
                "    $xml=[xml]$ev.ToXml(); "
                "    $map=@{}; foreach($d in $xml.Event.EventData.Data){ $map[$d.Name]=$d.'#text' }; "
                "    $u=[string]($map['TargetUserName']); if(-not [string]::IsNullOrWhiteSpace($u)){ if(-not $users.ContainsKey($u)){$users[$u]=0}; $users[$u]++ }; "
                "    $ip=[string]($map['IpAddress']); if(-not [string]::IsNullOrWhiteSpace($ip) -and $ip -ne '-' -and $ip -ne '::1' -and $ip -ne '127.0.0.1'){ if(-not $ips.ContainsKey($ip)){$ips[$ip]=0}; $ips[$ip]++ }; "
                "    $st=[string]($map['Status']); if(-not [string]::IsNullOrWhiteSpace($st)){ if(-not $statuses.ContainsKey($st)){$statuses[$st]=0}; $statuses[$st]++ }; "
                "    if($recent.Count -lt 40){ "
                "      $recent += [pscustomobject]@{time=$ev.TimeCreated;user=$u;ip=$ip;status=$st;substatus=[string]($map['SubStatus']);workstation=[string]($map['WorkstationName'])} "
                "    } "
                "  } catch {} "
                "}; "
                "$topUsers = $users.GetEnumerator() | Sort-Object Value -Descending | Select-Object -First 10 | ForEach-Object { [pscustomobject]@{user=$_.Key;count=$_.Value} }; "
                "$topIps = $ips.GetEnumerator() | Sort-Object Value -Descending | Select-Object -First 10 | ForEach-Object { [pscustomobject]@{ip=$_.Key;count=$_.Value} }; "
                "$topStatuses = $statuses.GetEnumerator() | Sort-Object Value -Descending | Select-Object -First 10 | ForEach-Object { [pscustomobject]@{status=$_.Key;count=$_.Value} }; "
                "$lvl='low'; if($events.Count -ge 80 -or $topIps.Count -ge 5){$lvl='high'} elseif($events.Count -ge 20){$lvl='medium'}; "
                "@{ok=$true; mode='intrusion_summary'; failed_count=$events.Count; risk_level=$lvl; top_users=$topUsers; top_ips=$topIps; top_status_codes=$topStatuses; recent=$recent} | ConvertTo-Json -Compress -Depth 6"
            )
            ok, out = _run_powershell(ps, timeout=35)
            if ok and out:
                return out
            return self._error(out or "intrusion summary failed")
        return self._error(f"unsupported security_tools mode: {mode_norm}")

    def _web_tools(self, mode: str, url: str = "", city: str = "") -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm == "open_url":
            target = (url or "").strip()
            if not target:
                return self._error("url is required")
            subprocess.Popen(["cmd", "/c", "start", "", target], shell=False)
            return _json({"ok": True, "mode": mode_norm, "url": target})
        if mode_norm == "download_file":
            target = (url or "").strip()
            if not target:
                return self._error("url is required")
            media_dir = get_media_dir()
            out_name = Path(target.split("?")[0]).name or f"download_{_timestamp_id()}"
            out_path = media_dir / out_name
            ok, out = _run_powershell(
                f"Invoke-WebRequest -Uri '{target}' -OutFile '{out_path}' -UseBasicParsing",
                timeout=60,
            )
            return _json({"ok": True, "mode": mode_norm, "url": target, "path": str(out_path)}) if ok else self._error(out or "download failed")
        if mode_norm == "weather":
            c = (city or "").strip() or "amman"
            api_url = f"https://wttr.in/{c}?format=j1"
            ok, out = _run_powershell(f"(Invoke-RestMethod -Uri '{api_url}' -TimeoutSec 10) | ConvertTo-Json -Compress", timeout=20)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "city": c, "data": json.loads(out)})
            return self._error(out or "weather lookup failed")
        return self._error(f"unsupported web_tools mode: {mode_norm}")

    def _hardware_tools(self, mode: str, drive: str = "") -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm == "cpu_info":
            ps = (
                "Get-CimInstance Win32_Processor | "
                "Select-Object Name,NumberOfCores,NumberOfLogicalProcessors,MaxClockSpeed "
                "| ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=12)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "cpu info failed")
        if mode_norm == "cores_info":
            ps = (
                "Get-CimInstance Win32_Processor | "
                "Select-Object Name,NumberOfCores,NumberOfLogicalProcessors,CurrentClockSpeed,MaxClockSpeed "
                "| ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=12)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "cores info failed")
        if mode_norm == "gpu_info":
            ps = "Get-CimInstance Win32_VideoController | Select-Object Name,AdapterRAM,DriverVersion | ConvertTo-Json -Compress"
            ok, out = _run_powershell(ps, timeout=12)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "gpu info failed")
        if mode_norm == "gpu_temp":
            return self._performance_tools("gpu_temp")
        if mode_norm == "mobo_serial":
            ps = "Get-CimInstance Win32_BaseBoard | Select-Object Manufacturer,Product,SerialNumber | ConvertTo-Json -Compress"
            ok, out = _run_powershell(ps, timeout=12)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "motherboard serial failed")
        if mode_norm in {"mobo_model", "motherboard_model"}:
            ps = "Get-CimInstance Win32_BaseBoard | Select-Object Manufacturer,Product,Version | ConvertTo-Json -Compress"
            ok, out = _run_powershell(ps, timeout=12)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "motherboard model failed")
        if mode_norm == "ram_info":
            ps = "Get-CimInstance Win32_PhysicalMemory | Select-Object Manufacturer,Speed,Capacity,PartNumber | ConvertTo-Json -Compress"
            ok, out = _run_powershell(ps, timeout=12)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "ram info failed")
        if mode_norm == "ram_speed_type":
            ps = (
                "Get-CimInstance Win32_PhysicalMemory | "
                "Select-Object Manufacturer,Speed,ConfiguredClockSpeed,SMBIOSMemoryType,Capacity,PartNumber | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=12)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "ram speed/type failed")
        if mode_norm == "battery_report":
            out_file = Path.home() / "battery_report.html"
            ok, out = _run_powershell(f"powercfg /batteryreport /output \"{out_file}\"", timeout=20)
            if ok:
                return _json({"ok": True, "mode": mode_norm, "path": str(out_file)})
            return self._error(out or "battery report failed")
        if mode_norm == "battery_minutes":
            raw = self._battery_status()
            if raw.lower().startswith("error:"):
                return raw
            try:
                data = json.loads(raw)
                secs = data.get("secs_left")
                mins = int(round(float(secs) / 60.0)) if secs is not None else None
                return _json({"ok": True, "mode": mode_norm, "minutes_left": mins, "percent": data.get("percent")})
            except Exception as exc:
                return self._error(f"battery minutes failed: {exc}")
        if mode_norm == "battery_cycle_count":
            out_file = Path.home() / "battery_report.html"
            ok, out = _run_powershell(f"powercfg /batteryreport /output \"{out_file}\"", timeout=20)
            if not ok:
                return self._error(out or "battery report failed")
            try:
                html = out_file.read_text(encoding="utf-8", errors="ignore")
                m = re.search(r"Cycle count\\s*</td>\\s*<td[^>]*>\\s*([0-9]+)\\s*</td>", html, flags=re.IGNORECASE)
                if not m:
                    m = re.search(r"Cycle Count\\s*:?\\s*([0-9]+)", html, flags=re.IGNORECASE)
                count = int(m.group(1)) if m else None
                return _json({"ok": True, "mode": mode_norm, "cycle_count": count, "path": str(out_file)})
            except Exception as exc:
                return self._error(f"battery cycle count parse failed: {exc}")
        if mode_norm == "smart_status":
            return self._disk_tools("smart_status", drive=drive)
        return self._error(f"unsupported hardware_tools mode: {mode_norm}")

    def _update_tools(self, mode: str, target: str = "") -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm == "list_updates":
            ok, out = _run_powershell("Get-HotFix | Select-Object -First 120 HotFixID,InstalledOn,Description | ConvertTo-Json -Compress", timeout=20)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "list updates failed")
        if mode_norm == "last_update_time":
            ps = (
                "Get-HotFix | Sort-Object InstalledOn -Descending | Select-Object -First 1 HotFixID,InstalledOn,Description | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=18)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "last update lookup failed")
        if mode_norm == "check_updates":
            ok, out = _run_powershell("UsoClient StartScan", timeout=12)
            return _json({"ok": True, "mode": mode_norm, "status": "scan_started"}) if ok else self._error(out or "check updates failed")
        if mode_norm == "install_kb":
            kb = (target or "").strip()
            if not kb:
                return self._error("target KB is required")
            cmd = f"wusa.exe /install /quiet /norestart {kb}"
            ok, out = _run_powershell(cmd, timeout=30)
            return _json({"ok": True, "mode": mode_norm, "kb": kb}) if ok else self._error(out or "install kb failed")
        if mode_norm == "winsxs_cleanup":
            cmd = "Dism.exe /online /Cleanup-Image /StartComponentCleanup"
            ok, out = _run_powershell(cmd, timeout=120)
            return _json({"ok": True, "mode": mode_norm, "output": out[:1200]}) if ok else self._error(out or "winsxs cleanup failed")
        if mode_norm in {"stop_background_updates", "stop_updates"}:
            ps = (
                "Stop-Service -Name wuauserv -Force -ErrorAction SilentlyContinue; "
                "Stop-Service -Name bits -Force -ErrorAction SilentlyContinue; "
                "Stop-Service -Name usosvc -Force -ErrorAction SilentlyContinue; "
                "@{ok=$true; mode='stop_background_updates'} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=20)
            if ok and out:
                return out
            return self._error(out or "stop background updates failed")
        return self._error(f"unsupported update_tools mode: {mode_norm}")

    def _ui_tools(self, mode: str) -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm in {"dark_mode", "light_mode"}:
            val = 0 if mode_norm == "dark_mode" else 1
            ps = (
                "$k='HKCU:\\Software\\Microsoft\\Windows\\CurrentVersion\\Themes\\Personalize'; "
                f"Set-ItemProperty -Path $k -Name AppsUseLightTheme -Value {val}; "
                f"Set-ItemProperty -Path $k -Name SystemUsesLightTheme -Value {val}; "
                f"@{{ok=$true; mode='{mode_norm}'}} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=12)
            return out if ok and out else self._error(out or "theme switch failed")
        if mode_norm in {"transparency_on", "transparency_off"}:
            val = 1 if mode_norm.endswith("_on") else 0
            ps = (
                "$k='HKCU:\\Software\\Microsoft\\Windows\\CurrentVersion\\Themes\\Personalize'; "
                f"Set-ItemProperty -Path $k -Name EnableTransparency -Value {val}; "
                f"@{{ok=$true; mode='{mode_norm}'}} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=12)
            return out if ok and out else self._error(out or "transparency toggle failed")
        if mode_norm in {"taskbar_autohide_on", "taskbar_autohide_off"}:
            # best-effort registry switch + explorer restart
            val = 3 if mode_norm.endswith("_on") else 2
            ps = (
                "$p='HKCU:\\Software\\Microsoft\\Windows\\CurrentVersion\\Explorer\\StuckRects3'; "
                "$v=(Get-ItemProperty -Path $p -Name Settings).Settings; "
                f"$v[8]={val}; Set-ItemProperty -Path $p -Name Settings -Value $v; "
                "Stop-Process -Name explorer -Force; "
                f"@{{ok=$true; mode='{mode_norm}'}} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=20)
            return out if ok and out else self._error(out or "taskbar autohide toggle failed")
        if mode_norm in {"night_light_on", "night_light_off"}:
            val = 1 if mode_norm.endswith("_on") else 0
            ps = (
                "$k='HKCU:\\Software\\Microsoft\\Windows\\CurrentVersion\\CloudStore\\Store\\DefaultAccount\\Current\\default$windows.data.bluelightreduction.bluelightreductionstate\\windows.data.bluelightreduction.bluelightreductionstate'; "
                f"@{{ok=$true; mode='{mode_norm}'; hint='registry-managed'}} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=10)
            return out if ok and out else self._error(out or "night light toggle failed")
        if mode_norm in {"desktop_icons_show", "desktop_icons_hide"}:
            show = "$true" if mode_norm.endswith("_show") else "$false"
            ps = (
                "$sig='[DllImport(\"user32.dll\")]public static extern IntPtr FindWindow(string lpClassName,string lpWindowName);"
                "[DllImport(\"user32.dll\")]public static extern IntPtr SendMessage(IntPtr hWnd,int Msg,IntPtr wParam,IntPtr lParam);'; "
                "Add-Type -Name NativeMethods -Namespace Win32 -MemberDefinition $sig; "
                "$h=[Win32.NativeMethods]::FindWindow('Progman',$null); "
                "[Win32.NativeMethods]::SendMessage($h,0x111,[intptr]0x7402,[intptr]0) | Out-Null; "
                f"@{{ok=$true; mode='{mode_norm}'}} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=10)
            return out if ok and out else self._error(out or "desktop icons toggle failed")
        if mode_norm in {"game_mode_on", "game_mode_off"}:
            val = 1 if mode_norm.endswith("_on") else 0
            ps = (
                "$k='HKCU:\\Software\\Microsoft\\GameBar'; "
                "if(!(Test-Path $k)){ New-Item -Path $k -Force | Out-Null }; "
                f"Set-ItemProperty -Path $k -Name AllowAutoGameMode -Value {val}; "
                f"Set-ItemProperty -Path $k -Name AutoGameModeEnabled -Value {val}; "
                f"@{{ok=$true; mode='{mode_norm}'}} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=12)
            return out if ok and out else self._error(out or "game mode toggle failed")
        return self._error(f"unsupported ui_tools mode: {mode_norm}")

    def _automation_tools(
        self,
        mode: str,
        seconds: Any = None,
        monitor_seconds: Any = None,
        text: str = "",
        key: str = "",
        repeat_count: Any = None,
        x: Any = None,
        y: Any = None,
        x2: Any = None,
        y2: Any = None,
        width: Any = None,
        height: Any = None,
    ) -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm == "delay":
            try:
                sec = float(seconds if seconds is not None else 1.0)
            except Exception:
                sec = 1.0
            sec = max(0.0, min(120.0, sec))
            time.sleep(sec)
            return _json({"ok": True, "mode": mode_norm, "seconds": sec})
        if mode_norm == "popup":
            msg = text.strip() or "Mudabbir"
            ok, out = _run_powershell(f"[System.Windows.MessageBox]::Show('{msg}') | Out-Null; 'ok'", timeout=20)
            return _json({"ok": True, "mode": mode_norm}) if ok else self._error(out or "popup failed")
        if mode_norm == "tts":
            msg = text.strip()
            if not msg:
                return self._error("text is required")
            escaped_msg = msg.replace("'", "''")
            ps = (
                "Add-Type -AssemblyName System.Speech; "
                "$s=New-Object System.Speech.Synthesis.SpeechSynthesizer; "
                f"$s.Speak('{escaped_msg}'); "
                "@{ok=$true; mode='tts'} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=30)
            return out if ok and out else self._error(out or "tts failed")
        if mode_norm == "repeat_key":
            k = (key or "").strip().lower()
            if not k:
                return self._error("key is required")
            try:
                count = int(repeat_count if repeat_count is not None else 1)
            except Exception:
                count = 1
            count = max(1, min(200, count))
            try:
                import pyautogui

                pyautogui.FAILSAFE = False
                for _ in range(count):
                    pyautogui.press(k)
                return _json({"ok": True, "mode": mode_norm, "key": k, "count": count})
            except Exception as exc:
                return self._error(f"repeat_key failed: {exc}")
        if mode_norm in {"mouse_down", "mouse_up"}:
            btn = (key or "left").strip().lower()
            if btn not in {"left", "right", "middle"}:
                btn = "left"
            try:
                import pyautogui

                pyautogui.FAILSAFE = False
                if mode_norm == "mouse_down":
                    pyautogui.mouseDown(button=btn)
                else:
                    pyautogui.mouseUp(button=btn)
                return _json({"ok": True, "mode": mode_norm, "button": btn})
            except Exception as exc:
                return self._error(f"{mode_norm} failed: {exc}")
        if mode_norm == "drag_drop":
            try:
                sx = int(x if x is not None else 0)
                sy = int(y if y is not None else 0)
                tx = int(x2 if x2 is not None else (width if width is not None else 0))
                ty = int(y2 if y2 is not None else (height if height is not None else 0))
            except Exception:
                return self._error("x,y,x2,y2 must be integers for drag_drop")
            try:
                import pyautogui

                pyautogui.FAILSAFE = False
                pyautogui.moveTo(sx, sy, duration=0.15)
                pyautogui.dragTo(tx, ty, duration=0.35, button="left")
                return _json({"ok": True, "mode": mode_norm, "from": {"x": sx, "y": sy}, "to": {"x": tx, "y": ty}})
            except Exception as exc:
                return self._error(f"drag_drop failed: {exc}")
        if mode_norm in {"scroll_up", "scroll_down"}:
            try:
                amount = int(repeat_count if repeat_count is not None else 4)
            except Exception:
                amount = 4
            amount = max(1, min(60, amount))
            wheel = amount * 120
            if mode_norm == "scroll_down":
                wheel *= -1
            try:
                import pyautogui

                pyautogui.FAILSAFE = False
                pyautogui.scroll(wheel)
                return _json({"ok": True, "mode": mode_norm, "amount": amount})
            except Exception as exc:
                return self._error(f"{mode_norm} failed: {exc}")
        if mode_norm == "move_corner":
            corner = (key or text or "top_left").strip().casefold()
            try:
                import pyautogui

                pyautogui.FAILSAFE = False
                sw, sh = pyautogui.size()
                mapping = {
                    "top_left": (0, 0),
                    "left_top": (0, 0),
                    "top_right": (max(0, sw - 1), 0),
                    "right_top": (max(0, sw - 1), 0),
                    "bottom_left": (0, max(0, sh - 1)),
                    "left_bottom": (0, max(0, sh - 1)),
                    "bottom_right": (max(0, sw - 1), max(0, sh - 1)),
                    "right_bottom": (max(0, sw - 1), max(0, sh - 1)),
                }
                px, py = mapping.get(corner, (0, 0))
                pyautogui.moveTo(px, py, duration=0.2)
                return _json({"ok": True, "mode": mode_norm, "corner": corner, "x": px, "y": py})
            except Exception as exc:
                return self._error(f"move_corner failed: {exc}")
        if mode_norm == "click_center":
            try:
                import pyautogui

                pyautogui.FAILSAFE = False
                sw, sh = pyautogui.size()
                px = int(sw // 2)
                py = int(sh // 2)
                pyautogui.click(px, py)
                return _json({"ok": True, "mode": mode_norm, "x": px, "y": py})
            except Exception as exc:
                return self._error(f"click_center failed: {exc}")
        if mode_norm == "mouse_speed_up":
            ps = (
                "$k='HKCU:\\Control Panel\\Mouse'; "
                "Set-ItemProperty -Path $k -Name MouseSensitivity -Value '20'; "
                "rundll32.exe user32.dll,UpdatePerUserSystemParameters; "
                "@{ok=$true; mode='mouse_speed_up'; value=20} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=12)
            return out if ok and out else self._error(out or "mouse speed update failed")
        if mode_norm in {"mouse_sonar_on", "mouse_sonar_off"}:
            sonar_value = 1 if mode_norm.endswith("_on") else 0
            ps = (
                "$k='HKCU:\\Control Panel\\Mouse'; "
                f"Set-ItemProperty -Path $k -Name Sonar -Value '{sonar_value}'; "
                "rundll32.exe user32.dll,UpdatePerUserSystemParameters; "
                f"@{{ok=$true; mode='{mode_norm}'; value={sonar_value}}} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=12)
            return out if ok and out else self._error(out or "mouse sonar update failed")
        if mode_norm == "mouse_keys_toggle":
            try:
                import pyautogui

                pyautogui.hotkey("left", "alt", "left", "shift", "numlock")
                return _json({"ok": True, "mode": mode_norm})
            except Exception as exc:
                return self._error(f"mouse_keys_toggle failed: {exc}")
        if mode_norm == "timer":
            try:
                sec = float(seconds if seconds is not None else 5.0)
            except Exception:
                sec = 5.0
            sec = max(1.0, min(3600.0, sec))
            time.sleep(sec)
            note = text.strip() or f"Timer finished ({int(sec)}s)"
            return self._automation_tools("popup", text=note)
        if mode_norm in {"mouse_lock_on", "mouse_lock_off"}:
            ps = (
                "$sig='[DllImport(\"user32.dll\")]public static extern bool ClipCursor(ref RECT rect);"
                "[DllImport(\"user32.dll\")]public static extern bool ClipCursor(IntPtr rect);"
                "[DllImport(\"user32.dll\")]public static extern int GetSystemMetrics(int nIndex);"
                "public struct RECT { public int Left; public int Top; public int Right; public int Bottom; }'; "
                "Add-Type -Name NativeMethods -Namespace Win32 -MemberDefinition $sig; "
            )
            if mode_norm.endswith("_on"):
                ps += (
                    "$r=New-Object Win32.NativeMethods+RECT; "
                    "$r.Left=0; $r.Top=0; "
                    "$r.Right=[Win32.NativeMethods]::GetSystemMetrics(78); "
                    "$r.Bottom=[Win32.NativeMethods]::GetSystemMetrics(79); "
                    "[void][Win32.NativeMethods]::ClipCursor([ref]$r); "
                    "@{ok=$true; mode='mouse_lock_on'} | ConvertTo-Json -Compress"
                )
            else:
                ps += (
                    "[void][Win32.NativeMethods]::ClipCursor([intptr]::Zero); "
                    "@{ok=$true; mode='mouse_lock_off'} | ConvertTo-Json -Compress"
                )
            ok, out = _run_powershell(ps, timeout=12)
            return out if ok and out else self._error(out or "mouse lock toggle failed")
        if mode_norm == "mouse_lock_region":
            try:
                rx = int(x if x is not None else 0)
                ry = int(y if y is not None else 0)
                rw = int(width if width is not None else 400)
                rh = int(height if height is not None else 300)
            except Exception:
                return self._error("x,y,width,height must be integers for mouse_lock_region")
            rw = max(10, rw)
            rh = max(10, rh)
            ps = (
                "$sig='[DllImport(\"user32.dll\")]public static extern bool ClipCursor(ref RECT rect);"
                "public struct RECT { public int Left; public int Top; public int Right; public int Bottom; }'; "
                "Add-Type -Name NativeMethods -Namespace Win32 -MemberDefinition $sig; "
                "$r=New-Object Win32.NativeMethods+RECT; "
                f"$r.Left={rx}; $r.Top={ry}; $r.Right={rx + rw}; $r.Bottom={ry + rh}; "
                "[void][Win32.NativeMethods]::ClipCursor([ref]$r); "
                "@{ok=$true; mode='mouse_lock_region'; x="
                + str(rx)
                + "; y="
                + str(ry)
                + "; width="
                + str(rw)
                + "; height="
                + str(rh)
                + "} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=12)
            return out if ok and out else self._error(out or "mouse lock region failed")
        if mode_norm == "battery_guard":
            try:
                threshold_pct = float(seconds if seconds is not None else 15.0)
            except Exception:
                threshold_pct = 15.0
            threshold_pct = max(1.0, min(99.0, threshold_pct))
            raw = self._battery_status()
            if raw.lower().startswith("error:"):
                return raw
            try:
                data = json.loads(raw)
                pct = float(data.get("percent") or 0.0)
                if pct <= threshold_pct and not bool(data.get("plugged")):
                    return self._system_power("power_plan_saver")
                return _json(
                    {
                        "ok": True,
                        "mode": mode_norm,
                        "changed": False,
                        "percent": pct,
                        "threshold": threshold_pct,
                        "plugged": bool(data.get("plugged")),
                    }
                )
            except Exception as exc:
                return self._error(f"battery_guard failed: {exc}")
        if mode_norm == "screen_off_mouse_lock":
            lock_result = self._automation_tools("mouse_lock_on")
            if isinstance(lock_result, str) and lock_result.lower().startswith("error:"):
                return lock_result
            return self._system_power("screen_off")
        if mode_norm in {"anti_idle_f5", "anti_idle", "f5_keepalive"}:
            try:
                interval_sec = float(seconds if seconds is not None else 60.0)
            except Exception:
                interval_sec = 60.0
            try:
                total_sec = float(monitor_seconds if monitor_seconds is not None else 300.0)
            except Exception:
                total_sec = 300.0
            interval_sec = max(5.0, min(600.0, interval_sec))
            total_sec = max(interval_sec, min(86400.0, total_sec))
            loops = max(1, int(total_sec // interval_sec))
            try:
                import pyautogui

                pyautogui.FAILSAFE = False
                for _ in range(loops):
                    pyautogui.press("f5")
                    time.sleep(interval_sec)
                return _json(
                    {
                        "ok": True,
                        "mode": mode_norm,
                        "interval_seconds": interval_sec,
                        "duration_seconds": total_sec,
                        "press_count": loops,
                    }
                )
            except Exception as exc:
                return self._error(f"anti idle failed: {exc}")
        return self._error(f"unsupported automation_tools mode: {mode_norm}")

    def _app_tools(self, mode: str, app: str = "", dry_run: bool = False, max_kill: Any = None) -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm == "open_default_browser":
            ok, out = _run_powershell("Start-Process 'https://www.bing.com'", timeout=10)
            return _json({"ok": True, "mode": mode_norm}) if ok else self._error(out or "open default browser failed")
        mapping = {
            "open_notepad": "notepad.exe",
            "open_calc": "calc.exe",
            "open_paint": "mspaint.exe",
            "open_task_manager": "taskmgr.exe",
            "open_control_panel": "control.exe",
            "open_store": "ms-windows-store:",
            "open_registry": "regedit.exe",
            "open_camera": "microsoft.windows.camera:",
            "open_calendar": "outlookcal:",
            "open_mail": "mailto:",
            "open_chrome": "chrome.exe",
            "open_edge": "msedge.exe",
            "open_add_remove_programs": "appwiz.cpl",
            "open_volume_mixer": "sndvol.exe",
            "open_mic_settings": "ms-settings:sound",
            "open_sound_cpl": "mmsys.cpl",
            "open_network_connections": "ncpa.cpl",
            "open_netconnections_cpl": "control netconnections",
            "open_time_date": "timedate.cpl",
            "open_system_properties": "sysdm.cpl",
            "open_power_options": "powercfg.cpl",
            "open_firewall_cpl": "firewall.cpl",
            "open_internet_options_cpl": "inetcpl.cpl",
            "open_display_cpl": "desk.cpl",
            "open_admin_tools_cpl": "control admintools",
            "open_schedtasks_cpl": "control schedtasks",
            "open_mouse_cpl": "main.cpl",
            "open_keyboard_cpl": "control keyboard",
            "open_fonts_cpl": "control fonts",
            "open_region_cpl": "intl.cpl",
            "open_folder_options_cpl": "control folders",
            "open_color_cpl": "control color",
            "open_desktop_cpl": "control desktop",
            "open_printers_cpl": "control printers",
            "open_user_accounts_cpl": "control userpasswords2",
            "open_bluetooth_cpl": "bthprops.cpl",
            "open_accessibility_cpl": "access.cpl",
        }
        if mode_norm in mapping:
            target = mapping[mode_norm]
            ok, out = _run_powershell(f"Start-Process '{target}'", timeout=12)
            return _json({"ok": True, "mode": mode_norm, "target": target}) if ok else self._error(out or f"{mode_norm} failed")
        if mode_norm == "close_all_apps":
            limit = 9999
            if max_kill is not None:
                try:
                    limit = max(0, int(max_kill))
                except Exception:
                    limit = 9999
            stop_cmd = "  try { Stop-Process -Id $p.Id -Force -ErrorAction Stop } catch {} " if not dry_run else ""
            ps = (
                "$items=Get-Process | Where-Object {$_.MainWindowTitle -ne '' -and $_.ProcessName -notin @('explorer','ShellExperienceHost','StartMenuExperienceHost')}; "
                "$items=$items | Sort-Object WorkingSet -Descending | Select-Object -First " + str(limit) + "; "
                "$affected=@(); "
                "foreach($p in $items){ "
                "  $affected += [pscustomobject]@{pid=$p.Id;name=$p.ProcessName;title=$p.MainWindowTitle}; "
                + stop_cmd
                + "}; "
                "@{ok=$true; mode='close_all_apps'; dry_run=" + ("$true" if dry_run else "$false") + "; max_kill=" + str(limit) + "; count=$affected.Count; affected=$affected} | ConvertTo-Json -Compress -Depth 5"
            )
            ok, out = _run_powershell(ps, timeout=25)
            return out if ok and out else self._error(out or "close all apps failed")
        if mode_norm == "open_app":
            q = (app or "").strip()
            if not q:
                return self._error("app is required")
            return self._launch_start_app(q)
        return self._error(f"unsupported app_tools mode: {mode_norm}")

    def _info_tools(self, mode: str, timezone: str = "") -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm == "timezone_get":
            ok, out = _run_powershell("Get-TimeZone | Select-Object Id,DisplayName | ConvertTo-Json -Compress", timeout=10)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "timezone get failed")
        if mode_norm == "timezone_set":
            tz = (timezone or "").strip()
            if not tz:
                return self._error("timezone is required")
            ok, out = _run_powershell(f"Set-TimeZone -Id '{tz}'", timeout=12)
            return _json({"ok": True, "mode": mode_norm, "timezone": tz}) if ok else self._error(out or "timezone set failed")
        if mode_norm == "system_language":
            ok, out = _run_powershell("Get-WinSystemLocale | Select-Object Name,DisplayName | ConvertTo-Json -Compress", timeout=10)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "system language failed")
        if mode_norm == "windows_product_key":
            ps = "(Get-CimInstance SoftwareLicensingService).OA3xOriginalProductKey | ConvertTo-Json -Compress"
            ok, out = _run_powershell(ps, timeout=10)
            if ok and out:
                try:
                    data = json.loads(out)
                except Exception:
                    data = out.strip()
                return _json({"ok": True, "mode": mode_norm, "key": data})
            return self._error(out or "product key lookup failed")
        if mode_norm == "model_info":
            ps = "Get-CimInstance Win32_ComputerSystem | Select-Object Manufacturer,Model,Name | ConvertTo-Json -Compress"
            ok, out = _run_powershell(ps, timeout=10)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "model info failed")
        if mode_norm == "windows_install_date":
            ps = (
                "Get-CimInstance Win32_OperatingSystem | "
                "Select-Object Caption,Version,InstallDate,LastBootUpTime | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=12)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "windows install date failed")
        if mode_norm == "refresh_rate":
            ps = (
                "Get-CimInstance Win32_VideoController | "
                "Select-Object Name,CurrentRefreshRate,CurrentHorizontalResolution,CurrentVerticalResolution | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=12)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "refresh rate lookup failed")
        return self._error(f"unsupported info_tools mode: {mode_norm}")

    def _dev_tools(
        self,
        mode: str,
        drive: str = "",
        path: str = "",
        editor: str = "",
        max_results: int = 20,
        target: str = "",
        text: str = "",
        execute: bool = True,
    ) -> str:
        mode_norm = (mode or "").strip().lower()
        mapping = {
            "open_cmd_admin": "Start-Process cmd.exe -Verb RunAs",
            "open_powershell_admin": "Start-Process powershell.exe -Verb RunAs",
            "open_disk_management": "Start-Process diskmgmt.msc",
            "open_device_manager": "Start-Process devmgmt.msc",
            "open_perfmon": "Start-Process perfmon.msc",
            "open_event_viewer": "Start-Process eventvwr.msc",
            "open_services": "Start-Process services.msc",
            "open_registry": "Start-Process regedit.exe -Verb RunAs",
            "open_task_scheduler": "Start-Process taskschd.msc",
            "open_computer_management": "Start-Process compmgmt.msc",
            "open_local_users_groups": "Start-Process lusrmgr.msc",
            "open_local_security_policy": "Start-Process secpol.msc",
            "open_print_management": "Start-Process printmanagement.msc",
        }
        if mode_norm in mapping:
            ok, out = _run_powershell(mapping[mode_norm], timeout=15)
            return _json({"ok": True, "mode": mode_norm}) if ok else self._error(out or f"{mode_norm} failed")
        if mode_norm in {"interpret_powershell", "nl_to_powershell", "human_to_ps"}:
            query = _normalize_query(text or target)
            if not query:
                return self._error("text or target is required")
            q = query.casefold()
            candidates: list[tuple[str, str]] = [
                (
                    "services_list",
                    "Get-Service | Select-Object -First 120 Name,Status,StartType | ConvertTo-Json -Compress",
                ),
                (
                    "startup_list",
                    "Get-CimInstance Win32_StartupCommand | Select-Object Name,Command,Location,User | ConvertTo-Json -Compress",
                ),
                (
                    "network_status",
                    "Get-NetIPConfiguration | Select-Object InterfaceAlias,IPv4Address,IPv6Address,IPv4DefaultGateway | ConvertTo-Json -Compress",
                ),
                (
                    "top_cpu",
                    "Get-Process | Sort-Object CPU -Descending | Select-Object -First 10 Id,ProcessName,CPU,WS | ConvertTo-Json -Compress",
                ),
                (
                    "top_ram",
                    "Get-Process | Sort-Object WS -Descending | Select-Object -First 10 Id,ProcessName,CPU,WS | ConvertTo-Json -Compress",
                ),
                (
                    "installed_updates",
                    "Get-HotFix | Sort-Object InstalledOn -Descending | Select-Object -First 60 HotFixID,InstalledOn,Description | ConvertTo-Json -Compress",
                ),
                (
                    "event_errors",
                    "Get-WinEvent -FilterHashtable @{LogName='System'; Level=2} -MaxEvents 30 | Select-Object TimeCreated,Id,ProviderName,Message | ConvertTo-Json -Compress",
                ),
                (
                    "open_ports",
                    "Get-NetTCPConnection -State Listen -ErrorAction SilentlyContinue | Select-Object LocalAddress,LocalPort,OwningProcess | ConvertTo-Json -Compress",
                ),
                (
                    "disk_usage",
                    "Get-PSDrive -PSProvider FileSystem | Select-Object Name,Used,Free | ConvertTo-Json -Compress",
                ),
                (
                    "battery_status",
                    "$b=(Get-CimInstance Win32_Battery -ErrorAction SilentlyContinue | Select-Object -First 1 EstimatedChargeRemaining,BatteryStatus,EstimatedRunTime); $b | ConvertTo-Json -Compress",
                ),
            ]
            intent = "generic"
            command = ""
            if any(tok in q for tok in ("startup", "بدء التشغيل", "ستارتب")):
                intent, command = candidates[1]
            elif any(tok in q for tok in ("service", "خدمة", "الخدمات")):
                intent, command = candidates[0]
            elif any(tok in q for tok in ("ip", "network", "شبكة", "انترنت", "إنترنت")):
                intent, command = candidates[2]
            elif any(tok in q for tok in ("top cpu", "اكثر معالج", "أكثر معالج", "المعالج")):
                intent, command = candidates[3]
            elif any(tok in q for tok in ("top ram", "اكثر رام", "أكثر رام", "الرام", "memory")):
                intent, command = candidates[4]
            elif any(tok in q for tok in ("update", "hotfix", "تحديث")):
                intent, command = candidates[5]
            elif any(tok in q for tok in ("error", "event", "سجل", "اخطاء", "أخطاء")):
                intent, command = candidates[6]
            elif any(tok in q for tok in ("open port", "listen", "منفذ", "بورت")):
                intent, command = candidates[7]
            elif any(tok in q for tok in ("disk", "storage", "قرص", "مساحة")):
                intent, command = candidates[8]
            elif any(tok in q for tok in ("battery", "بطارية")):
                intent, command = candidates[9]
            else:
                return self._error("could not map request to a safe PowerShell intent")

            response: dict[str, Any] = {
                "ok": True,
                "mode": mode_norm,
                "query": query,
                "intent": intent,
                "powershell": command,
                "executed": bool(execute),
            }
            if not execute:
                return _json(response)

            ok, out = _run_powershell(command, timeout=30)
            if not ok:
                return self._error(out or "interpreted command execution failed")
            parsed: Any = out
            try:
                parsed = json.loads(out)
            except Exception:
                parsed = out[:3000]
            response["result"] = parsed
            return _json(response)
        if mode_norm == "sfc_scan":
            ok, out = _run_powershell("sfc /scannow", timeout=600)
            return _json({"ok": True, "mode": mode_norm, "output": out[:2200]}) if ok else self._error(out or "sfc scan failed")
        if mode_norm == "chkdsk":
            drv = (drive or "C:").strip()
            return self._disk_tools("chkdsk_scan", drive=drv)
        if mode_norm in {"env_vars", "environment_variables"}:
            ok, out = _run_powershell("Get-ChildItem Env: | Sort-Object Name | ConvertTo-Json -Compress", timeout=12)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "environment variables query failed")
        if mode_norm in {"runtime_versions", "python_node_versions"}:
            py_ver = ""
            node_ver = ""
            try:
                py = subprocess.run(["python", "--version"], capture_output=True, text=True, timeout=5)
                py_ver = (py.stdout or py.stderr or "").strip()
            except Exception:
                py_ver = ""
            try:
                nd = subprocess.run(["node", "--version"], capture_output=True, text=True, timeout=5)
                node_ver = (nd.stdout or nd.stderr or "").strip()
            except Exception:
                node_ver = ""
            return _json({"ok": True, "mode": mode_norm, "python": py_ver or None, "node": node_ver or None})
        if mode_norm in {"git_last_log", "git_recent"}:
            base = Path(path).expanduser() if path else Path.cwd()
            if not base.exists():
                return self._error(f"path not found: {base}")
            try:
                cmd = ["git", "-C", str(base), "log", "--oneline", f"-n{_clamp(max_results, 1, 50)}"]
                proc = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
                if proc.returncode != 0:
                    return self._error((proc.stderr or proc.stdout or "git log failed").strip())
                lines = [ln.strip() for ln in (proc.stdout or "").splitlines() if ln.strip()]
                return _json({"ok": True, "mode": mode_norm, "path": str(base), "lines": lines})
            except Exception as exc:
                return self._error(f"git log failed: {exc}")
        if mode_norm in {"open_editor", "open_code_here"}:
            base = Path(path).expanduser() if path else Path.cwd()
            if not base.exists():
                return self._error(f"path not found: {base}")
            pref = (editor or os.getenv("MUDABBIR_EDITOR") or os.getenv("EDITOR") or "code").strip()
            choices = [pref, "code", "notepad++", "notepad"]
            last_err = ""
            for choice in choices:
                exe = choice.strip()
                if not exe:
                    continue
                try:
                    subprocess.Popen([exe, str(base)], shell=False)
                    return _json({"ok": True, "mode": mode_norm, "editor": exe, "path": str(base)})
                except Exception as exc:
                    last_err = str(exc)
                    continue
            return self._error(last_err or "failed to open editor")
        if mode_norm in {"event_errors", "bsod_summary"}:
            ps = (
                "Get-WinEvent -FilterHashtable @{LogName='System'; Level=2} -MaxEvents 40 | "
                "Select-Object TimeCreated,Id,ProviderName,LevelDisplayName,Message | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=25)
            if ok and out:
                try:
                    rows = json.loads(out)
                    if isinstance(rows, dict):
                        rows = [rows]
                    normalized = []
                    for r in rows[:40]:
                        if not isinstance(r, dict):
                            continue
                        normalized.append(
                            {
                                "time": str(r.get("TimeCreated", "") or ""),
                                "id": r.get("Id"),
                                "provider": str(r.get("ProviderName", "") or ""),
                                "level": str(r.get("LevelDisplayName", "") or ""),
                                "message": str(r.get("Message", "") or "")[:380],
                            }
                        )
                    return _json({"ok": True, "mode": mode_norm, "count": len(normalized), "items": normalized})
                except Exception:
                    return _json({"ok": True, "mode": mode_norm, "raw": out[:3000]})
            return self._error(out or "event log query failed")
        if mode_norm in {"analyze_bsod", "bsod_analyze", "debug_system_errors"}:
            ps = (
                "Get-WinEvent -FilterHashtable @{LogName='System'; ID=@(41,1001,6008)} -MaxEvents 80 | "
                "Select-Object TimeCreated,Id,ProviderName,Message | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=30)
            if not ok or not out:
                return self._error(out or "bsod analysis query failed")
            try:
                rows = json.loads(out)
                if isinstance(rows, dict):
                    rows = [rows]
                if not isinstance(rows, list):
                    rows = []

                normalized: list[dict[str, Any]] = []
                id_counts: dict[str, int] = {}
                bugcheck_codes: list[str] = []
                provider_counts: dict[str, int] = {}
                hints: set[str] = set()

                for row in rows:
                    if not isinstance(row, dict):
                        continue
                    msg = str(row.get("Message", "") or "")
                    event_id = int(row.get("Id") or 0)
                    provider = str(row.get("ProviderName", "") or "")
                    id_counts[str(event_id)] = int(id_counts.get(str(event_id), 0)) + 1
                    provider_counts[provider or "unknown"] = int(provider_counts.get(provider or "unknown", 0)) + 1
                    for match in re.findall(r"0x[0-9A-Fa-f]{4,16}", msg):
                        if match not in bugcheck_codes:
                            bugcheck_codes.append(match)
                    low = msg.casefold()
                    if event_id in {41, 6008}:
                        hints.add("Check power settings, PSU health, and unexpected shutdown causes.")
                    if event_id == 1001:
                        hints.add("Update GPU/chipset/storage drivers and review crash dumps in C:\\Windows\\Minidump.")
                    if "memory" in low or "ram" in low:
                        hints.add("Run Windows Memory Diagnostic (mdsched.exe) to verify RAM stability.")
                    if "disk" in low or "ntfs" in low or "storage" in low:
                        hints.add("Run chkdsk and inspect disk SMART status for storage-related faults.")
                    if "driver" in low or "whea" in low:
                        hints.add("Roll back or update recently changed drivers; verify BIOS/UEFI firmware updates.")
                    normalized.append(
                        {
                            "time": str(row.get("TimeCreated", "") or ""),
                            "id": event_id,
                            "provider": provider,
                            "message": msg[:520],
                        }
                    )

                normalized.sort(key=lambda item: str(item.get("time", "")), reverse=True)
                ordered_hints = sorted(hints)
                if not ordered_hints:
                    ordered_hints = [
                        "Collect full dump/minidump and correlate with recent driver or hardware changes.",
                        "Run `sfc /scannow` and `DISM /Online /Cleanup-Image /RestoreHealth`.",
                    ]
                return _json(
                    {
                        "ok": True,
                        "mode": mode_norm,
                        "event_count": len(normalized),
                        "event_id_counts": id_counts,
                        "provider_counts": provider_counts,
                        "bugcheck_codes": bugcheck_codes[:12],
                        "suggestions": ordered_hints[:8],
                        "recent_events": normalized[:25],
                    }
                )
            except Exception as exc:
                return self._error(f"bsod analysis parse failed: {exc}")
        return self._error(f"unsupported dev_tools mode: {mode_norm}")

    def _shell_tools(self, mode: str) -> str:
        mode_norm = (mode or "").strip().lower()
        mapping = {
            "quick_settings": ("win", "a"),
            "notifications": ("win", "n"),
            "search": ("win", "s"),
            "run": ("win", "r"),
            "file_explorer": ("win", "e"),
            "quick_link_menu": ("win", "x"),
            "task_view": ("win", "tab"),
            "new_virtual_desktop": ("win", "ctrl", "d"),
            "next_virtual_desktop": ("win", "ctrl", "right"),
            "prev_virtual_desktop": ("win", "ctrl", "left"),
            "close_virtual_desktop": ("win", "ctrl", "f4"),
            "emoji_panel": ("win", "."),
            "start_menu": ("win",),
            "refresh": ("f5",),
            "magnifier_open": ("win", "+"),
            "magnifier_zoom_out": ("win", "-"),
            "magnifier_close": ("win", "esc"),
            "narrator_toggle": ("win", "ctrl", "enter"),
            "clipboard_history": ("win", "v"),
        }
        if mode_norm in {"list_shortcuts", "shortcuts", "hotkeys"}:
            return _json({"ok": True, "mode": mode_norm, "shortcuts": {k: list(v) for k, v in mapping.items()}})
        hotkey = mapping.get(mode_norm)
        if not hotkey:
            return self._error(f"unsupported shell_tools mode: {mode_norm}")
        try:
            import pyautogui

            pyautogui.hotkey(*hotkey)
            return _json({"ok": True, "mode": mode_norm, "keys": list(hotkey)})
        except Exception as exc:
            return self._error(f"shell tools failed: {exc}")

    def _office_tools(self, mode: str, path: str = "", target: str = "") -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm == "open_word_new":
            ok, out = _run_powershell("Start-Process winword.exe", timeout=12)
            return _json({"ok": True, "mode": mode_norm}) if ok else self._error(out or "open word failed")
        if mode_norm == "silent_print":
            src = Path(path).expanduser()
            if not src.exists():
                return self._error(f"path not found: {src}")
            ok, out = _run_powershell(f'Start-Process -FilePath "{src}" -Verb Print', timeout=20)
            return _json({"ok": True, "mode": mode_norm, "path": str(src)}) if ok else self._error(out or "print failed")
        if mode_norm == "docx_to_pdf":
            src = Path(path).expanduser()
            dst = Path(target).expanduser() if target else src.with_suffix(".pdf")
            if not src.exists():
                return self._error(f"path not found: {src}")
            esc_src = str(src).replace("'", "''")
            esc_dst = str(dst).replace("'", "''")
            ps = (
                "$word = New-Object -ComObject Word.Application; "
                "$word.Visible = $false; "
                f"$doc = $word.Documents.Open('{esc_src}'); "
                f"$doc.SaveAs([ref] '{esc_dst}', [ref] 17); "
                "$doc.Close(); $word.Quit(); "
                f"@{{ok=$true; mode='docx_to_pdf'; source='{str(src)}'; target='{str(dst)}'}} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=90)
            return out if ok and out else self._error(out or "docx to pdf failed")
        return self._error(f"unsupported office_tools mode: {mode_norm}")

    def _remote_tools(self, mode: str, host: str = "") -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm == "rdp_open":
            ok, out = _run_powershell("Start-Process mstsc.exe", timeout=10)
            return _json({"ok": True, "mode": mode_norm}) if ok else self._error(out or "open rdp failed")
        if mode_norm == "vpn_connect":
            name = (host or "").strip()
            if not name:
                return self._error("host is required (VPN connection name)")
            ok, out = _run_powershell(f"rasdial \"{name}\"", timeout=20)
            return _json({"ok": True, "mode": mode_norm, "name": name, "output": out[:1200]}) if ok else self._error(out or "vpn connect failed")
        if mode_norm == "vpn_disconnect":
            name = (host or "").strip()
            cmd = f"rasdial \"{name}\" /disconnect" if name else "rasdial /disconnect"
            ok, out = _run_powershell(cmd, timeout=20)
            return _json({"ok": True, "mode": mode_norm, "name": name or "all", "output": out[:1200]}) if ok else self._error(out or "vpn disconnect failed")
        return self._error(f"unsupported remote_tools mode: {mode_norm}")

    def _search_tools(
        self, mode: str, folder: str = "", pattern: str = "", ext: str = "", size_mb: Any = None
    ) -> str:
        mode_norm = (mode or "").strip().lower()
        base = Path(folder).expanduser() if folder else Path.home()
        if not base.exists():
            return self._error(f"folder not found: {base}")
        if mode_norm == "search_text":
            token = (pattern or "").strip()
            if not token:
                return self._error("pattern is required")
            hits: list[dict[str, Any]] = []
            for p in base.rglob("*"):
                if len(hits) >= 120:
                    break
                if not p.is_file():
                    continue
                if p.suffix.lower() not in {".txt", ".md", ".py", ".js", ".json", ".log", ".csv"}:
                    continue
                try:
                    text = p.read_text(encoding="utf-8", errors="ignore")
                except Exception:
                    continue
                if token.casefold() in text.casefold():
                    hits.append({"path": str(p), "size": p.stat().st_size})
            return _json({"ok": True, "mode": mode_norm, "count": len(hits), "items": hits})
        if mode_norm == "files_larger_than":
            try:
                threshold = float(size_mb if size_mb is not None else 100.0) * 1024 * 1024
            except Exception:
                threshold = 100.0 * 1024 * 1024
            items: list[dict[str, Any]] = []
            for p in base.rglob("*"):
                if len(items) >= 200:
                    break
                if p.is_file():
                    try:
                        sz = p.stat().st_size
                    except Exception:
                        continue
                    if sz >= threshold:
                        items.append({"path": str(p), "size_mb": round(sz / (1024 * 1024), 2)})
            return _json({"ok": True, "mode": mode_norm, "count": len(items), "items": items})
        if mode_norm == "modified_today":
            import datetime as _dt

            today = _dt.date.today()
            items: list[str] = []
            for p in base.rglob("*"):
                if len(items) >= 200:
                    break
                if p.is_file():
                    try:
                        d = _dt.datetime.fromtimestamp(p.stat().st_mtime).date()
                    except Exception:
                        continue
                    if d == today:
                        items.append(str(p))
            return _json({"ok": True, "mode": mode_norm, "count": len(items), "items": items})
        if mode_norm == "find_images":
            exts = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".gif"}
            items = [str(p) for p in base.rglob("*") if p.is_file() and p.suffix.lower() in exts][:300]
            return _json({"ok": True, "mode": mode_norm, "count": len(items), "items": items})
        if mode_norm == "find_videos":
            exts = {".mp4", ".mkv", ".mov", ".avi", ".wmv", ".webm"}
            items = [str(p) for p in base.rglob("*") if p.is_file() and p.suffix.lower() in exts][:300]
            return _json({"ok": True, "mode": mode_norm, "count": len(items), "items": items})
        if mode_norm == "count_files":
            count = 0
            for p in base.rglob("*"):
                if p.is_file():
                    count += 1
            return _json({"ok": True, "mode": mode_norm, "count": count, "folder": str(base)})
        if mode_norm in {
            "search_open_windows",
            "search_windows",
            "search_open_windows_content",
            "search_windows_content",
            "search_window_text",
        }:
            query = (pattern or "").strip().casefold()
            if not query:
                return self._error("pattern is required")
            include_content = "content" in mode_norm or "text" in mode_norm
            wins = _serialize_windows(include_untitled=False, limit=400)
            hits: list[dict[str, Any]] = []
            for w in wins:
                title = str(w.get("title") or "")
                if query in title.casefold():
                    item = dict(w)
                    item["match_source"] = "title"
                    hits.append(item)
            if not include_content:
                return _json({"ok": True, "mode": mode_norm, "count": len(hits), "items": hits[:200]})
            try:
                import psutil
                from pywinauto import Desktop

                desktop = Desktop(backend="uia")
                scanned = 0
                for wrapper in desktop.windows():
                    if scanned >= 40:
                        break
                    scanned += 1
                    try:
                        title = str(wrapper.window_text() or "").strip()
                        if not title:
                            continue
                        info = wrapper.element_info
                        pid = int(getattr(info, "process_id", 0) or 0)
                        proc_name = ""
                        if pid > 0:
                            try:
                                proc_name = str(psutil.Process(pid).name() or "")
                            except Exception:
                                proc_name = ""
                        controls = self._enumerate_window_controls(wrapper, max_items=150)
                        control_hits: list[dict[str, Any]] = []
                        for ctrl in controls:
                            name_val = str(ctrl.get("name") or "").strip()
                            if not name_val:
                                continue
                            if query not in name_val.casefold():
                                continue
                            control_hits.append(
                                {
                                    "name": name_val,
                                    "control_type": str(ctrl.get("control_type") or ""),
                                    "auto_id": str(ctrl.get("auto_id") or ""),
                                    "class_name": str(ctrl.get("class_name") or ""),
                                }
                            )
                            if len(control_hits) >= 8:
                                break
                        if not control_hits:
                            continue
                        hits.append(
                            {
                                "title": title,
                                "pid": pid,
                                "process_name": proc_name,
                                "match_source": "window_content",
                                "control_match_count": len(control_hits),
                                "control_matches": control_hits,
                            }
                        )
                        if len(hits) >= 200:
                            break
                    except Exception:
                        continue
            except Exception as exc:
                return self._error(f"search_open_windows_content failed: {exc}")
            return _json({"ok": True, "mode": mode_norm, "count": len(hits), "items": hits[:200]})
        return self._error(f"unsupported search_tools mode: {mode_norm}")

    def _performance_tools(self, mode: str, threshold: Any = None) -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm in {"top_cpu", "top_ram"}:
            raw = self._process_tools(mode_norm)
            if isinstance(raw, str) and raw.lower().startswith("error:"):
                return raw
            try:
                data = json.loads(raw)
                if isinstance(data, dict) and isinstance(data.get("items"), list):
                    data["items"] = data.get("items", [])[:5]
                    data["count"] = len(data["items"])
                    return _json(data)
            except Exception:
                pass
            return raw
        if mode_norm == "top_disk":
            try:
                import psutil

                items = []
                for p in psutil.process_iter(["pid", "name", "io_counters"]):
                    try:
                        io = p.info.get("io_counters")
                        if not io:
                            continue
                        rb = float(getattr(io, "read_bytes", 0.0) or 0.0)
                        wb = float(getattr(io, "write_bytes", 0.0) or 0.0)
                        total_mb = round((rb + wb) / (1024 * 1024), 2)
                        items.append(
                            {
                                "pid": int(p.info.get("pid") or 0),
                                "name": str(p.info.get("name") or ""),
                                "read_mb": round(rb / (1024 * 1024), 2),
                                "write_mb": round(wb / (1024 * 1024), 2),
                                "total_mb": total_mb,
                            }
                        )
                    except Exception:
                        continue
                items.sort(key=lambda x: float(x.get("total_mb", 0.0)), reverse=True)
                return _json({"ok": True, "mode": mode_norm, "items": items[:5]})
            except Exception as exc:
                return self._error(f"top_disk failed: {exc}")
        if mode_norm == "total_ram_percent":
            try:
                import psutil

                return _json({"ok": True, "mode": mode_norm, "percent": float(psutil.virtual_memory().percent)})
            except Exception as exc:
                return self._error(f"total ram failed: {exc}")
        if mode_norm == "total_cpu_percent":
            try:
                import psutil

                return _json({"ok": True, "mode": mode_norm, "percent": float(psutil.cpu_percent(interval=0.5))})
            except Exception as exc:
                return self._error(f"total cpu failed: {exc}")
        if mode_norm == "cpu_clock":
            try:
                import psutil

                freq = psutil.cpu_freq()
                return _json(
                    {
                        "ok": True,
                        "mode": mode_norm,
                        "current_mhz": float(getattr(freq, "current", 0.0) or 0.0),
                        "min_mhz": float(getattr(freq, "min", 0.0) or 0.0),
                        "max_mhz": float(getattr(freq, "max", 0.0) or 0.0),
                    }
                )
            except Exception as exc:
                return self._error(f"cpu clock failed: {exc}")
        if mode_norm == "available_ram":
            try:
                import psutil

                vm = psutil.virtual_memory()
                return _json({"ok": True, "mode": mode_norm, "available_mb": round(float(vm.available) / (1024 * 1024), 2)})
            except Exception as exc:
                return self._error(f"available ram failed: {exc}")
        if mode_norm == "pagefile_used":
            try:
                import psutil

                sm = psutil.swap_memory()
                return _json(
                    {
                        "ok": True,
                        "mode": mode_norm,
                        "used_mb": round(float(sm.used) / (1024 * 1024), 2),
                        "total_mb": round(float(sm.total) / (1024 * 1024), 2),
                        "percent": float(sm.percent),
                    }
                )
            except Exception as exc:
                return self._error(f"pagefile query failed: {exc}")
        if mode_norm == "disk_io_rate":
            ps = (
                "$r=Get-Counter '\\PhysicalDisk(_Total)\\Disk Read Bytes/sec','\\PhysicalDisk(_Total)\\Disk Write Bytes/sec'; "
                "$vals=$r.CounterSamples | Select-Object Path,CookedValue; "
                "$vals | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=12)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "disk io rate failed")
        if mode_norm == "gpu_util":
            ps = (
                "Get-Counter '\\GPU Engine(*)\\Utilization Percentage' -ErrorAction SilentlyContinue | "
                "Select-Object -ExpandProperty CounterSamples | "
                "Sort-Object CookedValue -Descending | Select-Object -First 20 InstanceName,CookedValue | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=12)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "items": json.loads(out)})
            return self._error(out or "gpu utilization query failed")
        if mode_norm in {"top_gpu_processes", "gpu_processes"}:
            ps = (
                "$samples=Get-Counter '\\GPU Engine(*)\\Utilization Percentage' -ErrorAction SilentlyContinue | "
                "Select-Object -ExpandProperty CounterSamples; "
                "$agg=@{}; "
                "foreach($s in $samples){ "
                "  $inst=[string]$s.InstanceName; "
                "  $m=[regex]::Match($inst,'pid_([0-9]+)'); "
                "  if(-not $m.Success){ continue }; "
                "  $pid=[int]$m.Groups[1].Value; "
                "  if($pid -le 0){ continue }; "
                "  if(-not $agg.ContainsKey($pid)){ $agg[$pid]=0.0 }; "
                "  $agg[$pid] = [double]$agg[$pid] + [double]$s.CookedValue; "
                "}; "
                "$rows=@(); "
                "foreach($k in $agg.Keys){ "
                "  $p=Get-Process -Id $k -ErrorAction SilentlyContinue; "
                "  if($p){ $rows += [pscustomobject]@{pid=$k;name=$p.ProcessName;gpu_util=[math]::Round([double]$agg[$k],2)} } "
                "}; "
                "$rows | Sort-Object gpu_util -Descending | Select-Object -First 10 | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=18)
            if ok and out:
                parsed = json.loads(out)
                items = [parsed] if isinstance(parsed, dict) else (parsed if isinstance(parsed, list) else [])
                return _json({"ok": True, "mode": mode_norm, "count": len(items), "items": items[:10]})
            return self._error(out or "top gpu processes query failed")
        if mode_norm == "gpu_temp":
            ps = (
                "Get-CimInstance MSAcpi_ThermalZoneTemperature -Namespace root/wmi | "
                "Select-Object InstanceName,CurrentTemperature | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=12)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "items": json.loads(out)})
            return self._error(out or "gpu temp unavailable")
        if mode_norm == "empty_ram":
            return self._maintenance_tools("empty_ram")
        if mode_norm == "kill_high_cpu":
            return self._process_tools("kill_high_cpu", threshold=threshold)
        if mode_norm == "cpu_popup":
            try:
                import psutil

                cpu = psutil.cpu_percent(interval=0.8)
                return self._automation_tools("popup", text=f"CPU usage: {cpu}%")
            except Exception as exc:
                return self._error(f"cpu_popup failed: {exc}")
        return self._error(f"unsupported performance_tools mode: {mode_norm}")

    def _media_tools(self, mode: str, url: str = "") -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm == "stop_all_media":
            try:
                import pyautogui

                pyautogui.FAILSAFE = False
                pyautogui.press("playpause")
                return _json({"ok": True, "mode": mode_norm})
            except Exception as exc:
                return self._error(f"stop_all_media failed: {exc}")
        if mode_norm == "youtube_open":
            target = (url or "").strip()
            if not target:
                return self._error("url is required")
            return self._web_tools("open_url", url=target)
        if mode_norm == "media_next":
            return self._media_control("next")
        if mode_norm == "media_prev":
            return self._media_control("previous")
        if mode_norm == "play_pause":
            return self._media_control("play_pause")
        return self._error(f"unsupported media_tools mode: {mode_norm}")

    def _browser_deep_tools(self, mode: str, urls: Any = None) -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm == "multi_open":
            values = urls if isinstance(urls, list) else []
            opened = 0
            for u in values[:20]:
                target = str(u or "").strip()
                if not target:
                    continue
                subprocess.Popen(["cmd", "/c", "start", "", target], shell=False)
                opened += 1
            return _json({"ok": True, "mode": mode_norm, "opened": opened})
        if mode_norm == "clear_chrome_cache":
            cache = Path.home() / "AppData/Local/Google/Chrome/User Data/Default/Cache"
            if cache.exists():
                import shutil

                shutil.rmtree(cache, ignore_errors=True)
            return _json({"ok": True, "mode": mode_norm, "path": str(cache)})
        if mode_norm == "clear_edge_cache":
            cache = Path.home() / "AppData/Local/Microsoft/Edge/User Data/Default/Cache"
            if cache.exists():
                import shutil

                shutil.rmtree(cache, ignore_errors=True)
            return _json({"ok": True, "mode": mode_norm, "path": str(cache)})
        return self._error(f"unsupported browser_deep_tools mode: {mode_norm}")

    def _maintenance_tools(self, mode: str) -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm == "empty_ram":
            ps = (
                "$sig='[DllImport(\"psapi.dll\")]public static extern int EmptyWorkingSet(IntPtr hwProc);'; "
                "Add-Type -Name NativeMethods -Namespace Win32 -MemberDefinition $sig; "
                "Get-Process | ForEach-Object { try {[Win32.NativeMethods]::EmptyWorkingSet($_.Handle) | Out-Null} catch {} }; "
                "@{ok=$true; mode='empty_ram'} | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=40)
            return out if ok and out else self._error(out or "empty ram failed")
        if mode_norm == "winsxs_cleanup":
            return self._update_tools("winsxs_cleanup")
        if mode_norm == "temp_clean":
            return self._disk_tools("temp_files_clean")
        return self._error(f"unsupported maintenance_tools mode: {mode_norm}")

    def _driver_tools(self, mode: str) -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm == "drivers_list":
            ok, out = _run_powershell(
                "Get-CimInstance Win32_PnPSignedDriver | Select-Object -First 200 DeviceName,DriverVersion,Manufacturer,DriverDate | ConvertTo-Json -Compress",
                timeout=25,
            )
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "drivers list failed")
        if mode_norm == "drivers_backup":
            out_dir = Path.home() / f"drivers_backup_{_timestamp_id()}"
            ok, out = _run_powershell(f"pnputil /export-driver * \"{out_dir}\"", timeout=180)
            return _json({"ok": True, "mode": mode_norm, "path": str(out_dir), "output": out[:1200]}) if ok else self._error(out or "drivers backup failed")
        if mode_norm == "updates_pending":
            return self._update_tools("list_updates")
        if mode_norm == "drivers_issues":
            ps = (
                "Get-PnpDevice -PresentOnly | Where-Object {$_.Status -ne 'OK'} | "
                "Select-Object Class,FriendlyName,Status,Problem,InstanceId | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=25)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "drivers issues query failed")
        return self._error(f"unsupported driver_tools mode: {mode_norm}")

    def _power_user_tools(self, mode: str) -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm == "airplane_on":
            return self._shell_tools("quick_settings")
        if mode_norm == "airplane_off":
            return self._shell_tools("quick_settings")
        if mode_norm == "god_mode":
            desktop = Path.home() / "Desktop"
            gm = desktop / "GodMode.{ED7BA470-8E54-465E-825C-99712043E01C}"
            gm.mkdir(exist_ok=True)
            return _json({"ok": True, "mode": mode_norm, "path": str(gm)})
        if mode_norm == "invert_colors":
            try:
                import pyautogui

                pyautogui.hotkey("ctrl", "win", "c")
                return _json({"ok": True, "mode": mode_norm})
            except Exception as exc:
                return self._error(f"invert colors failed: {exc}")
        return self._error(f"unsupported power_user_tools mode: {mode_norm}")

    def _screenshot_tools(
        self,
        mode: str,
        x: Any = None,
        y: Any = None,
        width: Any = None,
        height: Any = None,
        path: str = "",
    ) -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm == "full":
            if not path:
                return self._screen_snapshot()
            try:
                from PIL import ImageGrab

                out_path = Path(path).expanduser()
                out_path.parent.mkdir(parents=True, exist_ok=True)
                ImageGrab.grab().save(out_path)
                return _json({"ok": True, "mode": mode_norm, "path": str(out_path)})
            except Exception as exc:
                return self._error(f"full screenshot failed: {exc}")
        if mode_norm == "snipping_tool":
            ok, out = _run_powershell("Start-Process snippingtool", timeout=8)
            return _json({"ok": True, "mode": mode_norm}) if ok else self._error(out or "snipping tool failed")
        if mode_norm == "window_active":
            try:
                from PIL import ImageGrab
                import pygetwindow as gw

                win = gw.getActiveWindow()
                if win is None:
                    return self._error("no active window")
                bbox = (int(win.left), int(win.top), int(win.left + win.width), int(win.top + win.height))
                img = ImageGrab.grab(bbox=bbox)
                out_path = Path(path).expanduser() if path else (get_media_dir() / f"window_{_timestamp_id()}.png")
                out_path.parent.mkdir(parents=True, exist_ok=True)
                img.save(out_path)
                return _json({"ok": True, "mode": mode_norm, "path": str(out_path)})
            except Exception as exc:
                return self._error(f"window screenshot failed: {exc}")
        if mode_norm == "region":
            try:
                from PIL import ImageGrab

                rx = int(x if x is not None else 0)
                ry = int(y if y is not None else 0)
                rw = int(width if width is not None else 300)
                rh = int(height if height is not None else 200)
                rw = max(10, rw)
                rh = max(10, rh)
                img = ImageGrab.grab(bbox=(rx, ry, rx + rw, ry + rh))
                out_path = Path(path).expanduser() if path else (get_media_dir() / f"region_{_timestamp_id()}.png")
                out_path.parent.mkdir(parents=True, exist_ok=True)
                img.save(out_path)
                return _json({"ok": True, "mode": mode_norm, "path": str(out_path), "x": rx, "y": ry, "width": rw, "height": rh})
            except Exception as exc:
                return self._error(f"region screenshot failed: {exc}")
        return self._error(f"unsupported screenshot_tools mode: {mode_norm}")

    def _text_tools(
        self,
        mode: str,
        path: str = "",
        content: str = "",
        folder: str = "",
        pattern: str = "",
        replace_with: str = "",
    ) -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm == "text_to_file":
            out_path = Path(path).expanduser() if path else (Path.home() / f"mudabbir_{_timestamp_id()}.txt")
            out_path.parent.mkdir(parents=True, exist_ok=True)
            out_path.write_text(content or "", encoding="utf-8")
            return _json({"ok": True, "mode": mode_norm, "path": str(out_path), "chars": len(content or "")})
        if mode_norm == "clipboard_to_file":
            out_path = Path(path).expanduser() if path else (Path.home() / f"clipboard_{_timestamp_id()}.txt")
            ok, out = _run_powershell("Get-Clipboard", timeout=8)
            if not ok:
                return self._error(out or "failed to read clipboard")
            out_path.parent.mkdir(parents=True, exist_ok=True)
            out_path.write_text(out or "", encoding="utf-8")
            return _json({"ok": True, "mode": mode_norm, "path": str(out_path), "chars": len(out or "")})
        if mode_norm == "word_count":
            src = Path(path).expanduser()
            if not src.exists():
                return self._error(f"path not found: {src}")
            text = src.read_text(encoding="utf-8", errors="ignore")
            words = re.findall(r"\S+", text)
            return _json({"ok": True, "mode": mode_norm, "path": str(src), "words": len(words), "chars": len(text)})
        if mode_norm == "create_batch":
            out_path = Path(path).expanduser() if path else (Path.home() / f"mudabbir_{_timestamp_id()}.bat")
            body = content.strip()
            if not body:
                return self._error("content is required for create_batch")
            if not body.lower().startswith("@echo"):
                body = "@echo off\n" + body
            out_path.parent.mkdir(parents=True, exist_ok=True)
            out_path.write_text(body + ("\n" if not body.endswith("\n") else ""), encoding="utf-8")
            return _json({"ok": True, "mode": mode_norm, "path": str(out_path)})
        if mode_norm == "search_replace_files":
            root = Path(folder).expanduser() if folder else Path.home()
            if not root.exists():
                return self._error(f"folder not found: {root}")
            if not pattern:
                return self._error("pattern is required")
            changed = 0
            touched: list[str] = []
            for p in root.rglob("*"):
                if len(touched) >= 200:
                    break
                if not p.is_file():
                    continue
                if p.suffix.lower() not in {".txt", ".md", ".py", ".js", ".json", ".log", ".csv"}:
                    continue
                try:
                    raw = p.read_text(encoding="utf-8", errors="ignore")
                except Exception:
                    continue
                if pattern in raw:
                    new = raw.replace(pattern, replace_with)
                    if new != raw:
                        p.write_text(new, encoding="utf-8")
                        changed += 1
                        touched.append(str(p))
            return _json({"ok": True, "mode": mode_norm, "changed": changed, "files": touched})
        return self._error(f"unsupported text_tools mode: {mode_norm}")

    def _api_tools(self, mode: str, target: str = "", city: str = "", text: str = "") -> str:
        mode_norm = (mode or "").strip().lower()
        if mode_norm == "currency":
            base = (target or "USD").strip().upper()
            ps = f"(Invoke-RestMethod -Uri 'https://open.er-api.com/v6/latest/{base}' -TimeoutSec 10) | ConvertTo-Json -Compress"
            ok, out = _run_powershell(ps, timeout=20)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "base": base, "data": json.loads(out)})
            return self._error(out or "currency api failed")
        if mode_norm == "weather_city":
            c = (city or target or "amman").strip()
            return self._web_tools("weather", city=c)
        if mode_norm == "translate_quick":
            payload = (text or target or "").strip()
            if not payload:
                return self._error("text is required")
            ps = (
                "$body=@{q='"
                + payload.replace("'", "''")
                + "';source='auto';target='en';format='text'} | ConvertTo-Json; "
                "(Invoke-RestMethod -Method Post -Uri 'https://libretranslate.de/translate' -Body $body -ContentType 'application/json' -TimeoutSec 15) | ConvertTo-Json -Compress"
            )
            ok, out = _run_powershell(ps, timeout=25)
            if ok and out:
                return _json({"ok": True, "mode": mode_norm, "data": json.loads(out)})
            return self._error(out or "translate api failed")
        return self._error(f"unsupported api_tools mode: {mode_norm}")

    def _vision_tools(
        self,
        mode: str,
        path: str = "",
        x: Any = None,
        y: Any = None,
        width: Any = None,
        height: Any = None,
    ) -> str:
        mode_norm = (mode or "").strip().lower()

        def _ocr_image(img_path: Path) -> tuple[bool, str, str]:
            try:
                import pytesseract
                from PIL import Image
            except Exception as exc:
                return False, "", f"ocr dependency missing: {exc}"
            try:
                text_out = str(pytesseract.image_to_string(Image.open(img_path)) or "").strip()
                return True, text_out, "pytesseract"
            except Exception as exc:
                return False, "", f"ocr failed: {exc}"

        def _capture(mode_name: str) -> tuple[Path | None, str | None]:
            try:
                from PIL import ImageGrab
            except Exception as exc:
                return None, f"PIL/ImageGrab unavailable: {exc}"

            try:
                if mode_name in {"ocr_image", "extract_text_image"}:
                    src = Path(path).expanduser()
                    if not src.exists():
                        return None, f"path not found: {src}"
                    return src, None

                if mode_name in {"ocr_active_window", "extract_text_active_window"}:
                    import pygetwindow as gw

                    win = gw.getActiveWindow()
                    if win is None:
                        return None, "no active window"
                    bbox = (int(win.left), int(win.top), int(win.left + win.width), int(win.top + win.height))
                    img = ImageGrab.grab(bbox=bbox)
                    out_path = get_media_dir() / f"ocr_window_{_timestamp_id()}.png"
                    img.save(out_path)
                    return out_path, None

                if mode_name in {"ocr_region", "extract_text_region"}:
                    rx = int(x if x is not None else 0)
                    ry = int(y if y is not None else 0)
                    rw = max(10, int(width if width is not None else 400))
                    rh = max(10, int(height if height is not None else 280))
                    img = ImageGrab.grab(bbox=(rx, ry, rx + rw, ry + rh))
                    out_path = get_media_dir() / f"ocr_region_{_timestamp_id()}.png"
                    img.save(out_path)
                    return out_path, None

                img = ImageGrab.grab()
                out_path = get_media_dir() / f"ocr_screen_{_timestamp_id()}.png"
                img.save(out_path)
                return out_path, None
            except Exception as exc:
                return None, str(exc)

        def _describe_with_cloud_vision(img_path: Path) -> tuple[bool, dict[str, Any], str | None]:
            settings = get_settings()
            provider = str(getattr(settings, "vision_provider", "auto") or "auto").strip().lower()
            prompt = (
                "Analyze this Windows desktop screenshot and return STRICT JSON with keys: "
                "ui_summary (short string), top_app (string), detected_elements (array of short strings), confidence (0..1). "
                "No markdown."
            )
            image_bytes = img_path.read_bytes()
            b64 = base64.b64encode(image_bytes).decode("ascii")

            def _parse_json_text(raw_text: str) -> dict[str, Any]:
                text = str(raw_text or "").strip()
                if text.startswith("```"):
                    text = re.sub(r"^```(?:json)?", "", text, flags=re.IGNORECASE).strip()
                    text = re.sub(r"```$", "", text).strip()
                parsed = json.loads(text)
                if not isinstance(parsed, dict):
                    return {}
                return parsed

            if provider in {"openai", "auto"}:
                api_key = str(getattr(settings, "openai_api_key", "") or "").strip()
                if api_key:
                    try:
                        import httpx

                        model = str(getattr(settings, "vision_model", "gpt-4o") or "gpt-4o")
                        payload = {
                            "model": model,
                            "messages": [
                                {
                                    "role": "user",
                                    "content": [
                                        {"type": "text", "text": prompt},
                                        {
                                            "type": "image_url",
                                            "image_url": {"url": f"data:image/png;base64,{b64}"},
                                        },
                                    ],
                                }
                            ],
                            "temperature": 0.1,
                            "max_tokens": 500,
                        }
                        with httpx.Client(timeout=45.0) as client:
                            resp = client.post(
                                "https://api.openai.com/v1/chat/completions",
                                headers={
                                    "Authorization": f"Bearer {api_key}",
                                    "Content-Type": "application/json",
                                },
                                json=payload,
                            )
                            resp.raise_for_status()
                        data = resp.json()
                        text_out = (
                            (((data or {}).get("choices") or [{}])[0].get("message") or {}).get("content", "")
                        )
                        parsed = _parse_json_text(str(text_out))
                        if parsed:
                            return True, parsed, "openai"
                    except Exception:
                        pass

            if provider in {"gemini", "auto"}:
                api_key = str(getattr(settings, "google_api_key", "") or "").strip()
                if api_key:
                    try:
                        import httpx

                        model = str(getattr(settings, "vision_model", "gemini-2.5-flash") or "gemini-2.5-flash")
                        endpoint = (
                            "https://generativelanguage.googleapis.com/v1beta/models/"
                            f"{model}:generateContent?key={api_key}"
                        )
                        payload = {
                            "contents": [
                                {
                                    "parts": [
                                        {"text": prompt},
                                        {
                                            "inline_data": {
                                                "mime_type": "image/png",
                                                "data": b64,
                                            }
                                        },
                                    ]
                                }
                            ],
                            "generationConfig": {"temperature": 0.1, "maxOutputTokens": 500},
                        }
                        with httpx.Client(timeout=45.0) as client:
                            resp = client.post(endpoint, json=payload)
                            resp.raise_for_status()
                        data = resp.json()
                        candidates = (data or {}).get("candidates") or []
                        parts = (((candidates[0] if candidates else {}).get("content") or {}).get("parts") or [])
                        text_out = ""
                        for p in parts:
                            if isinstance(p, dict) and p.get("text"):
                                text_out += str(p.get("text"))
                        parsed = _parse_json_text(text_out)
                        if parsed:
                            return True, parsed, "gemini"
                    except Exception:
                        pass

            return False, {}, None

        if mode_norm in {"describe_screen"}:
            image_path, err = _capture("ocr_screen")
            if image_path is None:
                return self._error(err or "screen capture failed")
            ok_vision, vision_data, vision_engine = _describe_with_cloud_vision(image_path)
            if ok_vision:
                ui_summary = str(vision_data.get("ui_summary") or "").strip()
                top_app = str(vision_data.get("top_app") or "").strip()
                detected = vision_data.get("detected_elements")
                if not isinstance(detected, list):
                    detected = []
                confidence = vision_data.get("confidence")
                return _json(
                    {
                        "ok": True,
                        "mode": mode_norm,
                        "screenshot": str(image_path),
                        "source": "vision",
                        "vision_engine": vision_engine,
                        "ui_summary": ui_summary,
                        "top_app": top_app,
                        "detected_elements": [str(x) for x in detected[:20]],
                        "confidence": confidence,
                    }
                )
            if not bool(getattr(get_settings(), "vision_fallback_ocr_enabled", True)):
                return self._error("vision analysis failed and OCR fallback is disabled")
            ok_ocr, text_out, engine = _ocr_image(image_path)
            windows = _serialize_windows(include_untitled=False, limit=25)
            return _json(
                {
                    "ok": True,
                    "mode": mode_norm,
                    "screenshot": str(image_path),
                    "source": "ocr_fallback",
                    "window_count": len(windows),
                    "top_windows": [str(w.get("title", "") or "") for w in windows[:8]],
                    "ocr_engine": engine if ok_ocr else None,
                    "ocr_text_preview": text_out[:800] if ok_ocr else "",
                    "ocr_error": None if ok_ocr else engine,
                }
            )

        if mode_norm in {
            "ocr_screen",
            "extract_text_screen",
            "ocr_image",
            "extract_text_image",
            "ocr_active_window",
            "extract_text_active_window",
            "ocr_region",
            "extract_text_region",
            "copy_ocr_to_clipboard",
        }:
            capture_mode = "ocr_screen"
            if mode_norm in {"ocr_image", "extract_text_image"}:
                capture_mode = "ocr_image"
            elif mode_norm in {"ocr_active_window", "extract_text_active_window"}:
                capture_mode = "ocr_active_window"
            elif mode_norm in {"ocr_region", "extract_text_region"}:
                capture_mode = "ocr_region"
            image_path, err = _capture(capture_mode)
            if image_path is None:
                return self._error(err or "capture failed")
            ok_ocr, text_out, engine = _ocr_image(image_path)
            if not ok_ocr:
                return self._error(engine)
            if mode_norm == "copy_ocr_to_clipboard":
                tmp = get_media_dir() / f"ocr_text_{_timestamp_id()}.txt"
                tmp.write_text(text_out, encoding="utf-8", errors="ignore")
                ok, out = _run_powershell(f"Get-Content -Raw '{tmp}' | Set-Clipboard", timeout=8)
                if not ok:
                    return self._error(out or "failed to copy OCR text to clipboard")
                return _json(
                    {
                        "ok": True,
                        "mode": mode_norm,
                        "copied": True,
                        "chars": len(text_out),
                        "image_path": str(image_path),
                    }
                )
            return _json(
                {
                    "ok": True,
                    "mode": mode_norm,
                    "image_path": str(image_path),
                    "ocr_engine": engine,
                    "chars": len(text_out),
                    "text": text_out,
                }
            )

        return self._error(f"unsupported vision_tools mode: {mode_norm}")

    def _threat_tools(
        self,
        mode: str,
        path: str = "",
        target: str = "",
        max_results: int = 50,
    ) -> str:
        mode_norm = (mode or "").strip().lower()
        max_results = _clamp(max_results, 1, 200)

        if mode_norm in {"file_hash", "sha256"}:
            src = Path(path or target).expanduser()
            if not src.exists() or not src.is_file():
                return self._error(f"path not found: {src}")
            h = hashlib.sha256()
            with src.open("rb") as f:
                for chunk in iter(lambda: f.read(1024 * 1024), b""):
                    h.update(chunk)
            return _json({"ok": True, "mode": mode_norm, "path": str(src), "sha256": h.hexdigest()})

        if mode_norm in {"vt_lookup", "virustotal_lookup"}:
            api_key = (os.getenv("VIRUSTOTAL_API_KEY") or "").strip()
            if not api_key:
                return self._error("VIRUSTOTAL_API_KEY is not set")
            hash_value = (target or "").strip().lower()
            if not hash_value:
                src = Path(path).expanduser()
                if not src.exists() or not src.is_file():
                    return self._error("provide file hash in target or file path in path")
                h = hashlib.sha256()
                with src.open("rb") as f:
                    for chunk in iter(lambda: f.read(1024 * 1024), b""):
                        h.update(chunk)
                hash_value = h.hexdigest()
            try:
                import urllib.error
                import urllib.request

                req = urllib.request.Request(
                    f"https://www.virustotal.com/api/v3/files/{hash_value}",
                    headers={"x-apikey": api_key},
                    method="GET",
                )
                with urllib.request.urlopen(req, timeout=20) as resp:
                    payload = json.loads(resp.read().decode("utf-8", errors="ignore"))
                attrs = (((payload or {}).get("data") or {}).get("attributes") or {})
                stats = attrs.get("last_analysis_stats") or {}
                return _json(
                    {
                        "ok": True,
                        "mode": mode_norm,
                        "hash": hash_value,
                        "last_analysis_stats": stats,
                        "reputation": attrs.get("reputation"),
                        "meaningful_name": attrs.get("meaningful_name"),
                    }
                )
            except urllib.error.HTTPError as exc:
                body = ""
                try:
                    body = exc.read().decode("utf-8", errors="ignore")
                except Exception:
                    body = str(exc)
                return self._error(f"virustotal http error: {exc.code} {body[:300]}")
            except Exception as exc:
                return self._error(f"virustotal lookup failed: {exc}")

        if mode_norm in {"external_ips", "public_connections"}:
            try:
                import psutil
            except Exception as exc:
                return self._error(f"psutil unavailable: {exc}")
            items: list[dict[str, Any]] = []
            for conn in psutil.net_connections(kind="inet"):
                try:
                    if str(conn.status) != "ESTABLISHED" or not conn.raddr:
                        continue
                    remote_ip = str(conn.raddr.ip)
                    ip_obj = ipaddress.ip_address(remote_ip)
                    if ip_obj.is_private or ip_obj.is_loopback or ip_obj.is_link_local:
                        continue
                    items.append(
                        {
                            "pid": int(conn.pid or 0),
                            "remote_ip": remote_ip,
                            "remote_port": int(conn.raddr.port or 0),
                            "local_ip": str(conn.laddr.ip),
                            "local_port": int(conn.laddr.port or 0),
                        }
                    )
                except Exception:
                    continue
            return _json({"ok": True, "mode": mode_norm, "count": len(items), "items": items[:max_results]})

        if mode_norm in {"suspicious_connections", "suspicious_apps"}:
            try:
                import psutil
            except Exception as exc:
                return self._error(f"psutil unavailable: {exc}")
            by_pid: dict[int, dict[str, Any]] = {}
            for conn in psutil.net_connections(kind="inet"):
                try:
                    if str(conn.status) != "ESTABLISHED" or not conn.raddr:
                        continue
                    remote_ip = str(conn.raddr.ip)
                    ip_obj = ipaddress.ip_address(remote_ip)
                    if ip_obj.is_private or ip_obj.is_loopback or ip_obj.is_link_local:
                        continue
                    pid = int(conn.pid or 0)
                    if pid <= 0:
                        continue
                    slot = by_pid.setdefault(pid, {"remote_ips": set(), "connections": 0})
                    slot["connections"] = int(slot["connections"]) + 1
                    slot["remote_ips"].add(remote_ip)
                except Exception:
                    continue

            rows: list[dict[str, Any]] = []
            for pid, data in by_pid.items():
                try:
                    proc = psutil.Process(pid)
                    pname = str(proc.name() or "")
                    ppath = str(proc.exe() or "")
                    conn_count = int(data.get("connections", 0))
                    remote_ips = sorted(list(data.get("remote_ips", set())))
                    score = 0
                    reasons: list[str] = []
                    if conn_count >= 12:
                        score += 4
                        reasons.append("many_public_connections")
                    elif conn_count >= 6:
                        score += 2
                        reasons.append("multiple_public_connections")
                    low_path = ppath.lower()
                    if "\\appdata\\local\\temp" in low_path or low_path.endswith(".tmp"):
                        score += 3
                        reasons.append("running_from_temp_path")
                    if not ppath:
                        score += 1
                        reasons.append("unknown_executable_path")
                    if score <= 0:
                        continue
                    rows.append(
                        {
                            "pid": pid,
                            "name": pname,
                            "path": ppath,
                            "public_connection_count": conn_count,
                            "public_remote_ips": remote_ips[:10],
                            "suspicion_score": score,
                            "reasons": reasons,
                        }
                    )
                except Exception:
                    continue
            rows.sort(key=lambda x: int(x.get("suspicion_score", 0)), reverse=True)
            return _json({"ok": True, "mode": mode_norm, "count": len(rows), "items": rows[:max_results]})

        if mode_norm in {"behavior_scan", "suspicious_behavior", "ai_security_scan"}:
            suspicious_raw = self._threat_tools("suspicious_connections", max_results=max_results)
            if isinstance(suspicious_raw, str) and suspicious_raw.lower().startswith("error:"):
                return suspicious_raw
            startup_raw = self._startup_tools("signature_check")
            failed_raw = self._security_tools("failed_audit_logins")

            suspicious_items: list[dict[str, Any]] = []
            startup_items: list[dict[str, Any]] = []
            failed_items: list[dict[str, Any]] = []

            try:
                suspicious_data = json.loads(suspicious_raw)
                suspicious_items = list(suspicious_data.get("items") or [])
            except Exception:
                suspicious_items = []
            try:
                startup_data = json.loads(startup_raw)
                startup_items = list(startup_data.get("items") or [])
            except Exception:
                startup_items = []
            try:
                failed_data = json.loads(failed_raw)
                failed_items = list(failed_data.get("items") or [])
            except Exception:
                failed_items = []

            unsigned_startups = []
            for item in startup_items:
                try:
                    status = str(item.get("signature_status") or "").strip().lower()
                    if status not in {"valid", "unknown"}:
                        unsigned_startups.append(item)
                except Exception:
                    continue

            risk_score = 0
            if suspicious_items:
                risk_score += min(6, len(suspicious_items))
            if unsigned_startups:
                risk_score += min(3, len(unsigned_startups))
            if len(failed_items) >= 15:
                risk_score += 2
            elif len(failed_items) >= 5:
                risk_score += 1

            if risk_score >= 7:
                risk_level = "high"
            elif risk_score >= 4:
                risk_level = "medium"
            else:
                risk_level = "low"

            recommendations: list[str] = []
            if suspicious_items:
                recommendations.append("Review suspicious public network connections and kill unknown processes.")
            if unsigned_startups:
                recommendations.append("Disable unsigned or error-signature startup items and verify file reputation.")
            if len(failed_items) >= 5:
                recommendations.append("Inspect Security logs for repeated failed logons and rotate credentials if needed.")
            if not recommendations:
                recommendations.append("No major anomalies detected; keep periodic monitoring enabled.")

            return _json(
                {
                    "ok": True,
                    "mode": mode_norm,
                    "risk_level": risk_level,
                    "risk_score": risk_score,
                    "suspicious_process_count": len(suspicious_items),
                    "startup_signature_issues": len(unsigned_startups),
                    "failed_login_events": len(failed_items),
                    "recommendations": recommendations,
                    "top_suspicious_processes": suspicious_items[:10],
                    "startup_issues": unsigned_startups[:20],
                }
            )

        return self._error(f"unsupported threat_tools mode: {mode_norm}")

    def _content_tools(self, mode: str, content: str = "", path: str = "", target: str = "") -> str:
        mode_norm = (mode or "").strip().lower()
        payload = (content or "").strip()

        def _prepare_output_path(default_name: str, supplied: str = "") -> Path:
            candidate = Path(supplied).expanduser() if supplied else (get_media_dir() / default_name)
            try:
                candidate.parent.mkdir(parents=True, exist_ok=True)
                probe = candidate.parent / f".mudabbir_write_probe_{_timestamp_id()}.tmp"
                probe.write_text("ok", encoding="utf-8")
                probe.unlink(missing_ok=True)
                return candidate
            except Exception:
                fallback = Path.cwd() / default_name
                fallback.parent.mkdir(parents=True, exist_ok=True)
                return fallback

        if mode_norm in {"draft_reply", "email_draft"}:
            if not payload:
                return self._error("content is required")
            arabic = _contains_arabic(payload)
            if arabic:
                draft = (
                    "السادة الكرام،\n\n"
                    "شكرًا لتواصلكم. اطلعت على رسالتكم، وأؤكد أنني أعمل على الموضوع وسأوافيكم بتحديث واضح خلال أقرب وقت.\n\n"
                    "إذا كان هناك أي تفاصيل إضافية مطلوبة، يرجى مشاركتها حتى أضمن لكم معالجة أدق وأسرع.\n\n"
                    "مع خالص التقدير."
                )
            else:
                draft = (
                    "Dear Team,\n\n"
                    "Thank you for your message. I have reviewed your request and I am currently working on it.\n\n"
                    "I will share a clear update as soon as possible. If there are any additional details you want me to include, please let me know.\n\n"
                    "Best regards."
                )
            out_path = _prepare_output_path(f"reply_draft_{_timestamp_id()}.txt", supplied=path)
            out_path.write_text(draft, encoding="utf-8")
            return _json({"ok": True, "mode": mode_norm, "path": str(out_path), "text": draft})

        if mode_norm in {"email_auto_reply_docx", "auto_reply_word"}:
            if not payload:
                return self._error("content is required")
            arabic = _contains_arabic(payload)
            if arabic:
                draft = (
                    "الموضوع: رد على رسالتكم\n\n"
                    "السادة الكرام،\n\n"
                    "شكرًا على رسالتكم. تم استلام طلبكم ومراجعته، ونحن نعمل حاليًا على تنفيذه.\n\n"
                    "سنوافيكم بتحديث واضح يتضمن الحالة الحالية والخطوات التالية في أقرب وقت ممكن.\n\n"
                    "إذا كانت هناك تفاصيل إضافية أو مرفقات داعمة، يرجى إرسالها حتى نضمن معالجة أسرع وأدق.\n\n"
                    "مع خالص التقدير."
                )
            else:
                draft = (
                    "Subject: Re: Your Request\n\n"
                    "Dear Team,\n\n"
                    "Thank you for your message. We have received and reviewed your request, and we are currently working on it.\n\n"
                    "We will share a clear status update and next steps as soon as possible.\n\n"
                    "If there are any additional details or attachments you would like us to consider, please send them.\n\n"
                    "Best regards."
                )
            out_path = _prepare_output_path(f"email_reply_{_timestamp_id()}.docx", supplied=path)
            try:
                from docx import Document

                doc = Document()
                for paragraph in draft.splitlines():
                    doc.add_paragraph(paragraph)
                doc.save(out_path)
                return _json({"ok": True, "mode": mode_norm, "path": str(out_path), "format": "docx", "text": draft})
            except Exception:
                txt_path = out_path.with_suffix(".txt")
                txt_path.write_text(draft, encoding="utf-8")
                return _json({"ok": True, "mode": mode_norm, "path": str(txt_path), "format": "txt_fallback", "text": draft})

        if mode_norm in {"draft_to_word", "save_word_draft"}:
            if not payload:
                return self._error("content is required")
            out_path = _prepare_output_path(f"draft_{_timestamp_id()}.docx", supplied=path)
            try:
                from docx import Document

                doc = Document()
                for paragraph in payload.splitlines() or [payload]:
                    doc.add_paragraph(paragraph)
                doc.save(out_path)
                return _json({"ok": True, "mode": mode_norm, "path": str(out_path), "format": "docx"})
            except Exception:
                txt_path = out_path.with_suffix(".txt")
                txt_path.write_text(payload, encoding="utf-8")
                return _json({"ok": True, "mode": mode_norm, "path": str(txt_path), "format": "txt_fallback"})

        if mode_norm in {"text_numbers_to_excel", "text_to_excel"}:
            text_source = payload
            if not text_source and path:
                src = Path(path).expanduser()
                if not src.exists():
                    return self._error(f"path not found: {src}")
                text_source = src.read_text(encoding="utf-8", errors="ignore")
            if not text_source:
                return self._error("content or path is required")
            numbers = re.findall(r"-?\d+(?:\.\d+)?", text_source)
            if not numbers:
                return self._error("no numeric values found in input text")

            out_xlsx = _prepare_output_path(f"numbers_{_timestamp_id()}.xlsx", supplied=target)
            try:
                from openpyxl import Workbook

                wb = Workbook()
                ws = wb.active
                ws.title = "Numbers"
                ws["A1"] = "Index"
                ws["B1"] = "Value"
                for i, n in enumerate(numbers, start=1):
                    ws.cell(row=i + 1, column=1, value=i)
                    try:
                        value = float(n)
                    except Exception:
                        value = n
                    ws.cell(row=i + 1, column=2, value=value)
                wb.save(out_xlsx)
                return _json({"ok": True, "mode": mode_norm, "path": str(out_xlsx), "count": len(numbers)})
            except Exception:
                out_csv = out_xlsx.with_suffix(".csv")
                lines = ["Index,Value"]
                for i, n in enumerate(numbers, start=1):
                    lines.append(f"{i},{n}")
                out_csv.write_text("\n".join(lines) + "\n", encoding="utf-8")
                return _json({"ok": True, "mode": mode_norm, "path": str(out_csv), "count": len(numbers), "format": "csv_fallback"})

        return self._error(f"unsupported content_tools mode: {mode_norm}")


class ScreenshotTool(BaseTool):
    @property
    def name(self) -> str:
        return "screenshot"

    @property
    def description(self) -> str:
        return "Capture one screenshot of the current desktop."

    async def execute(self, **params: Any) -> str:
        return await DesktopTool().execute(action="screen_snapshot")


class StatusTool(BaseTool):
    @property
    def name(self) -> str:
        return "desktop_status"

    @property
    def description(self) -> str:
        return "Return a lightweight desktop overview."

    async def execute(self, **params: Any) -> str:
        return await DesktopTool().execute(action="desktop_overview")
